#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Module de traitement de texte pour Vynal Docs Automator
Fournit des fonctions pour le prétraitement, le nettoyage et l'analyse du texte
des documents avant extraction des données.
"""

import re
import unicodedata
import logging
from typing import List, Dict, Any, Tuple, Optional, Union
import os

# Configuration du logger
logger = logging.getLogger("Vynal Docs Automator.doc_analyzer.utils.text_processor")

# Importer PyPDF2 pour l'extraction de texte des PDF
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logging.warning("PyPDF2 n'est pas installé. L'extraction de texte des PDF pourrait être limitée.")

# Importer python-docx pour l'extraction de texte des DOCX
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx n'est pas installé. L'extraction de texte des DOCX ne sera pas possible.")

# Importer les modules OCR si disponibles
try:
    from .ocr import extract_text_from_image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("Module OCR non disponible. L'extraction de texte des images ne sera pas possible.")

def clean_text(text: str) -> str:
    """
    Nettoie le texte brut en supprimant les caractères non imprimables,
    en normalisant les espaces et en corrigeant les problèmes courants d'OCR.
    
    Args:
        text (str): Texte brut à nettoyer
        
    Returns:
        str: Texte nettoyé
    """
    return TextProcessor.clean_text(text)

def preprocess_text(text: str) -> str:
    """
    Prétraite le texte pour l'analyse.
    Nettoie et normalise le texte pour faciliter l'extraction d'informations.
    
    Args:
        text (str): Texte à prétraiter
        
    Returns:
        str: Texte prétraité
    """
    return TextProcessor.preprocess_text(text)

class TextProcessor:
    """
    Classe de traitement de texte pour l'analyse de documents.
    Fournit des méthodes pour le nettoyage, le prétraitement et l'analyse de texte.
    """
    
    def __init__(self):
        """Initialise le processeur de texte."""
        logger.info("Initialisation du TextProcessor")
    
    @staticmethod
    def clean_text(text: str) -> str:
        """
        Nettoie le texte brut en supprimant les caractères non imprimables,
        en normalisant les espaces et en corrigeant les problèmes courants d'OCR.
        
        Args:
            text (str): Texte brut à nettoyer
            
        Returns:
            str: Texte nettoyé
        """
        if not text:
            return ""
        
        # Conserver le texte original pour la comparaison
        original = text
        
        # Normalisation Unicode (conversion des caractères spéciaux)
        text = unicodedata.normalize('NFKC', text)
        
        # Suppression des caractères de contrôle (sauf sauts de ligne et tabulations)
        text = ''.join(ch for ch in text if unicodedata.category(ch)[0] != 'C' 
                      or ch in ['\n', '\t', '\r'])
        
        # Normalisation des sauts de ligne
        text = re.sub(r'(\r\n|\r)', '\n', text)
        
        # Correction des problèmes courants d'OCR
        ocr_corrections = {
            r'l\s*["\']': "d'",      # l'homme -> d'homme
            r'0(\d)': r'O\1',        # 01 -> O1 (pour les identifiants)
            r'(\d)0': r'\1O',        # 10 -> 1O (pour les identifiants)
            r'l(\d)': r'1\1',        # l2 -> 12
            r'(\d)l': r'\1l',        # 2l -> 2l
            'rn': 'm',               # rn -> m
            'vv': 'w',               # vv -> w
            r'\bI([A-Z])': r'1\1',   # IEAN -> 1EAN
            r'5(\d)': r'S\1',        # 51 -> S1 (dans certains codes)
            '[\u2018\u2019]': "'",   # Apostrophes typographiques -> apostrophe simple
            '[\u201C\u201D]': '"',   # Guillemets typographiques -> guillemets simples
            '·': '.',                # Point médian -> point
            '≤': '<=',               # Remplacement des symboles mathématiques
            '≥': '>=',
            '—': '-',                # Tiret cadratin -> tiret simple
            '–': '-',                # Tiret demi-cadratin -> tiret simple
        }
        
        for pattern, replacement in ocr_corrections.items():
            text = re.sub(pattern, replacement, text)
        
        # Normalisation des espaces (pas plus d'un espace consécutif)
        text = re.sub(r' +', ' ', text)
        
        # Suppression des espaces en début et fin de ligne
        text = re.sub(r'^ +| +$', '', text, flags=re.MULTILINE)
        
        # Suppression des lignes vides multiples (max 2 sauts de ligne consécutifs)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Log du résultat si le texte a changé significativement
        if len(original) > 0 and abs(len(text) - len(original)) / len(original) > 0.1:
            logger.info(f"Nettoyage du texte: longueur modifiée de {len(original)} à {len(text)} caractères")
        
        return text

    @staticmethod
    def preprocess_text(text: str) -> str:
        """
        Prétraite le texte pour l'analyse en effectuant des opérations comme
        la correction d'espacement autour de la ponctuation, la normalisation
        des dates, des montants, etc.
        
        Args:
            text: Texte à prétraiter
            
        Returns:
            str: Texte prétraité optimisé pour l'extraction
        """
        if not text:
            return ""
        
        # Supprimer les mentions de numéros de page
        page_patterns = [
            r'(?i)page\s*\d+(?:\s*(?:sur|of|de)\s*\d+)?',
            r'(?i)page\s*\d+/\d+',
            r'(?i)-\s*\d+\s*-',  # Format courant pour les numéros de page centrés
            r'(?i)^\s*\d+\s*$'   # Numéro de page seul sur une ligne
        ]
        
        for pattern in page_patterns:
            text = re.sub(pattern, '', text)
        
        # Nettoyer les lignes vides multiples résultant de la suppression des numéros de page
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Pas plus de 2 sauts de ligne consécutifs
        text = re.sub(r'^\s*\n', '', text)  # Supprimer les sauts de ligne au début
        text = re.sub(r'\n\s*$', '', text)  # Supprimer les sauts de ligne à la fin
        
        # Normalisation des montants
        # Regrouper les chiffres séparés par des espaces (10 000 -> 10000)
        # mais uniquement dans les contextes de montants monétaires
        money_patterns = [
            (r'(\d{1,3}(?:\s\d{3})+)(?=\s*(?:€|EUR|euros?|F\s*CFA|XOF|MAD|DZD|XAF|FCFA))', 
             lambda m: m.group(1).replace(' ', '')),
            
            # Format des montants avec virgule/point
            (r'(\d+)[,\.](\d{1,2})(?=\s*(?:€|EUR|euros?|F\s*CFA|XOF|MAD|DZD|XAF|FCFA))', 
             r'\1.\2')
        ]
        
        for pattern, replacement in money_patterns:
            if callable(replacement):
                # Si le remplacement est une fonction lambda
                text = re.sub(pattern, replacement, text)
            else:
                # Si le remplacement est une chaîne
                text = re.sub(pattern, replacement, text)
        
        # Normalisation des dates
        date_patterns = [
            # Format JJ/MM/AAAA ou JJ-MM-AAAA avec année à 2 ou 4 chiffres
            (r'(\d{1,2})[-/\.](\d{1,2})[-/\.](\d{2})(?!\d)', 
             lambda m: f"{m.group(1).zfill(2)}/{m.group(2).zfill(2)}/20{m.group(3)}" 
             if int(m.group(3)) < 50 else f"{m.group(1).zfill(2)}/{m.group(2).zfill(2)}/19{m.group(3)}"),
            
            # Format JJ/MM/AAAA ou JJ-MM-AAAA avec année à 4 chiffres
            (r'(\d{1,2})[-/\.](\d{1,2})[-/\.](\d{4})', 
             lambda m: f"{m.group(1).zfill(2)}/{m.group(2).zfill(2)}/{m.group(3)}"),
            
            # Date écrite avec mois en lettres (15 janvier 2023)
            (r'(\d{1,2})\s+(janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre)\s+(\d{4})',
             lambda m: f"{m.group(1).zfill(2)}/{TextProcessor.month_to_number(m.group(2))}/{m.group(3)}")
        ]
        
        for pattern, replacement in date_patterns:
            if callable(replacement):
                text = re.sub(pattern, replacement, text)
            else:
                text = re.sub(pattern, replacement, text)
        
        # Normalisation des numéros de téléphone français
        # +33 6 12 34 56 78 ou 06 12 34 56 78 -> +33612345678
        phone_pattern = r'(?:\+33|0)\s*[1-9](?:[\s.-]?\d{2}){4}'
        
        for match in re.finditer(phone_pattern, text):
            original = match.group(0)
            # Nettoyer le numéro (supprimer espaces, tirets, points)
            cleaned = re.sub(r'[\s.-]', '', original)
            # Convertir format national en international
            if cleaned.startswith('0'):
                cleaned = '+33' + cleaned[1:]
            # Remplacer dans le texte
            text = text.replace(original, cleaned)
        
        # Simplification des numéros SIRET/SIREN (supprimer espaces)
        siret_pattern = r'(?:SIRET|SIREN)\s*:?\s*(\d{3}\s*\d{3}\s*\d{3}\s*\d{5}|\d{3}\s*\d{3}\s*\d{3})'
        
        for match in re.finditer(siret_pattern, text, re.IGNORECASE):
            original = match.group(0)
            # Garder "SIRET:" ou "SIREN:" et nettoyer le numéro
            prefix = "SIRET:" if "SIRET" in original.upper() else "SIREN:"
            number_part = re.sub(r'(?:SIRET|SIREN)\s*:?\s*', '', original, flags=re.IGNORECASE)
            cleaned = prefix + re.sub(r'\s', '', number_part)
            # Remplacer dans le texte
            text = text.replace(original, cleaned)
        
        # Ajout d'espaces après la ponctuation si manquants
        punctuation_fixes = [
            (r'([.,:;!?])([^\s\d])', r'\1 \2'),  # Ajouter espace après ponctuation
            (r'([^\s])([\[(])', r'\1 \2'),       # Ajouter espace avant parenthèse ouvrante
            (r'([\])])([^\s,.:;!?])', r'\1 \2')  # Ajouter espace après parenthèse fermante
        ]
        
        for pattern, replacement in punctuation_fixes:
            text = re.sub(pattern, replacement, text)
        
        # Normalisation du numéro de TVA intracommunautaire
        tva_intra_pattern = r'(?:TVA\s+intracommunautaire|N°\s*TVA|VAT\s+Number)\s*:?\s*([A-Z]{2}[0-9A-Z\s]+)'
        
        for match in re.finditer(tva_intra_pattern, text, re.IGNORECASE):
            original = match.group(0)
            prefix_match = re.search(r'(TVA\s+intracommunautaire|N°\s*TVA|VAT\s+Number)\s*:?', original, re.IGNORECASE)
            prefix = prefix_match.group(0) if prefix_match else ""
            number_part = match.group(1)
            # Nettoyer le numéro de TVA (supprimer espaces)
            cleaned = prefix + " " + re.sub(r'\s', '', number_part)
            # Remplacer dans le texte
            text = text.replace(original, cleaned)
        
        # Normalisation des mots-clés et labels importants pour l'extraction
        # Standardiser les variantes d'orthographe
        keyword_standardization = {
            r'\bno(?:\.|°)\s': 'numéro ',
            r'\bn(?:\.|°)\s': 'numéro ',
            r'\bref(?:\.|°)?\s': 'référence ',
            r'\b(?:tel|tél)(?:\.|°)?\s': 'téléphone ',
            r'\bfax(?:\.|°)?\s': 'fax ',
            r'\be-?mail\b': 'email',
            r'\badrs?(?:\.|°)?\s': 'adresse ',
            r'\bcp\b': 'code postal',
            r'\bhors\s+taxe(?:s)?\b': 'HT',
            r'\btoutes\s+taxes\s+comprises\b': 'TTC',
            r'\btaxe\s+sur\s+la\s+valeur\s+ajoutée\b': 'TVA'
        }
        
        for pattern, replacement in keyword_standardization.items():
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        return text

    @staticmethod
    def segment_document(text: str) -> Dict[str, str]:
        """
        Segmente le document en sections (en-tête, corps, signature, etc.)
        pour faciliter l'extraction ciblée.
        
        Args:
            text (str): Texte complet du document
            
        Returns:
            dict: Dictionnaire des sections du document
        """
        sections = {
            'header': '',
            'body': '',
            'footer': '',
            'signature': ''
        }
        
        lines = text.split('\n')
        total_lines = len(lines)
        
        if total_lines <= 5:
            # Document trop court, tout est considéré comme corps
            sections['body'] = text
            return sections
        
        # Considérer les 5 premières lignes comme en-tête
        header_lines = min(5, total_lines // 5)  # Max 20% du document ou 5 lignes
        sections['header'] = '\n'.join(lines[:header_lines])
        
        # Rechercher la section signature
        signature_keywords = ['signature', 'fait à', 'le mandataire', 'le client', 
                              'date et signature', 'lu et approuvé']
        signature_start = total_lines
        
        for i in range(total_lines - 1, total_lines - 15, -1):
            if i < 0:
                break
            
            line_lower = lines[i].lower()
            if any(keyword in line_lower for keyword in signature_keywords):
                signature_start = i
                break
        
        # Si une signature est détectée
        if signature_start < total_lines:
            sections['signature'] = '\n'.join(lines[signature_start:])
            # Le pied de page est juste avant la signature (3-5 lignes)
            footer_start = max(header_lines, signature_start - 5)
            sections['footer'] = '\n'.join(lines[footer_start:signature_start])
            # Le corps est entre l'en-tête et le pied de page
            sections['body'] = '\n'.join(lines[header_lines:footer_start])
        else:
            # Pas de signature détectée
            # Les 10% dernières lignes sont considérées comme pied de page
            footer_lines = max(2, total_lines // 10)
            footer_start = total_lines - footer_lines
            sections['footer'] = '\n'.join(lines[footer_start:])
            # Le corps est entre l'en-tête et le pied de page
            sections['body'] = '\n'.join(lines[header_lines:footer_start])
        
        return sections

    @staticmethod
    def extract_tables(text: str) -> List[List[str]]:
        """
        Extrait les tableaux potentiels du texte basé sur l'alignement et les délimiteurs.
        
        Args:
            text (str): Texte contenant potentiellement des tableaux
            
        Returns:
            list: Liste de tableaux (chaque tableau est une liste de lignes)
        """
        tables = []
        lines = text.split('\n')
        current_table = []
        in_table = False
        
        # Patterns pour détecter les délimiteurs de tableau
        table_delimiters = [
            r'[-+]{3,}',             # Lignes de séparation (----, +---+---+)
            r'[|]{1}.*[|]{1}',       # Texte entre barres verticales (|texte|texte|)
            r'^\s*\w+(?:\s{2,}\w+)+' # Mots séparés par au moins 2 espaces (alignement de colonnes)
        ]
        
        for line in lines:
            is_table_row = False
            for pattern in table_delimiters:
                if re.match(pattern, line):
                    is_table_row = True
                    break
            
            # Détection basée sur l'alignement des espaces
            if not is_table_row and '  ' in line:
                # Compter les positions des doubles espaces dans cette ligne
                space_positions = [m.start() for m in re.finditer(r'  +', line)]
                if len(space_positions) >= 2:
                    is_table_row = True
            
            # Gestion des lignes du tableau
            if is_table_row:
                if not in_table:
                    in_table = True
                    current_table = []
                current_table.append(line)
            else:
                if in_table:
                    if len(current_table) >= 2:  # Un tableau a au moins 2 lignes
                        tables.append(current_table)
                    in_table = False
                    current_table = []
        
        # Ne pas oublier le dernier tableau
        if in_table and len(current_table) >= 2:
            tables.append(current_table)
        
        return tables

    @staticmethod
    def table_to_dict(table_lines: List[str]) -> List[Dict[str, str]]:
        """
        Convertit un tableau texte en liste de dictionnaires.
        
        Args:
            table_lines (list): Liste des lignes du tableau
            
        Returns:
            list: Liste de dictionnaires (chaque dictionnaire représente une ligne du tableau)
        """
        # Supprimer les lignes de séparation et vides
        clean_lines = [line for line in table_lines if line.strip() and not re.match(r'^[-+\s]+$', line)]
        
        if not clean_lines:
            return []
        
        # Déterminer si le tableau utilise des barres verticales comme séparateurs
        uses_pipe = '|' in clean_lines[0]
        
        if uses_pipe:
            # Tableau avec des barres verticales
            headers = [h.strip() for h in clean_lines[0].split('|') if h.strip()]
            result = []
            
            for i in range(1, len(clean_lines)):
                values = [v.strip() for v in clean_lines[i].split('|') if v]
                
                if len(values) == len(headers):
                    row_dict = {headers[j]: values[j] for j in range(len(headers))}
                    result.append(row_dict)
        else:
            # Tableau avec espacement comme délimiteur
            # Détection des colonnes basée sur les espaces
            result = TextProcessor.parse_space_delimited_table(clean_lines)
        
        return result

    @staticmethod
    def parse_space_delimited_table(lines: List[str]) -> List[Dict[str, str]]:
        """
        Parse un tableau délimité par des espaces.
        
        Args:
            lines (list): Lignes du tableau
            
        Returns:
            list: Liste de dictionnaires par ligne
        """
        if not lines:
            return []
        
        # Déterminer les positions de colonne
        # Nous cherchons les zones avec des espaces multiples
        first_line = lines[0]
        positions = [0]  # Début de la première colonne
        
        for match in re.finditer(r'\s{2,}', first_line):
            positions.append(match.start())
        
        positions.append(len(first_line))  # Fin de la dernière colonne
        
        # Extraire les en-têtes
        headers = []
        for i in range(len(positions) - 1):
            start = positions[i]
            end = positions[i + 1]
            if i == len(positions) - 2:  # Dernière colonne
                header = first_line[start:].strip()
            else:
                header = first_line[start:end].strip()
            if header:
                headers.append(header)
        
        # Si nous n'avons pas d'en-têtes valides, nous ne pouvons pas continuer
        if not headers:
            return []
        
        # Extraire les données pour chaque ligne
        result = []
        for i in range(1, len(lines)):
            line = lines[i]
            values = []
            
            for j in range(len(positions) - 1):
                start = positions[j]
                end = positions[j + 1] if j < len(positions) - 2 else len(line)
                
                if start < len(line):
                    if j == len(positions) - 2:  # Dernière colonne
                        value = line[start:].strip()
                    else:
                        value = line[start:end].strip()
                    values.append(value)
                else:
                    values.append("")
            
            # Vérifier que nous avons suffisamment de valeurs
            if len(values) >= len(headers):
                row_dict = {headers[j]: values[j] for j in range(len(headers))}
                result.append(row_dict)
        
        return result

    @staticmethod
    def find_document_title(text: str) -> str:
        """
        Tente d'extraire le titre du document.
        
        Args:
            text (str): Texte du document
            
        Returns:
            str: Titre du document ou chaîne vide si non trouvé
        """
        # Obtenir les 20 premières lignes non vides
        lines = [line.strip() for line in text.split('\n') if line.strip()][:20]
        
        # Critères pour un titre potentiel
        for i, line in enumerate(lines):
            # Les titres sont généralement en MAJUSCULES ou commencent par une majuscule
            # et sont relativement courts
            if len(line) <= 100 and (
                line.isupper() or  # Tout en majuscules
                (line[0].isupper() and not line.endswith(('.', ':', ',', ';'))) or  # Commence par majuscule, pas de ponctuation finale
                re.match(r'^(contrat|convention|accord|proposition|devis|facture|attestation|certificat)',
                        line, re.IGNORECASE)  # Commence par un mot-clé de document
            ):
                # Vérifier si la ligne est isolée (lignes vides avant/après)
                # ou si elle est formatée différemment
                if i == 0 or not lines[i-1].strip() or len(line) < len(lines[i-1]) // 2:
                    return line
        
        # Si aucun titre n'est trouvé, essayer les patterns spécifiques
        title_patterns = [
            r'^((?:CONTRAT|CONVENTION|ACCORD|PROPOSITION|DEVIS|FACTURE|ATTESTATION)[A-Z\s]+)',
            r'^((?:[A-Z][A-Z\s]+){3,})',  # Au moins 3 mots en majuscules
            r'^([^.\n]{5,60})\n\s*={3,}',  # Texte suivi d'une ligne de soulignement
            r'^=[^=]*=\n(.*?)\n=[^=]*='    # Texte entre lignes de =
        ]
        
        for pattern in title_patterns:
            match = re.search(pattern, text, re.MULTILINE)
            if match:
                return match.group(1).strip()
        
        return ""

    @staticmethod
    def extract_paragraphs(text: str, min_length: int = 50) -> List[str]:
        """
        Extrait les paragraphes significatifs du texte.
        
        Args:
            text (str): Texte à analyser
            min_length (int): Longueur minimale pour considérer un paragraphe
            
        Returns:
            list: Liste des paragraphes significatifs
        """
        # Diviser par lignes vides
        raw_paragraphs = re.split(r'\n\s*\n', text)
        
        # Filtrer les paragraphes trop courts ou les en-têtes/titres
        paragraphs = []
        for para in raw_paragraphs:
            # Remplacer les sauts de ligne simples par des espaces
            clean_para = re.sub(r'\n', ' ', para)
            # Normaliser les espaces multiples
            clean_para = re.sub(r'\s+', ' ', clean_para).strip()
            
            # Vérifier la longueur et exclure ce qui semble être des titres
            if len(clean_para) >= min_length and not clean_para.isupper():
                paragraphs.append(clean_para)
        
        return paragraphs

    @staticmethod
    def extract_keywords(text: str, count: int = 10) -> List[Tuple[str, int]]:
        """
        Extrait les mots-clés les plus fréquents du document, en excluant les mots courants.
        
        Args:
            text (str): Texte à analyser
            count (int): Nombre de mots-clés à extraire
            
        Returns:
            list: Liste de tuples (mot-clé, nombre d'occurrences)
        """
        # Liste de mots courants français à exclure
        stop_words = set([
            "le", "la", "les", "un", "une", "des", "et", "ou", "à", "au", "aux", "ce", "ces",
            "cette", "en", "dans", "par", "pour", "sur", "avec", "sans", "de", "du", "je", "tu",
            "il", "elle", "nous", "vous", "ils", "elles", "qui", "que", "quoi", "dont", "où",
            "est", "sont", "sera", "seront", "a", "ont", "avait", "avaient", "comme", "si",
            "plus", "moins", "très", "tout", "tous", "toute", "toutes", "autre", "autres",
            "même", "mêmes", "leur", "leurs", "mon", "ton", "son", "notre", "votre", "mais",
            "donc", "car", "ainsi", "alors", "aussi", "bien", "mal", "non", "oui", "se", "ne",
            "pas", "plus", "afin", "dont", "donc"
        ])
        
        # Nettoyer et tokeniser le texte
        clean = re.sub(r'[^\w\s]', ' ', text.lower())  # Remplacer la ponctuation par des espaces
        words = clean.split()
        
        # Filtrer les mots courts, les chiffres seuls et les mots courants
        filtered_words = [word for word in words if len(word) > 2 and not word.isdigit() and word not in stop_words]
        
        # Compter les occurrences
        word_counts = {}
        for word in filtered_words:
            word_counts[word] = word_counts.get(word, 0) + 1
        
        # Trier par fréquence et retourner les plus fréquents
        sorted_words = sorted(word_counts.items(), key=lambda x: x[1], reverse=True)
        return sorted_words[:count]

    @staticmethod
    def month_to_number(month_name: str) -> str:
        """
        Convertit un nom de mois français en numéro de mois (format 2 chiffres).
        
        Args:
            month_name (str): Nom du mois en français
            
        Returns:
            str: Numéro du mois au format 2 chiffres
        """
        months = {
            'janvier': '01', 'février': '02', 'mars': '03', 'avril': '04',
            'mai': '05', 'juin': '06', 'juillet': '07', 'août': '08',
            'septembre': '09', 'octobre': '10', 'novembre': '11', 'décembre': '12'
        }
        
        month_lower = month_name.lower()
        return months.get(month_lower, '00')

    @staticmethod
    def detect_document_language(text: str) -> str:
        """
        Détecte la langue principale du document.
        
        Args:
            text (str): Texte à analyser
            
        Returns:
            str: Code de langue ISO ('fr', 'en', 'ar', etc.)
        """
        # Compteurs de mots-clés par langue
        lang_counters = {
            'fr': 0,  # Français
            'en': 0,  # Anglais
            'ar': 0,  # Arabe
            'es': 0   # Espagnol
        }
        
        # Mots-clés spécifiques par langue
        keywords = {
            'fr': ['et', 'le', 'la', 'les', 'un', 'une', 'des', 'ce', 'cette', 'pour', 'dans',
                   'sur', 'avec', 'contrat', 'convention', 'accord', 'monsieur', 'madame', 'date'],
            'en': ['and', 'the', 'to', 'of', 'a', 'an', 'in', 'on', 'with', 'for', 'by', 'this',
                   'that', 'contract', 'agreement', 'mr', 'mrs', 'date'],
            'ar': ['و', 'في', 'من', 'على', 'إلى', 'عن', 'مع', 'هذا', 'هذه', 'التي', 'الذي',
                   'اتفاقية', 'عقد', 'السيد', 'السيدة', 'تاريخ'],
            'es': ['y', 'el', 'la', 'los', 'las', 'un', 'una', 'unos', 'unas', 'por', 'para',
                   'en', 'con', 'contrato', 'acuerdo', 'señor', 'señora', 'fecha']
        }
        
        # Convertir en minuscules et retirer la ponctuation
        text_lower = text.lower()
        for lang, words in keywords.items():
            for word in words:
                pattern = r'\b' + re.escape(word) + r'\b'
                matches = re.findall(pattern, text_lower)
                lang_counters[lang] += len(matches)
        
        # Détecter les caractères spécifiques arabes
        if re.search(r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]', text):
            lang_counters['ar'] += 50  # Bonus pour la présence de caractères arabes
        
        # Retourner la langue avec le plus grand score
        max_lang = max(lang_counters, key=lang_counters.get)
        
        # Si le score est trop faible, considérer que c'est du français par défaut
        if lang_counters[max_lang] < 5:
            return 'fr'
        
        return max_lang

    @staticmethod
    def identify_key_terms(text: str, domain: str = None) -> List[Tuple[str, str]]:
        """
        Identifie les termes techniques ou juridiques importants dans le document
        
        Args:
            text (str): Texte à analyser
            domain (str): Domaine optionnel pour cibler l'analyse (juridique, commercial, etc.)
            
        Returns:
            list: Liste de tuples (terme, contexte)
        """
        terms = []
        
        # Domaines spécifiques et leur terminologie
        domain_terms = {
            'juridique': [
                r'\b(?:clause|article)s?\s+de\s+confidentialité\b',
                r'\b(?:clause|article)s?\s+de\s+non[\s-]concurrence\b',
                r'\brésiliation\b',
                r'\bforce\s+majeure\b',
                r'\bpropriété\s+intellectuelle\b',
                r'\bdroit\s+applicable\b',
                r'\bjuridiction\s+compétente\b',
                r'\blitige\b',
                r'\bindemnités?\b',
                r'\bpréjudice\b'
            ],
            'commercial': [
                r'\bprix\b',
                r'\bmontant\b',
                r'\bdevis\b',
                r'\bfacture\b',
                r'\btarif\b',
                r'\bremise\b',
                r'\bréduction\b',
                r'\blivraison\b',
                r'\bcommande\b',
                r'\brèglement\b'
            ],
            'administratif': [
                r'\bdemande\s+d[\'e]\s*(?:autorisation|agrément)\b',
                r'\bdéclaration\b',
                r'\battestations?\b',
                r'\bformulaire\b',
                r'\bréférence\b',
                r'\bdossier\b',
                r'\bidentifiant\b',
                r'\bnuméro\s+d[\'e]\s*enregistrement\b'
            ]
        }
        
        # Si aucun domaine n'est spécifié, utiliser tous les domaines
        patterns_to_use = []
        if domain and domain in domain_terms:
            patterns_to_use = domain_terms[domain]
        else:
            for patterns in domain_terms.values():
                patterns_to_use.extend(patterns)
        
        # Rechercher les termes dans le texte
        for pattern in patterns_to_use:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                term = match.group(0)
                
                # Extraire le contexte (30 caractères avant et après)
                start = max(0, match.start() - 30)
                end = min(len(text), match.end() + 30)
                context = text[start:end].strip()
                
                # Ajouter le terme et son contexte
                terms.append((term, context))
        
        return terms

    @staticmethod
    def detect_headings(text: str) -> List[Tuple[int, str]]:
        """
        Détecte les titres et sous-titres dans un document
        
        Args:
            text (str): Texte à analyser
            
        Returns:
            list: Liste de tuples (niveau, texte du titre)
        """
        headings = []
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Patterns pour les titres numérotés
            if re.match(r'^(?:ARTICLE|Article)\s+\d+\s*[:.]\s*(.+)', line):
                # Titre d'article avec numéro (Article 1: Objet du contrat)
                match = re.match(r'^(?:ARTICLE|Article)\s+\d+\s*[:.]\s*(.+)', line)
                headings.append((1, match.group(0)))
            
            elif re.match(r'^\d+\.\s+[A-Z]', line):
                # Titre principal numéroté (1. TITRE)
                headings.append((1, line))
            
            elif re.match(r'^\d+\.\d+\.\s+[A-Z]', line):
                # Sous-titre niveau 2 (1.1. Sous-titre)
                headings.append((2, line))
            
            elif re.match(r'^\d+\.\d+\.\d+\.\s+[A-Z]', line):
                # Sous-titre niveau 3 (1.1.1. Sous-sous-titre)
                headings.append((3, line))
            
            # Patterns pour les titres non numérotés
            elif line.isupper() and len(line) < 60 and not re.match(r'^[A-Z]{1,5}\s*:', line):
                # Ligne entièrement en majuscules (TITRE)
                # Éviter les petits labels comme "NOM:" ou "TEL:"
                headings.append((1, line))
            
            elif re.match(r'^[A-Z][^.!?]*', line) and len(line) < 60:
                # Ligne commençant par une majuscule, sans ponctuation finale
                # Si la ligne précédente est vide ou la ligne suivante est vide
                if (i == 0 or not lines[i-1].strip()) and (i == len(lines)-1 or not lines[i+1].strip()):
                    headings.append((2, line))
            
            # Patterns pour les titres stylisés
            elif re.match(r'^={3,}\s*([^=]+)\s*={3,}', line):
                # Titre encadré par des = (=== Titre ===)
                match = re.match(r'^={3,}\s*([^=]+)\s*={3,}', line)
                headings.append((1, match.group(1).strip()))
            
            elif re.match(r'^-{3,}\s*([^-]+)\s*-{3,}', line):
                # Titre encadré par des - (--- Titre ---)
                match = re.match(r'^-{3,}\s*([^-]+)\s*-{3,}', line)
                headings.append((2, match.group(1).strip()))
            
            elif i > 0 and re.match(r'^[=]+', lines[i-1]) and line.strip():
                # Titre souligné par des = (Titre\n====)
                headings.append((1, line))
            
            elif i > 0 and re.match(r'^[-]+', lines[i-1]) and line.strip():
                # Titre souligné par des - (Titre\n----)
                headings.append((2, line))
        
        return headings

    @staticmethod
    def detect_form_fields(text: str) -> List[Tuple[str, str]]:
        """
        Détecte les champs de formulaire potentiels dans un document
        
        Args:
            text (str): Texte à analyser
            
        Returns:
            list: Liste de tuples (label, valeur)
        """
        fields = []
        
        # Pattern pour les champs sous forme "Label: Valeur"
        label_value_pattern = r'([A-Za-zÀ-ÖØ-öø-ÿ\s]{2,30})\s*[:]\s*([^:\n]{1,100})'
        for match in re.finditer(label_value_pattern, text):
            label = match.group(1).strip()
            value = match.group(2).strip()
            
            # Éviter les faux positifs
            if len(label) > 2 and not label.isupper() and ":" not in label and ":" not in value:
                fields.append((label, value))
        
        # Pattern pour les champs avec soulignement ou pointillés
        underline_patterns = [
            r'([A-Za-zÀ-ÖØ-öø-ÿ\s]{2,30})\s*:\s*[_]{3,}',  # Label: _____
            r'([A-Za-zÀ-ÖØ-öø-ÿ\s]{2,30})\s*:\s*[\.]{3,}'  # Label: .....
        ]
        
        for pattern in underline_patterns:
            for match in re.finditer(pattern, text):
                label = match.group(1).strip()
                
                # Éviter les faux positifs
                if len(label) > 2 and not label.isupper() and ":" not in label:
                    fields.append((label, ""))
        
        return fields

    @staticmethod
    def analyze_document_structure(text: str) -> Dict[str, Any]:
        """
        Analyse la structure globale du document
        
        Args:
            text (str): Texte à analyser
            
        Returns:
            dict: Informations sur la structure du document
        """
        structure = {
            "title": TextProcessor.find_document_title(text),
            "sections": TextProcessor.segment_document(text),
            "headings": TextProcessor.detect_headings(text),
            "paragraphs_count": len(TextProcessor.extract_paragraphs(text)),
            "tables_count": len(TextProcessor.extract_tables(text)),
            "form_fields": TextProcessor.detect_form_fields(text),
            "keywords": TextProcessor.extract_keywords(text, 15),
            "language": TextProcessor.detect_document_language(text)
        }
        
        # Estimer le type de document
        structure["document_type"] = TextProcessor.estimate_document_type(text, structure)
        
        return structure

    @staticmethod
    def estimate_document_type(text: str, structure: Dict[str, Any]) -> str:
        """
        Estime le type de document en fonction de sa structure et de son contenu
        
        Args:
            text (str): Texte du document
            structure (dict): Structure du document
            
        Returns:
            str: Type de document estimé
        """
        # Mots-clés spécifiques à chaque type de document
        document_types = {
            "contrat": ["contrat", "convention", "accord", "entre les soussignés", "s'engage à", 
                       "obligations", "signataires", "parties"],
            "facture": ["facture", "montant", "total", "ht", "tva", "ttc", "règlement", "paiement", 
                       "client", "prestataire"],
            "devis": ["devis", "estimation", "proposition", "montant", "ht", "tva", "délai", 
                     "validité", "acceptation"],
            "lettre": ["madame", "monsieur", "je vous", "veuillez", "cordialement", "salutations"],
            "rapport": ["rapport", "analyse", "résultats", "étude", "conclusion", "recommandations"],
            "procès-verbal": ["procès-verbal", "réunion", "séance", "présents", "absents", "ordre du jour"],
            "formulaire": ["formulaire", "veuillez remplir", "cochez", "nom", "prénom", "adresse"],
            "attestation": ["attestation", "certifie", "déclare", "témoigne", "fait à", "le soussigné"]
        }
        
        # Initialiser les scores
        scores = {doc_type: 0 for doc_type in document_types}
        
        # Texte en minuscules pour la recherche
        text_lower = text.lower()
        
        # 1. Analyse des mots-clés
        for doc_type, keywords in document_types.items():
            for keyword in keywords:
                if keyword in text_lower:
                    scores[doc_type] += 1
        
        # 2. Analyse de la structure
        # Si le titre contient un type de document
        title = structure.get("title", "").lower()
        for doc_type in document_types:
            if doc_type in title:
                scores[doc_type] += 5
        
        # 3. Caractéristiques spécifiques
        # Contrats ont généralement des articles
        if len([h for h in structure.get("headings", []) if "article" in h[1].lower()]) > 0:
            scores["contrat"] += 3
        
        # Factures ont généralement des tables avec des montants
        if structure.get("tables_count", 0) > 0 and any(kw[0] in ["montant", "prix", "total", "somme"] 
                                                      for kw in structure.get("keywords", [])):
            scores["facture"] += 3
        
        # Devis ont souvent une date de validité
        if "validité" in text_lower or "valable jusqu" in text_lower:
            scores["devis"] += 2
        
        # Formulaires ont généralement des champs à remplir
        if len(structure.get("form_fields", [])) > 5:
            scores["formulaire"] += 3
        
        # Chercher le type de document avec le score le plus élevé
        max_score = 0
        doc_type = "document_general"
        
        for d_type, score in scores.items():
            if score > max_score:
                max_score = score
                doc_type = d_type
        
        return doc_type

    @staticmethod
    def find_client_matches(text: str, clients_db: List[Dict[str, Any]]) -> List[Tuple[Dict[str, Any], float]]:
        """
        Recherche des correspondances potentielles avec des clients existants
        
        Args:
            text (str): Texte du document
            clients_db (list): Liste des clients existants
            
        Returns:
            list: Liste de tuples (client, score de correspondance)
        """
        matches = []
        
        # Recherche de noms, d'emails et de numéros de téléphone dans le texte
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        
        phone_pattern = r'(?:\+33|0)\s*[1-9](?:[\s.-]?\d{2}){4}'
        phones = re.findall(phone_pattern, text)
        
        # Normaliser les numéros de téléphone trouvés
        normalized_phones = []
        for phone in phones:
            clean = re.sub(r'[\s.-]', '', phone)
            if clean.startswith('0'):
                clean = '+33' + clean[1:]
            normalized_phones.append(clean)
        
        # Rechercher des correspondances pour chaque client
        for client in clients_db:
            score = 0
            
            # Correspondance par email
            client_email = client.get('email', '').lower()
            if client_email and any(email.lower() == client_email for email in emails):
                score += 5  # Score élevé pour une correspondance d'email
            
            # Correspondance par téléphone
            client_phone = client.get('phone', '')
            normalized_client_phone = re.sub(r'[\s.-]', '', client_phone)
            if normalized_client_phone.startswith('0'):
                normalized_client_phone = '+33' + normalized_client_phone[1:]
            
            if normalized_client_phone and normalized_client_phone in normalized_phones:
                score += 4  # Score élevé pour une correspondance de téléphone
            
            # Correspondance par nom
            client_name = client.get('name', '').lower()
            if client_name and client_name in text.lower():
                score += 3  # Score modéré pour une correspondance de nom
            
            # Correspondance par entreprise
            client_company = client.get('company', '').lower()
            if client_company and len(client_company) > 2 and client_company in text.lower():
                score += 2  # Score pour une correspondance d'entreprise
            
            # Si le score est suffisant, ajouter à la liste des correspondances
            if score > 0:
                matches.append((client, score))
        
        # Trier par score décroissant
        matches.sort(key=lambda x: x[1], reverse=True)
        
        return matches

    @staticmethod
    def extract_document_signatures(text: str) -> List[Dict[str, Any]]:
        """
        Extrait les informations relatives aux signatures du document
        
        Args:
            text (str): Texte du document
            
        Returns:
            list: Liste des signatures détectées
        """
        signatures = []
        
        # Rechercher les sections de signature
        signature_sections = []
        
        # 1. Recherche de la mention "Fait à ... le ..."
        fait_a_patterns = [
            r'Fait\s+à\s+([^,\n]+)[,\s]+le\s+(\d{1,2}[/.-]\d{1,2}[/.-]\d{2,4})',
            r'Fait\s+en\s+(\d+)\s+exemplaires?\s+(?:originaux\s+)?(?:à\s+([^,\n]+)[,\s]+)?le\s+(\d{1,2}[/.-]\d{1,2}[/.-]\d{2,4})'
        ]
        
        for pattern in fait_a_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                signature_sections.append(match.group(0))
        
        # 2. Recherche de la section de signatures
        signature_keywords = [
            r'(?:Pour|Le)\s+([^,\n:]+)[,\s:]+(?:représenté\s+par\s+)?([^,\n]+)(?:,\s+([^,\n]+))?',
            r'Signature\s+(?:du|de la)?\s+([^:\n]+)(?:\s*:\s*)?',
            r'(?:Lu\s+et\s+approuvé|Bon\s+pour\s+accord)[,\s.]*(?:\n|\s)+([^,\n]+)',
            r'(?:Le\s+client|Le\s+prestataire)[\s:]*([^,\n]+)'
        ]
        
        for pattern in signature_keywords:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                signature_sections.append(match.group(0))
        
        # Analyser chaque section de signature
        for section in signature_sections:
            signature = {
                "text": section,
                "signataire": None,
                "qualite": None,
                "date": None,
                "lieu": None
            }
            
            # Extraire le nom du signataire
            signataire_patterns = [
                r'(?:Pour|Le)\s+(?:[^,\n:]+)[,\s:]+(?:représenté\s+par\s+)?([^,\n]+)',
                r'Signature\s+(?:du|de la)?\s+([^:\n]+)',
                r'(?:Lu\s+et\s+approuvé|Bon\s+pour\s+accord)[,\s.]*(?:\n|\s)+([^,\n]+)',
                r'(?:Le\s+client|Le\s+prestataire)[\s:]*([^,\n]+)'
            ]
            
            for pattern in signataire_patterns:
                signataire_match = re.search(pattern, section, re.IGNORECASE)
                if signataire_match:
                    signature["signataire"] = signataire_match.group(1).strip()
                    break
            
            # Extraire la qualité/fonction du signataire
            qualite_match = re.search(r'en\s+(?:sa|ma)\s+qualité\s+de\s+([^,\n]+)', section, re.IGNORECASE)
            if qualite_match:
                signature["qualite"] = qualite_match.group(1).strip()
            
            # Extraire la date
            date_match = re.search(r'le\s+(\d{1,2}[/.-]\d{1,2}[/.-]\d{2,4})', section, re.IGNORECASE)
            if date_match:
                signature["date"] = date_match.group(1).strip()
            
            # Extraire le lieu
            lieu_match = re.search(r'(?:Fait|Signé)\s+à\s+([^,\n]+)', section, re.IGNORECASE)
            if lieu_match:
                signature["lieu"] = lieu_match.group(1).strip()
            
            signatures.append(signature)
        
        return signatures

    def analyze_text(self, text: str) -> Dict[str, Any]:
        """
        Analyse un texte pour en extraire les informations pertinentes.
        
        Args:
            text (str): Texte à analyser
            
        Returns:
            Dict[str, Any]: Résultats de l'analyse avec les champs extraits et les métadonnées
        """
        try:
            # Nettoyer et prétraiter le texte
            cleaned_text = self.clean_text(text)
            processed_text = self.preprocess_text(cleaned_text)
            
            # Segmenter le document
            sections = self.segment_document(processed_text)
            
            # Extraire les champs
            fields = {}
            
            # Extraire les tableaux
            tables = self.extract_tables(processed_text)
            if tables:
                fields['tables'] = [self.table_to_dict(table) for table in tables]
            
            # Extraire les paragraphes
            paragraphs = self.extract_paragraphs(processed_text)
            if paragraphs:
                fields['paragraphs'] = paragraphs
            
            # Extraire les mots-clés
            keywords = self.extract_keywords(processed_text)
            if keywords:
                fields['keywords'] = keywords
            
            # Détecter les en-têtes
            headings = self.detect_headings(processed_text)
            if headings:
                fields['headings'] = headings
            
            # Détecter les champs de formulaire
            form_fields = self.detect_form_fields(processed_text)
            if form_fields:
                fields['form_fields'] = form_fields
            
            # Analyser la structure du document
            structure = self.analyze_document_structure(processed_text)
            
            # Estimer le type de document
            doc_type = self.estimate_document_type(processed_text, structure)
            
            # Extraire les signatures
            signatures = self.extract_document_signatures(processed_text)
            if signatures:
                fields['signatures'] = signatures
            
            # Détecter la langue
            language = self.detect_document_language(processed_text)
            
            # Identifier les termes clés
            key_terms = self.identify_key_terms(processed_text)
            if key_terms:
                fields['key_terms'] = key_terms
            
            # Organiser les résultats
            results = {
                'fields': fields,
                'metadata': {
                    'language': language,
                    'document_type': doc_type,
                    'structure': structure,
                    'sections': sections
                }
            }
            
            logger.info("Analyse du texte terminée avec succès")
            return results
            
        except Exception as e:
            logger.error(f"Erreur lors de l'analyse du texte: {e}")
            return {
                'fields': {},
                'metadata': {
                    'error': str(e)
                }
            }

    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extrait le texte d'un fichier PDF sans utiliser l'OCR
        
        Args:
            file_path: Chemin vers le fichier PDF
            
        Returns:
            str: Texte extrait du PDF
        """
        if not PYPDF2_AVAILABLE:
            return "PyPDF2 n'est pas installé. Impossible d'extraire le texte du PDF."
        
        try:
            text = ""
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                num_pages = len(reader.pages)
                
                for page_num in range(num_pages):
                    page = reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            
            if not text.strip():
                logger.warning(f"Aucun texte extrait du PDF {file_path} avec PyPDF2")
                return "Aucun texte n'a pu être extrait de ce PDF. Le document pourrait être numérisé ou protégé."
            
            logger.info(f"Texte extrait avec succès du PDF {file_path} avec PyPDF2")
            return text
            
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction du texte du PDF {file_path}: {e}")
            return f"Erreur lors de l'extraction du texte: {str(e)}"

    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extrait le texte d'un fichier DOCX
        
        Args:
            file_path: Chemin vers le fichier DOCX
            
        Returns:
            str: Texte extrait du DOCX
        """
        if not DOCX_AVAILABLE:
            return "python-docx n'est pas installé. Impossible d'extraire le texte du DOCX."
        
        try:
            text = []
            doc = Document(file_path)
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extraire aussi le texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append("\t".join(row_text))
            
            result = "\n\n".join(text)
            
            if not result.strip():
                logger.warning(f"Aucun texte extrait du DOCX {file_path}")
                return "Aucun texte n'a pu être extrait de ce DOCX."
            
            logger.info(f"Texte extrait avec succès du DOCX {file_path}")
            return result
            
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction du texte du DOCX {file_path}: {e}")
            return f"Erreur lors de l'extraction du texte: {str(e)}"

    def process_document(self, file_path: str) -> Union[str, Dict[str, Any]]:
        """
        Prétraite un document pour en extraire le texte
        
        Args:
            file_path: Chemin vers le document à prétraiter
            
        Returns:
            Union[str, Dict[str, Any]]: Texte extrait ou dictionnaire avec erreur
        """
        if not os.path.exists(file_path):
            error_msg = f"Le fichier {file_path} n'existe pas"
            logger.error(error_msg)
            return {"error": error_msg}
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # Traitement des fichiers texte
        if file_extension in ['.txt', '.md', '.csv']:
            try:
                encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
                text = None
                
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as file:
                            text = file.read()
                            break
                    except UnicodeDecodeError:
                        continue
                
                if text is None:
                    error_msg = "Impossible de détecter l'encodage du fichier texte"
                    logger.error(error_msg)
                    return {"error": error_msg}
                
                return text
                
            except Exception as e:
                error_msg = f"Erreur lors de la lecture du fichier texte: {str(e)}"
                logger.error(error_msg)
                return {"error": error_msg}
        
        # Traitement des fichiers DOCX
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file_path)
        
        # Traitement des fichiers PDF
        elif file_extension == '.pdf':
            # Essayer d'extraire le texte directement du PDF d'abord
            text = self.extract_text_from_pdf(file_path)
            
            # Si le texte est vide ou contient un message d'erreur et que l'OCR est disponible, essayer l'OCR
            if (not text or text.startswith("Erreur") or text.startswith("Aucun texte")) and OCR_AVAILABLE:
                logger.info(f"Tentative d'extraction de texte avec OCR pour {file_path}")
                try:
                    text = extract_text_from_image(file_path)
                    if text:
                        return text
                    else:
                        error_msg = "Échec de l'extraction de texte avec OCR"
                        logger.error(error_msg)
                        return {"error": error_msg, "text": "Aucun texte n'a pu être extrait de ce document."}
                except Exception as e:
                    error_msg = f"Erreur lors de l'extraction OCR: {str(e)}"
                    logger.error(error_msg)
                    return {"error": error_msg, "text": text}
            
            return text
        
        # Traitement des images
        elif file_extension in ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp']:
            if not OCR_AVAILABLE:
                error_msg = "Module OCR non disponible. Impossible d'extraire le texte des images."
                logger.error(error_msg)
                return {"error": error_msg}
            
            try:
                text = extract_text_from_image(file_path)
                if text:
                    return text
                else:
                    error_msg = "Échec de l'extraction de texte avec OCR"
                    logger.error(error_msg)
                    return {"error": error_msg}
            except Exception as e:
                error_msg = f"Erreur lors de l'extraction OCR: {str(e)}"
                logger.error(error_msg)
                return {"error": error_msg}
        
        # Format de fichier non pris en charge
        else:
            error_msg = f"Format de fichier non pris en charge: {file_extension}"
            logger.error(error_msg)
            return {"error": error_msg}