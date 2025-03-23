#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Contrôleur de gestion des documents pour l'application Vynal Docs Automator
"""

import os
import re
import shutil
import logging
import tkinter as tk
import customtkinter as ctk
from tkinter import messagebox, filedialog, simpledialog
from datetime import datetime
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from doc_analyzer.analyzer import DocumentAnalyzer
from docx.shared import Inches

logger = logging.getLogger("VynalDocsAutomator.DocumentController")

class DialogUtils:
    """
    Utilitaires pour créer des boîtes de dialogue cohérentes dans l'application
    """
    
    @staticmethod
    def show_confirmation(parent, title, message, on_yes=None, on_no=None):
        """
        Affiche une boîte de dialogue de confirmation
        
        Args:
            parent: Widget parent
            title: Titre de la boîte de dialogue
            message: Message à afficher
            on_yes: Fonction à appeler si l'utilisateur confirme
            on_no: Fonction à appeler si l'utilisateur annule
            
        Returns:
            bool: True si confirmé, False sinon
        """
        dialog = ctk.CTkToplevel(parent)
        dialog.title(title)
        dialog.geometry("400x200")
        dialog.resizable(False, False)
        dialog.grab_set()
        dialog.focus_set()
        
        # Centrer la fenêtre
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() - dialog.winfo_width()) // 2
        y = (dialog.winfo_screenheight() - dialog.winfo_height()) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # Résultat par défaut
        result = [False]
        
        # Cadre principal
        frame = ctk.CTkFrame(dialog, fg_color="transparent")
        frame.pack(fill=ctk.BOTH, expand=True, padx=20, pady=20)
        
        # Titre avec icône
        ctk.CTkLabel(
            frame,
            text=f"⚠️ {title}",
            font=ctk.CTkFont(size=16, weight="bold")
        ).pack(pady=(0, 10))
        
        # Message
        ctk.CTkLabel(
            frame,
            text=message,
            wraplength=360
        ).pack(pady=10)
        
        # Cadre pour les boutons
        button_frame = ctk.CTkFrame(frame, fg_color="transparent")
        button_frame.pack(pady=10)
        
        # Fonctions de callback
        def yes_action():
            result[0] = True
            dialog.destroy()
            if on_yes:
                on_yes()
        
        def no_action():
            result[0] = False
            dialog.destroy()
            if on_no:
                on_no()
        
        # Bouton Non
        ctk.CTkButton(
            button_frame,
            text="Non",
            command=no_action,
            width=100,
            fg_color="#e74c3c",
            hover_color="#c0392b"
        ).pack(side=ctk.LEFT, padx=10)
        
        # Bouton Oui
        ctk.CTkButton(
            button_frame,
            text="Oui",
            command=yes_action,
            width=100,
            fg_color="#2ecc71",
            hover_color="#27ae60"
        ).pack(side=ctk.LEFT, padx=10)
        
        # Attendre que la fenêtre soit fermée
        parent.wait_window(dialog)
        
        return result[0]
    
    @staticmethod
    def show_message(parent, title, message, message_type="info"):
        """
        Affiche une boîte de dialogue avec un message
        
        Args:
            parent: Widget parent
            title: Titre de la boîte de dialogue
            message: Message à afficher
            message_type: Type de message ('info', 'error', 'warning', 'success')
        """
        dialog = ctk.CTkToplevel(parent)
        dialog.title(title)
        dialog.geometry("400x200")
        dialog.resizable(False, False)
        dialog.grab_set()
        dialog.focus_set()
        
        # Centrer la fenêtre
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() - dialog.winfo_width()) // 2
        y = (dialog.winfo_screenheight() - dialog.winfo_height()) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # Icône selon le type
        icon = "ℹ️"
        button_color = "#3498db"
        hover_color = "#2980b9"
        
        if message_type == "error":
            icon = "❌"
            button_color = "#e74c3c"
            hover_color = "#c0392b"
        elif message_type == "warning":
            icon = "⚠️"
            button_color = "#f39c12"
            hover_color = "#d35400"
        elif message_type == "success":
            icon = "✅"
            button_color = "#2ecc71"
            hover_color = "#27ae60"
        
        # Cadre principal
        frame = ctk.CTkFrame(dialog, fg_color="transparent")
        frame.pack(fill=ctk.BOTH, expand=True, padx=20, pady=20)
        
        # Titre avec icône
        ctk.CTkLabel(
            frame,
            text=f"{icon} {title}",
            font=ctk.CTkFont(size=16, weight="bold")
        ).pack(pady=(0, 10))
        
        # Message
        ctk.CTkLabel(
            frame,
            text=message,
            wraplength=360
        ).pack(pady=10)
        
        # Bouton OK
        ctk.CTkButton(
            frame,
            text="OK",
            command=dialog.destroy,
            width=100,
            fg_color=button_color,
            hover_color=hover_color
        ).pack(pady=10)

class DocumentController:
    """
    Contrôleur de gestion des documents
    Gère la logique métier liée aux documents
    """
    
    def __init__(self, app_model, document_view):
        """
        Initialise le contrôleur des documents
        
        Args:
            app_model: Modèle de l'application
            document_view: Vue de gestion des documents
        """
        self.model = app_model
        self.view = document_view
        
        # Initialiser l'analyseur de documents
        self.document_analyzer = DocumentAnalyzer()
        
        # Connecter les événements de la vue aux méthodes du contrôleur
        self.connect_events()
        
        logger.info("DocumentController initialisé")
    
    def connect_events(self):
        """
        Connecte les événements de la vue aux méthodes du contrôleur
        """
        try:
            # Connexion des événements de la vue document
            self.view.on_new_document = self.new_document
            self.view.on_open_document = self.open_document
            self.view.on_filter_documents = self.filter_documents
            
            logger.info("Événements du contrôleur de documents connectés")
        except Exception as e:
            logger.error(f"Erreur lors de la connexion des événements: {e}")
            
    def handle_auto_fill(self, template_id: str, client_id: str = None) -> bool:
        """
        Gère l'auto-remplissage d'un document
        
        Args:
            template_id: ID du modèle à utiliser
            client_id: ID du client (optionnel)
            
        Returns:
            bool: True si l'auto-remplissage a réussi, False sinon
        """
        try:
            # Récupérer le modèle
            template = self.model.get_template(template_id)
            if not template:
                DialogUtils.show_message(self.view.parent, "Erreur", "Modèle non trouvé", "error")
                return False
            
            # Récupérer le client
            client = None
            if client_id:
                client = self.model.get_client(client_id)
            else:
                # Si aucun client n'est spécifié, utiliser le client sélectionné
                selected_client = self.view.get_selected_client()
                if selected_client:
                    client = self.model.get_client(selected_client)
            
            if not client:
                DialogUtils.show_message(self.view.parent, "Attention", "Veuillez sélectionner un client", "warning")
                return False
            
            # Analyser le modèle
            template_path = os.path.join(self.model.paths['templates'], template['file_path'])
            if not os.path.exists(template_path):
                DialogUtils.show_message(self.view.parent, "Erreur", "Fichier modèle non trouvé", "error")
                return False
            
            # Analyser le document
            analysis_result = self.document_analyzer.analyze_document(template_path)
            
            if 'error' in analysis_result:
                DialogUtils.show_message(self.view.parent, "Erreur", f"Erreur lors de l'analyse: {analysis_result['error']}", "error")
                return False
            
            # Extraire les données pertinentes
            extracted_data = analysis_result.get('data', {})
            
            # Créer une boîte de dialogue pour l'auto-remplissage
            dialog = ctk.CTkToplevel(self.view.parent)
            dialog.title("Auto-remplissage")
            dialog.geometry("600x400")
            dialog.resizable(True, True)
            dialog.grab_set()
            dialog.focus_set()
            
            # Centrer la fenêtre
            dialog.update_idletasks()
            x = (dialog.winfo_screenwidth() - dialog.winfo_width()) // 2
            y = (dialog.winfo_screenheight() - dialog.winfo_height()) // 2
            dialog.geometry(f"+{x}+{y}")
            
            # Cadre principal
            main_frame = ctk.CTkFrame(dialog)
            main_frame.pack(fill=ctk.BOTH, expand=True, padx=20, pady=20)
            
            # Titre
            ctk.CTkLabel(
                main_frame,
                text="Souhaitez-vous automatiser le remplissage de ce document?",
                font=ctk.CTkFont(size=16, weight="bold")
            ).pack(pady=10)
            
            # Liste des champs extraits
            fields_frame = ctk.CTkFrame(main_frame)
            fields_frame.pack(fill=ctk.BOTH, expand=True, pady=10)
            
            # Créer une liste scrollable pour les champs
            canvas = tk.Canvas(fields_frame)
            scrollbar = ctk.CTkScrollbar(fields_frame, orientation="vertical", command=canvas.yview)
            scrollable_frame = ctk.CTkFrame(canvas)
            
            scrollable_frame.bind(
                "<Configure>",
                lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
            )
            
            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)
            
            # Afficher les champs extraits
            for section, fields in extracted_data.items():
                if isinstance(fields, dict):
                    ctk.CTkLabel(
                        scrollable_frame,
                        text=section.title(),
                        font=ctk.CTkFont(weight="bold")
                    ).pack(pady=(10, 5), padx=10, anchor="w")
                    
                    for field, value in fields.items():
                        if value:
                            ctk.CTkLabel(
                                scrollable_frame,
                                text=f"{field}: {value}",
                                wraplength=400
                            ).pack(pady=2, padx=20, anchor="w")
            
            # Pack le canvas et la scrollbar
            canvas.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")
            
            # Boutons
            button_frame = ctk.CTkFrame(main_frame)
            button_frame.pack(fill=ctk.X, pady=10)
            
            def on_confirm():
                # Créer le document avec les données extraites
                document_data = {
                    'template_id': template_id,
                    'client_id': client['id'],
                    'title': template['title'],
                    'type': template['type'],
                    'created_at': datetime.now().isoformat(),
                    'extracted_data': extracted_data
                }
                
                # Générer le document
                if self.generate_document_content(template, client, extracted_data):
                    dialog.destroy()
                    DialogUtils.show_message(self.view.parent, "Succès", "Document généré avec succès", "success")
                else:
                    dialog.destroy()
                    DialogUtils.show_message(self.view.parent, "Erreur", "Erreur lors de la génération du document", "error")
            
            def on_cancel():
                dialog.destroy()
            
            ctk.CTkButton(
                button_frame,
                text="Annuler",
                command=on_cancel,
                width=100,
                fg_color="#e74c3c",
                hover_color="#c0392b"
            ).pack(side=ctk.RIGHT, padx=10)
            
            ctk.CTkButton(
                button_frame,
                text="Confirmer",
                command=on_confirm,
                width=100,
                fg_color="#2ecc71",
                hover_color="#27ae60"
            ).pack(side=ctk.RIGHT, padx=10)
            
            # Attendre que la fenêtre soit fermée
            self.view.parent.wait_window(dialog)
            
            return True
            
        except Exception as e:
            logger.error(f"Erreur lors de l'auto-remplissage: {e}")
            DialogUtils.show_message(self.view.parent, "Erreur", f"Une erreur est survenue: {str(e)}", "error")
            return False
    
    def new_document(self, template_id=None):
        """
        Crée un nouveau document à partir d'un modèle
        
        Args:
            template_id: ID du modèle à utiliser (optionnel)
        """
        try:
            # Vérifier s'il y a des modèles disponibles
            if not self.model.templates:
                DialogUtils.show_message(self.view.parent, "Attention", "Aucun modèle disponible. Veuillez d'abord créer un modèle.", "warning")
                return
            
            # Vérifier s'il y a des clients disponibles
            if not self.model.clients:
                DialogUtils.show_message(self.view.parent, "Attention", "Aucun client disponible. Veuillez d'abord ajouter un client.", "warning")
                return
            
            # Créer une nouvelle instance du formulaire de document
            try:
                from views.document_form_view import DocumentFormView
            except ImportError:
                # Fallback: essayer l'ancien chemin d'importation
                from views.document_view import DocumentFormView
            
            form = DocumentFormView(self.view.parent, self.model, on_save_callback=self.view.update_view)
            
            # Si un modèle est spécifié, le charger
            if template_id:
                # Trouver le modèle
                template = next((t for t in self.model.templates if t.get("id") == template_id), None)
                if template:
                    # Pré-remplir le formulaire avec les données du modèle
                    form.template_data = template
                    
                    # Sélectionner le modèle dans le combobox après la création de la vue
                    def select_template():
                        template_name = f"{template.get('name')} ({template.get('type', '')})"
                        form.template_var.set(template_name)
                        # Mettre à jour les informations du modèle manuellement
                        form._update_template_info()
                    
                    # Exécuter après un court délai pour s'assurer que l'interface est prête
                    form.dialog.after(100, select_template)
                else:
                    DialogUtils.show_message(self.view.parent, "Erreur", "Modèle non trouvé", "error")
                    return
            
            logger.info(f"Nouveau document créé" + (f" avec le modèle {template_id}" if template_id else ""))
            
        except Exception as e:
            logger.error(f"Erreur lors de la création du document: {e}")
            DialogUtils.show_message(self.view.parent, "Erreur", f"Une erreur est survenue: {str(e)}", "error")
    
    def open_document(self, document_id):
        """
        Ouvre un document pour le visualiser
        
        Args:
            document_id: ID du document à ouvrir
        """
        try:
            # Récupérer le document
            document = next((d for d in self.model.documents if d.get('id') == document_id), None)
            
            if not document:
                DialogUtils.show_message(self.view.parent, "Erreur", "Document non trouvé", "error")
                return
            
            # Vérifier si le fichier existe avec gestion des chemins multiples
            file_path = None
            
            # Vérifier d'abord si nous avons une structure file_paths
            if "file_paths" in document:
                paths = document["file_paths"]
                # Essayer les chemins dans cet ordre: main, by_client, by_type, by_date
                for key in ["main", "by_client", "by_type", "by_date"]:
                    if key in paths and os.path.exists(paths[key]):
                        file_path = paths[key]
                        break
            
            # Sinon, vérifier le chemin simple
            if not file_path and "file_path" in document and os.path.exists(document["file_path"]):
                file_path = document["file_path"]
            
            if not file_path or not os.path.exists(file_path):
                # Le fichier n'existe pas, on va le régénérer
                template = next((t for t in self.model.templates if t.get('id') == document.get('template_id')), None)
                client = next((c for c in self.model.clients if c.get('id') == document.get('client_id')), None)
                
                if not template or not client:
                    DialogUtils.show_message(self.view.parent, "Erreur", "Impossible de régénérer le document : modèle ou client non trouvé", "error")
                    return
                
                # Générer le document
                success, result = self.generate_document_content(template, client, document.get('variables', {}))
                
                if not success:
                    DialogUtils.show_message(self.view.parent, "Erreur", f"Erreur lors de la régénération du document : {result}", "error")
                    return
                
                # Mettre à jour les chemins du fichier
                if isinstance(result, dict):
                    document['file_paths'] = result
                    file_path = result.get('main')
                else:
                    file_path = result
                    document['file_path'] = file_path
                
                self.model.save_documents()
            
            # Ouvrir le fichier avec l'application par défaut du système
            try:
                import subprocess
                
                if os.name == 'nt':  # Windows
                    os.startfile(file_path)
                elif os.name == 'posix':  # macOS ou Linux
                    subprocess.call(('open' if os.uname().sysname == 'Darwin' else 'xdg-open', file_path))
                
                logger.info(f"Document ouvert: {document_id} - {file_path}")
                
            except Exception as e:
                logger.error(f"Erreur lors de l'ouverture du document: {e}")
                DialogUtils.show_message(self.view.parent, "Erreur", f"Erreur lors de l'ouverture du document: {str(e)}", "error")
                
        except Exception as e:
            logger.error(f"Erreur lors de l'ouverture du document: {e}")
            DialogUtils.show_message(self.view.parent, "Erreur", f"Une erreur est survenue: {str(e)}", "error")
    
    def download_document(self, document_id):
        """
        Télécharge (copie) un document vers un emplacement choisi par l'utilisateur
        
        Args:
            document_id: ID du document à télécharger
        """
        # Récupérer le document
        document = next((d for d in self.model.documents if d.get('id') == document_id), None)
        
        if not document:
            DialogUtils.show_message(self.view.parent, "Erreur", "Document non trouvé", "error")
            return
        
        # Trouver le fichier source
        file_path = None
        
        # Vérifier d'abord si nous avons une structure file_paths
        if "file_paths" in document:
            paths = document["file_paths"]
            # Essayer les chemins dans cet ordre: main, by_client, by_type, by_date
            for key in ["main", "by_client", "by_type", "by_date"]:
                if key in paths and os.path.exists(paths[key]):
                    file_path = paths[key]
                    break
        
        # Sinon, vérifier le chemin simple
        if not file_path and "file_path" in document and os.path.exists(document["file_path"]):
            file_path = document["file_path"]
        
        if not file_path or not os.path.exists(file_path):
            DialogUtils.show_message(self.view.parent, "Erreur", "Le fichier du document est introuvable", "error")
            return
        
        # Déterminer l'extension du fichier
        _, ext = os.path.splitext(file_path)
        
        # Ouvrir une boîte de dialogue pour choisir l'emplacement de sauvegarde
        dest_path = filedialog.asksaveasfilename(
            title="Enregistrer le document",
            defaultextension=ext,
            initialfile=os.path.basename(file_path),
            filetypes=[(f"Fichiers {ext.upper()}", f"*{ext}"), ("Tous les fichiers", "*.*")]
        )
        
        if not dest_path:
            return
        
        try:
            # Copier le fichier
            shutil.copy2(file_path, dest_path)
            
            logger.info(f"Document téléchargé: {document_id} - {dest_path}")
            DialogUtils.show_message(self.view.parent, "Succès", "Document téléchargé avec succès", "success")
        except Exception as e:
            logger.error(f"Erreur lors du téléchargement du document: {e}")
            DialogUtils.show_message(self.view.parent, "Erreur", f"Erreur lors du téléchargement du document: {str(e)}", "error")
    
    def filter_documents(self, *args):
        """
        Filtre les documents selon les critères sélectionnés
        """
        # Vérifier si la vue a une méthode _apply_filters
        if hasattr(self.view, '_apply_filters'):
            self.view._apply_filters(*args)
        # Sinon, essayer avec _apply_document_filters
        elif hasattr(self.view, '_apply_document_filters'):
            self.view._apply_document_filters(*args)
        # En dernier recours, mettre à jour la vue
        else:
            self.view.update_view()
        
        logger.debug("Documents filtrés")
    
    def generate_document_content(self, template, client, variables):
        """
        Génère le contenu d'un document en utilisant le générateur de documents
        
        Args:
            template: Modèle de document
            client: Informations du client
            variables: Valeurs des variables spécifiques
            
        Returns:
            tuple: (bool, str) - Succès et chemin du fichier généré ou message d'erreur
        """
        try:
            # Créer le générateur de documents
            from utils.document_generator import DocumentGenerator
            generator = DocumentGenerator(self.model.config)
            
            # Préparer les informations de l'entreprise
            company_info = {
                "name": self.model.config.get("app.company_name", ""),
                "address": self.model.config.get("app.company_address", ""),
                "email": self.model.config.get("app.company_email", ""),
                "phone": self.model.config.get("app.company_phone", ""),
                "website": self.model.config.get("app.company_website", "")
            }
            
            # Générer le nom du fichier
            doc_filename = self.generate_filename(
                template.get("type", "document"),
                client.get("name", "client"),
                datetime.now().strftime("%Y-%m-%d")
            )
            
            # Déterminer le format et le chemin de sortie
            format_type = self.model.config.get("document.default_format", "pdf")
            output_path = os.path.join(self.model.paths['documents'], f"{doc_filename}.{format_type}")
            
            # Générer le document
            success = generator.generate_document(
                output_path,
                template,
                client,
                company_info,
                variables,
                format_type
            )
            
            if success:
                return True, output_path
            else:
                return False, "Erreur lors de la génération du document"
            
        except Exception as e:
            logger.error(f"Erreur lors de la génération du document: {e}")
            return False, str(e)
    
    def generate_filename(self, doc_type, client_name, date):
        """
        Génère un nom de fichier pour le document
        
        Args:
            doc_type: Type de document
            client_name: Nom du client
            date: Date du document
            
        Returns:
            str: Nom de fichier normalisé
        """
        # Nettoyer les noms
        doc_type = self.clean_filename(doc_type)
        client_name = self.clean_filename(client_name)
        
        # Formatter la date
        date_format = self.model.config.get("document.date_format", "%Y-%m-%d")
        try:
            date_obj = datetime.strptime(date, "%Y-%m-%d")
            formatted_date = date_obj.strftime(date_format)
        except:
            formatted_date = datetime.now().strftime(date_format)
        
        # Construire le nom du fichier selon le pattern configuré
        pattern = self.model.config.get("document.filename_pattern", "{document_type}_{client_name}_{date}")
        
        filename = pattern.replace("{document_type}", doc_type)
        filename = filename.replace("{client_name}", client_name)
        filename = filename.replace("{date}", formatted_date)
        
        return filename
    
    def clean_filename(self, name):
        """
        Nettoie un nom pour qu'il soit utilisable dans un nom de fichier
        
        Args:
            name: Nom à nettoyer
            
        Returns:
            str: Nom nettoyé
        """
        # Supprimer les caractères spéciaux et remplacer les espaces par des underscores
        name = re.sub(r'[\\/*?:"<>|]', '', name)
        name = name.replace(' ', '_')
        name = name.replace('/', '_')
        name = name.replace('\\', '_')
        
        return name
    
    def generate_pdf(self, file_path, content, title, client, company_name):
        """
        Génère un fichier PDF
        
        Args:
            file_path: Chemin du fichier à créer
            content: Contenu du document
            title: Titre du document
            client: Informations du client
            company_name: Nom de l'entreprise
        """
        try:
            # Créer le répertoire parent si nécessaire
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Créer un PDF
            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4
            
            # Définir la police à utiliser - Helvetica par défaut
            font_name = "Helvetica"
            
            # Essayer d'utiliser une police qui supporte mieux l'Unicode si disponible
            try:
                # Importer les modules nécessaires
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                
                # Vérifier l'existence de DejaVuSans.ttf
                dejavu_path = None
                
                # Chemins possibles pour DejaVuSans.ttf
                possible_paths = [
                    os.path.join(os.path.dirname(__file__), 'fonts', 'DejaVuSans.ttf'),  # Dans le dossier fonts/ relatif à ce script
                    os.path.join(os.path.dirname(os.path.dirname(__file__)), 'fonts', 'DejaVuSans.ttf'),  # Dans le dossier fonts/ du projet
                    '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',  # Chemin Linux commun
                    'C:\\Windows\\Fonts\\DejaVuSans.ttf',  # Chemin Windows
                    os.path.expanduser('~/Library/Fonts/DejaVuSans.ttf')  # Chemin macOS
                ]
                
                for path in possible_paths:
                    if os.path.exists(path):
                        dejavu_path = path
                        break
                
                if dejavu_path:
                    pdfmetrics.registerFont(TTFont('DejaVuSans', dejavu_path))
                    c.setFont("DejaVuSans", 10)
                    font_name = "DejaVuSans"
                    logger.info(f"Utilisation de la police DejaVuSans pour le PDF")
            except Exception as font_error:
                logger.warning(f"Impossible d'utiliser la police DejaVuSans, utilisation de Helvetica: {font_error}")
            
            # Définir les marges et la largeur disponible
            margin_left = 50
            margin_right = 50
            available_width = width - margin_left - margin_right
            font_size = 10
            
            # Vérifier si le titre est trop long
            c.setFont(font_name + "-Bold" if font_name == "Helvetica" else font_name, 16)
            title = self._sanitize_text(title)
            
            if c.stringWidth(title, font_name + "-Bold" if font_name == "Helvetica" else font_name, 16) > available_width:
                # Découper le titre en plusieurs lignes
                words = title.split(' ')
                current_line = ""
                y_position = height - 70
                
                for word in words:
                    test_line = current_line + " " + word if current_line else word
                    if c.stringWidth(test_line, font_name + "-Bold" if font_name == "Helvetica" else font_name, 16) <= available_width:
                        current_line = test_line
                    else:
                        # Dessiner la partie actuelle du titre
                        c.drawString(margin_left, y_position, current_line)
                        y_position -= 20  # Espacement pour le titre
                        current_line = word
                
                # Dessiner la dernière partie du titre
                if current_line:
                    c.drawString(margin_left, y_position, current_line)
                    y_position -= 20
                
                # Ligne séparatrice après la dernière ligne du titre
                c.line(50, y_position - 5, width - 50, y_position - 5)
                y_position -= 20  # Espacement après la ligne
            else:
                # Titre normal sur une seule ligne
                c.drawString(margin_left, height - 70, title)
                
                # Ligne séparatrice
                c.line(50, height - 85, width - 50, height - 85)
                
                y_position = height - 105
            
            # Contenu du document
            c.setFont(font_name, 10)
            
            # Diviser le contenu en lignes
            lines = content.split('\n')
            
            for line in lines:
                # Sanitize line text
                line = self._sanitize_text(line)
                
                # Vérifier si la ligne est trop longue
                if c.stringWidth(line, font_name, font_size) > available_width:
                    # Découpage de la ligne en plusieurs lignes
                    words = line.split(' ')
                    current_line = ""
                    
                    for word in words:
                        test_line = current_line + " " + word if current_line else word
                        if c.stringWidth(test_line, font_name, font_size) <= available_width:
                            current_line = test_line
                        else:
                            # Vérifier s'il reste assez d'espace sur la page
                            if y_position < 70:
                                c.showPage()
                                c.setFont(font_name, font_size)
                                y_position = height - 70
                            
                            # Écrire la ligne actuelle
                            c.drawString(margin_left, y_position, current_line)
                            y_position -= 15
                            current_line = word
                    
                    # Écrire la dernière partie de la ligne
                    if current_line:
                        if y_position < 70:
                            c.showPage()
                            c.setFont(font_name, font_size)
                            y_position = height - 70
                        
                        c.drawString(margin_left, y_position, current_line)
                        y_position -= 15
                else:
                    if y_position < 70:
                        # Créer une nouvelle page
                        c.showPage()
                        c.setFont(font_name, font_size)
                        y_position = height - 70
                    
                    # Écrire la ligne
                    c.drawString(margin_left, y_position, line)
                    y_position -= 15
            
            # Finaliser le PDF
            c.showPage()
            c.save()
            
            logger.info(f"PDF généré: {file_path}")
            
        except Exception as e:
            logger.error(f"Erreur lors de la génération du PDF: {e}")
            raise
            
    def _sanitize_text(self, text):
        """
        Nettoie le texte pour éviter les problèmes avec les caractères spéciaux dans les PDF
        
        Args:
            text: Texte à nettoyer
            
        Returns:
            str: Texte nettoyé
        """
        if not isinstance(text, str):
            text = str(text)
        
        # Remplacer les caractères problématiques par leurs équivalents compatibles
        special_chars_map = {
            '…': '...',
            '–': '-',
            '—': '-',
            ''': "'",
            ''': "'",
            '"': '"',
            '"': '"',
            '«': '"',
            '»': '"',
            '•': '*'
        }
        
        for char, replacement in special_chars_map.items():
            text = text.replace(char, replacement)
        
        # Supprimer les caractères non imprimables ou non supportés
        text = ''.join(c for c in text if ord(c) >= 32 or c in '\n\r\t')
        
        return text
    
    def generate_docx(self, file_path, content, title, client, company_name):
        """
        Génère un fichier DOCX
        
        Args:
            file_path: Chemin du fichier à créer
            content: Contenu du document
            title: Titre du document
            client: Informations du client
            company_name: Nom de l'entreprise
        """
        try:
            # Créer le répertoire parent si nécessaire
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Créer un document Word
            doc = docx.Document()
            
            # Définir la largeur maximale des marges
            for section in doc.sections:
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)
                section.top_margin = Inches(1.2)
                section.bottom_margin = Inches(1.2)
            
            # Définir une police compatible avec les caractères spéciaux
            # Liste des polices Unicode communes qui fonctionnent bien avec les caractères spéciaux
            unicode_fonts = ['Arial Unicode MS', 'Calibri', 'Times New Roman', 'Arial', 'Segoe UI']
            
            # Définir le style par défaut
            doc_style = doc.styles['Normal']
            for font_name in unicode_fonts:
                try:
                    doc_style.font.name = font_name
                    break
                except:
                    continue
            
            # Titre
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(16)
            
            # S'assurer que le titre utilise aussi une police compatible
            for font_name in unicode_fonts:
                try:
                    title_run.font.name = font_name
                    break
                except:
                    continue
            
            # Ligne séparatrice
            doc.add_paragraph("_______________________________________________________________")
            
            # Contenu
            # Diviser le contenu en paragraphes
            paragraphs = content.split('\n')
            
            for paragraph in paragraphs:
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph)
                    # Limiter la largeur des paragraphes en ajustant les retraits
                    p.paragraph_format.first_line_indent = Inches(0)
                    p.paragraph_format.left_indent = Inches(0)
                    p.paragraph_format.right_indent = Inches(0)
                    
                    # S'assurer que chaque paragraphe utilise une police compatible
                    for run in p.runs:
                        for font_name in unicode_fonts:
                            try:
                                run.font.name = font_name
                                break
                            except:
                                continue
            
            # Enregistrer le document
            doc.save(file_path)
            
            logger.info(f"DOCX généré: {file_path}")
            
        except Exception as e:
            logger.error(f"Erreur lors de la génération du DOCX: {e}")
            raise

    def process_external_document(self, file_path, processor=None):
        """
        Traite un document externe uploadé par l'utilisateur
        
        Args:
            file_path: Chemin vers le fichier à traiter
            processor: Processeur de document à utiliser (optionnel)
        """
        try:
            logger.info(f"Traitement du document externe: {file_path}")
            
            # Récupérer le nom du fichier sans l'extension
            file_name = os.path.basename(file_path)
            file_base, file_ext = os.path.splitext(file_name)
            
            # Si aucun processeur n'est fourni, créer une instance
            if processor is None:
                from ai.document_processor import AIDocumentProcessor
                processor = AIDocumentProcessor()
                
            # Lire le contenu du document
            try:
                document_content = processor._read_file_safely(file_path)
                logger.info(f"Document lu avec succès: {len(document_content)} caractères")
            except Exception as e:
                logger.error(f"Erreur lors de la lecture du document: {e}")
                raise ValueError(f"Impossible de lire le document: {e}")
                
            # Créer une fenêtre d'analyse en cours
            progress_dialog = ctk.CTkToplevel(self.model.root)
            progress_dialog.title("Analyse en cours")
            progress_dialog.geometry("400x200")
            progress_dialog.resizable(False, False)
            progress_dialog.grab_set()
            
            # Centrer la fenêtre
            progress_dialog.update_idletasks()
            x = (progress_dialog.winfo_screenwidth() - progress_dialog.winfo_width()) // 2
            y = (progress_dialog.winfo_screenheight() - progress_dialog.winfo_height()) // 2
            progress_dialog.geometry(f"+{x}+{y}")
            
            # Cadre principal
            frame = ctk.CTkFrame(progress_dialog, fg_color="transparent")
            frame.pack(fill=ctk.BOTH, expand=True, padx=20, pady=20)
            
            # Titre
            ctk.CTkLabel(
                frame,
                text="Analyse du document en cours",
                font=ctk.CTkFont(size=16, weight="bold")
            ).pack(pady=(0, 10))
            
            # Message
            message_label = ctk.CTkLabel(
                frame,
                text=f"Extraction des informations du document...\nVeuillez patienter.",
                wraplength=360
            )
            message_label.pack(pady=10)
            
            # Barre de progression indéterminée
            progress = ctk.CTkProgressBar(frame, width=300, mode="indeterminate")
            progress.pack(pady=10)
            progress.start()
            
            # Bouton d'annulation
            cancel_button = ctk.CTkButton(
                frame,
                text="Annuler",
                width=100,
                fg_color="#e74c3c",
                hover_color="#c0392b",
                command=progress_dialog.destroy
            )
            cancel_button.pack(pady=10)
            
            # Mettre à jour l'interface
            progress_dialog.update()
            
            # Fonction pour l'analyse asynchrone
            def analyze_document():
                try:
                    # Analyser le document avec l'IA pour extraire les informations
                    analysis_result = processor.analyze_document(file_path)
                    
                    # Sauvegarder le document original dans le dossier du projet si nécessaire
                    documents_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "documents", "uploads")
                    os.makedirs(documents_dir, exist_ok=True)
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    saved_file_path = os.path.join(documents_dir, f"{file_base}_{timestamp}{file_ext}")
                    shutil.copy(file_path, saved_file_path)
                    
                    # Fermer la fenêtre de progression
                    progress_dialog.destroy()
                    
                    # Créer une entrée dans la base de données pour ce document
                    document_id = self.model.documents.add_document({
                        "name": file_base,
                        "type": "uploaded",
                        "path": saved_file_path,
                        "date_created": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "content": document_content[:500] + "..." if len(document_content) > 500 else document_content,
                        "metadata": analysis_result if analysis_result else {}
                    })
                    
                    # Ouvrir le document dans l'éditeur
                    if document_id:
                        self.open_document(document_id)
                    
                    logger.info(f"Document externe traité avec succès: {file_path}")
                    
                except Exception as e:
                    logger.error(f"Erreur lors de l'analyse du document: {e}")
                    # Fermer la fenêtre de progression
                    progress_dialog.destroy()
                    # Afficher l'erreur
                    DialogUtils.show_message(
                        self.model.root,
                        "Erreur d'analyse",
                        f"Une erreur est survenue lors de l'analyse du document:\n{e}",
                        "error"
                    )
            
            # Lancer l'analyse en arrière-plan
            import threading
            analysis_thread = threading.Thread(target=analyze_document)
            analysis_thread.daemon = True
            analysis_thread.start()
            
        except Exception as e:
            logger.error(f"Erreur lors du traitement du document externe: {e}")
            DialogUtils.show_message(
                self.model.root,
                "Erreur",
                f"Une erreur est survenue lors du traitement du document: {e}",
                "error"
            )