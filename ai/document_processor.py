#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Module de traitement des documents avec l'IA
Utilise Llama pour l'analyse et la personnalisation des documents
"""

import logging
import json
import os
import re
import chardet
import requests
import time
import traceback
from typing import Dict, List, Optional, Tuple, Any

logging.basicConfig(level=logging.DEBUG)

logger = logging.getLogger("VynalDocsAutomator.AIDocumentProcessor")

class AIDocumentProcessor:
    """
    Classe pour le traitement des documents avec l'IA
    Utilise Llama pour analyser et personnaliser les documents
    """
    
    def __init__(self):
        """
        Initialise le processeur de documents
        """
        self.config = self._load_config()
        self.cache = {}
        self.cache_size = 100  # Réduit de 1000 à 100
        self.fallback_mode = False  # Ajout de l'attribut fallback_mode initialisé à False
        
    def _load_config(self):
        """
        Charge la configuration depuis le fichier config.json
        Gère les différents encodages pour éviter les problèmes de décodage
        """
        try:
            # Utiliser le chemin absolu pour éviter les problèmes de chemin relatif
            config_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "config", "ai_config.json")
            logger.info(f"Tentative de chargement de la configuration depuis: {config_path}")
            
            if not os.path.exists(config_path):
                logger.warning(f"Le fichier de configuration n'existe pas: {config_path}")
                # Si le fichier n'existe pas, retourner une configuration par défaut
                return {
                    "api_url": "http://localhost:11434/api/generate",
                    "model": "llama3",
                    "options": {
                        "temperature": 0.1,
                        "num_predict": 1024,
                        "top_p": 0.85,
                        "frequency_penalty": 0.1,
                        "stop": ["\n\n\n", "###", "```"],
                        "timeout": 20
                    }
                }
            
            # Essayer différents encodages
            encodings = ['utf-8', 'latin-1', 'cp1252', 'utf-16', 'utf-16-le', 'utf-16-be']
            
            for encoding in encodings:
                try:
                    with open(config_path, 'r', encoding=encoding) as f:
                        config = json.load(f)
                    logger.info(f"Configuration chargée avec l'encodage {encoding} depuis {config_path}")
                    return config
                except UnicodeDecodeError:
                    logger.debug(f"Échec du décodage avec l'encodage {encoding}")
                    continue  # Essayer le prochain encodage
                except json.JSONDecodeError:
                    # Si le fichier s'ouvre mais que le JSON est invalide
                    logger.warning(f"Format JSON invalide dans le fichier de configuration avec l'encodage {encoding}")
                    continue
            
            # Si on arrive ici, tous les encodages ont échoué
            # Essai en mode binaire avec détection d'encodage
            logger.warning("Tous les encodages standards ont échoué, tentative de détection automatique")
            import chardet
            try:
                with open(config_path, 'rb') as f:
                    raw_data = f.read()
                    result = chardet.detect(raw_data)
                    detected_encoding = result['encoding']
                    logger.info(f"Encodage détecté: {detected_encoding} (confiance: {result['confidence']})")
                    
                    if detected_encoding:
                        try:
                            config_str = raw_data.decode(detected_encoding)
                            config = json.loads(config_str)
                            logger.info(f"Configuration chargée avec l'encodage détecté {detected_encoding}")
                            return config
                        except Exception as e:
                            logger.warning(f"Échec avec l'encodage détecté {detected_encoding}: {e}")
            except Exception as e:
                logger.warning(f"Échec de la détection d'encodage: {e}")
                
            # Créer une nouvelle configuration par défaut
            logger.warning("Création d'une nouvelle configuration par défaut")
            config = {
                "api_url": "http://localhost:11434/api/generate",
                "model": "llama3",
                "options": {
                    "temperature": 0.1,
                    "num_predict": 1024,
                    "top_p": 0.85,
                    "frequency_penalty": 0.1,
                    "stop": ["\n\n\n", "###", "```"],
                    "timeout": 20
                }
            }
            
            # Essayer de sauvegarder cette configuration
            try:
                os.makedirs(os.path.dirname(config_path), exist_ok=True)
                with open(config_path, 'w', encoding='utf-8') as f:
                    json.dump(config, f, indent=4)
                logger.info("Nouvelle configuration par défaut créée")
            except Exception as e:
                logger.error(f"Impossible de créer la configuration par défaut: {e}")
                
            return config
            
        except Exception as e:
            logger.warning(f"Impossible de charger la configuration: {e}")
            # Configuration par défaut
            return {
                "api_url": "http://localhost:11434/api/generate",
                "model": "llama3",
                "options": {
                    "temperature": 0.1,
                    "num_predict": 1024,
                    "top_p": 0.85,
                    "frequency_penalty": 0.1,
                    "stop": ["\n\n\n", "###", "```"],
                    "timeout": 20
                }
            }
        
    def _read_file_safely(self, file_path: str) -> str:
        """
        Lit un fichier de manière sécurisée avec plusieurs tentatives d'encodage
        
        Args:
            file_path: Chemin du fichier à lire
            
        Returns:
            Contenu du fichier ou chaîne vide en cas d'erreur
        """
        logger.info(f"====== LECTURE DU FICHIER: {file_path} ======")
        
        if not file_path or not os.path.exists(file_path):
            logger.error(f"ERREUR: Fichier introuvable: {file_path}")
            return ""
            
        try:
            # Vérifier la taille du fichier
            file_size = os.path.getsize(file_path)
            logger.info(f"Taille du fichier: {file_size} octets")
            
            if file_size == 0:
                logger.error(f"ERREUR: Fichier vide: {file_path}")
                return ""
                
            # Liste des encodages à essayer dans l'ordre
            encodings = ['utf-8', 'latin-1', 'cp1252', 'ISO-8859-1', 'ascii']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                        logger.info(f"Fichier lu avec succès en utilisant l'encodage {encoding}")
                        logger.info(f"Longueur du contenu: {len(content)} caractères")
                        logger.info(f"Début du contenu: {content[:min(100, len(content))]}...")
                        return content
                except UnicodeDecodeError:
                    logger.info(f"L'encodage {encoding} a échoué, essai avec le suivant...")
                except Exception as e:
                    logger.warning(f"Erreur lors de la lecture avec l'encodage {encoding}: {str(e)}")
            
            # Si tous les encodages ont échoué, essayer une lecture binaire
            logger.warning("Tous les encodages ont échoué, tentative de lecture binaire...")
            with open(file_path, 'rb') as f:
                binary_content = f.read()
                # Essayer de décoder avec remplacement des caractères non reconnus
                content = binary_content.decode('utf-8', errors='replace')
                logger.warning(f"Contenu lu en binaire et décodé avec remplacement des caractères: {len(content)} caractères")
                
            if not content:
                logger.error(f"ÉCHEC TOTAL: Impossible de lire le fichier {file_path} avec tous les encodages essayés")
                return ""
                
            return content
            
        except Exception as e:
            logger.error(f"ERREUR CRITIQUE lors de la lecture du fichier {file_path}: {str(e)}")
            traceback.print_exc()
            return ""
    
    def analyze_document(self, file_path: str, file_type: str = None, encoding: str = 'utf-8') -> Dict[str, Any]:
        """
        Analyse un document pour en extraire des informations structurées
        
        Args:
            file_path: Chemin du fichier à analyser
            file_type: Type de fichier (optionnel, détecté automatiquement si non fourni)
            encoding: Encodage du fichier
            
        Returns:
            Dict contenant les informations extraites
        """
        try:
            start_time = time.time()
            logger.info(f"Début de l'analyse du document: {file_path}")
            
            # Convertir le document en texte
            content = self.convert_to_text(file_path, file_type, encoding)
            
            if not content:
                logger.warning("Aucun contenu textuel n'a pu être extrait du document")
                return {"error": "Aucun contenu textuel n'a pu être extrait du document", "variables": {}}
                
            # Première approche: Extraction directe par expressions régulières
            # Cette approche est très rapide et fonctionne souvent bien pour les documents simples
            direct_variables = self._extract_basic_variables_from_text(content)
            
            # Si on a déjà trouvé un nombre suffisant de variables, retourner directement
            if len(direct_variables) >= 4:
                logger.info(f"Analyse rapide réussie - {len(direct_variables)} variables extraites")
                return {
                    "variables": direct_variables,
                    "_meta": {
                        "extraction_method": "direct_regex",
                        "processing_time": round(time.time() - start_time, 2),
                        "document_size": len(content)
                    }
                }
            
            # Deuxième approche: Analyse par sections avec IA
            try:
                logger.info("Tentative d'analyse avancée par IA")
                result = self._analyze_document_by_sections(content)
                
                # Fusionner avec l'extraction directe pour enrichir les résultats
                if "variables" in result:
                    # Priorité aux variables déjà trouvées par l'IA
                    for var_name, value in direct_variables.items():
                        if var_name not in result["variables"]:
                            result["variables"][var_name] = value
                            
                # Ajouter des métadonnées utiles
                if "_meta" not in result:
                    result["_meta"] = {}
                result["_meta"]["processing_time"] = round(time.time() - start_time, 2)
                result["_meta"]["document_size"] = len(content)
                
                return result
                
            except Exception as e:
                logger.error(f"Erreur lors de l'analyse avec IA: {str(e)}")
                # En cas d'échec, retourner les variables extraites par regex
                return {
                    "variables": direct_variables,
                    "error": str(e),
                    "_meta": {
                        "extraction_method": "fallback_to_regex",
                        "processing_time": round(time.time() - start_time, 2),
                        "document_size": len(content)
                    }
                }
                
        except Exception as e:
            logger.error(f"Erreur lors de l'analyse du document: {str(e)}")
            return {"error": str(e), "variables": {}}
    
    def _smart_reduce_content(self, content: str, max_size: int) -> str:
        """
        Réduit intelligemment le contenu pour l'analyse en préservant les parties importantes
        
        Args:
            content: Contenu à réduire
            max_size: Taille maximale souhaitée
            
        Returns:
            Contenu réduit intelligemment
        """
        if len(content) <= max_size:
            return content
            
        # Zones prioritaires
        start_size = min(max_size // 3, 800)  # Début du document (en-têtes, intro)
        end_size = min(max_size // 3, 600)    # Fin du document (conclusion, signatures)
        
        # Extraire les sections importantes
        content_start = content[:start_size]
        content_end = content[-end_size:]
        
        # Espace restant pour les sections importantes
        remaining_size = max_size - (start_size + end_size + 100)  # 100 pour le texte de séparation
        
        # Mots-clés importants selon le type de document
        important_keywords = [
            # Identités
            "client", "nom", "prénom", "société", "entreprise", 
            # Coordonnées
            "adresse", "email", "téléphone", "portable", "contact",
            # Données financières 
            "montant", "prix", "total", "somme", "€", "euros", "eur",
            # Références et dates
            "référence", "facture", "commande", "date", "numéro",
            # Termes légaux
            "contrat", "signature", "légal", "obligation"
        ]
        
        # Découper en paragraphes et noter leur importance
        paragraphs = re.split(r'\n\s*\n', content)
        scored_paragraphs = []
        
        for para in paragraphs:
            # Ignorer les paragraphes vides ou trop courts
            if len(para.strip()) < 5:
                continue
                
            # Calculer un score d'importance
            score = 0
            para_lower = para.lower()
            
            # Points pour chaque mot-clé trouvé
            for keyword in important_keywords:
                if keyword in para_lower:
                    score += 2
                    
            # Points supplémentaires pour les formats spéciaux
            # Dates
            if re.search(r'\b\d{1,2}[/.-]\d{1,2}[/.-]\d{2,4}\b', para):
                score += 3
            # Montants
            if re.search(r'\b\d+(?:[,.]\d{1,2})?(?:\s?[€$]|EUR)?\b', para):
                score += 3
            # Emails
            if re.search(r'\b[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+\b', para):
                score += 3
            # Téléphones
            if re.search(r'\b(?:0\d[\s.-]?){4,5}\d{2}\b', para):
                score += 3
            
            # Ajouter le paragraphe avec son score
            scored_paragraphs.append((para, score))
        
        # Trier les paragraphes par score d'importance décroissant
        scored_paragraphs.sort(key=lambda x: x[1], reverse=True)
        
        # Sélectionner les paragraphes les plus importants
        selected_paragraphs = []
        total_length = 0
        
        for para, score in scored_paragraphs:
            if total_length + len(para) <= remaining_size:
                selected_paragraphs.append(para)
                total_length += len(para) + 2  # +2 pour les sauts de ligne
            else:
                break
        
        # Construire le contenu final
        reduced_content = content_start + "\n\n"
        if selected_paragraphs:
            reduced_content += "--- INFORMATIONS IMPORTANTES ---\n\n" + "\n\n".join(selected_paragraphs) + "\n\n"
        reduced_content += "--- FIN DU DOCUMENT ---\n\n" + content_end
        
        logger.info(f"Contenu réduit de {len(content)} à {len(reduced_content)} caractères")
        return reduced_content

    def convert_to_text(self, file_path, file_type=None, encoding='utf-8'):
        """
        Convertit un document en texte brut, quel que soit son format
        
        Args:
            file_path: Chemin du fichier à convertir
            file_type: Type du fichier (extension)
            encoding: Encodage du fichier
            
        Returns:
            str: Contenu textuel du document
        """
        if not file_type:
            file_type = os.path.splitext(file_path)[1].lower()[1:]
            
        try:
            # Si le fichier est un .docx (Word)
            if file_type in ['docx', 'doc']:
                try:
                    import docx
                    doc = docx.Document(file_path)
                    full_text = []
                    # Extraire le texte de chaque paragraphe
                    for para in doc.paragraphs:
                        full_text.append(para.text)
                    # Extraire le texte des tableaux
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    full_text.append(para.text)
                    return '\n'.join(full_text)
                except ImportError:
                    logger.warning("Module python-docx non disponible, tentative de fallback")
                    # Fallback à une lecture textuelle
                    try:
                        with open(file_path, 'rb') as f:
                            import chardet
                            result = chardet.detect(f.read())
                            detected_encoding = result['encoding']
                        with open(file_path, 'r', encoding=detected_encoding or 'utf-8', errors='replace') as f:
                            return f.read()
                    except Exception as e:
                        logger.error(f"Erreur lors de la lecture du fichier Word: {e}")
                        return f"Erreur: {str(e)}"
                except Exception as e:
                    logger.error(f"Erreur lors de l'extraction du texte du fichier Word: {e}")
                    return f"Erreur: {str(e)}"
                    
            # Si le fichier est un .pdf
            elif file_type == 'pdf':
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        text = ''
                        for page_num in range(len(pdf_reader.pages)):
                            page = pdf_reader.pages[page_num]
                            text += page.extract_text() + '\n'
                        return text
                except ImportError:
                    logger.warning("Module PyPDF2 non disponible, tentative de fallback")
                    # Fallback vers une lecture textuelle basique
                    try:
                        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                            return f.read()
                    except Exception as e:
                        logger.error(f"Erreur lors de la lecture du fichier PDF: {e}")
                        return f"Erreur: {str(e)}"
                except Exception as e:
                    logger.error(f"Erreur lors de l'extraction du texte du fichier PDF: {e}")
                    return f"Erreur: {str(e)}"
                    
            # Si le fichier est une image
            elif file_type in ['jpg', 'jpeg', 'png', 'gif', 'bmp']:
                try:
                    from PIL import Image
                    import pytesseract
                    image = Image.open(file_path)
                    text = pytesseract.image_to_string(image, lang='fra+eng')
                    return text
                except ImportError:
                    logger.warning("Module pytesseract non disponible")
                    return "Erreur: L'extraction de texte depuis des images nécessite le module pytesseract"
                except Exception as e:
                    logger.error(f"Erreur lors de l'OCR sur l'image: {e}")
                    return f"Erreur: {str(e)}"
                    
            # Pour les fichiers texte (.txt, .md, .html, etc.)
            else:
                # Essayer différents encodages pour gérer tous les cas de figure
                encodings = ['utf-8', 'latin-1', 'cp1252', 'utf-16', 'utf-16-le', 'utf-16-be']
                content = None
                last_error = None
                
                # Si un encodage spécifique est demandé, le tester en premier
                if encoding and encoding not in encodings:
                    encodings.insert(0, encoding)
                
                for enc in encodings:
                    try:
                        with open(file_path, 'r', encoding=enc) as f:
                            content = f.read()
                        logger.info(f"Fichier lu avec succès en utilisant l'encodage {enc}")
                        break
                    except UnicodeDecodeError as e:
                        logger.debug(f"Échec de lecture avec l'encodage {enc}: {e}")
                        last_error = e
                    except Exception as e:
                        logger.warning(f"Erreur lors de la lecture du fichier avec l'encodage {enc}: {e}")
                        last_error = e
                
                # Si tous les encodages ont échoué, tenter une détection automatique
                if content is None:
                    try:
                        import chardet
                        with open(file_path, 'rb') as f:
                            raw_data = f.read()
                            result = chardet.detect(raw_data)
                            detected_encoding = result['encoding']
                            
                            if detected_encoding:
                                logger.info(f"Encodage détecté: {detected_encoding}")
                                with open(file_path, 'r', encoding=detected_encoding, errors='replace') as f:
                                    content = f.read()
                            else:
                                # Dernier recours: lecture en mode binaire et conversion forcée
                                content = raw_data.decode('utf-8', errors='replace')
                    except Exception as e:
                        logger.error(f"Erreur lors de la détection de l'encodage: {e}")
                        if last_error:
                            raise last_error
                        else:
                            raise e
                
                return content
                
        except Exception as e:
            logger.error(f"Erreur lors de la conversion du document en texte: {e}")
            return f"Erreur lors de la lecture du document: {str(e)}"
            
    def fill_template(self, template_path: str, client_info: Dict, encoding: str = 'utf-8') -> Tuple[str, Dict]:
        """
        Remplit un modèle avec les informations client
        
        Args:
            template_path: Chemin vers le modèle
            client_info: Informations client
            encoding: Encodage du fichier
            
        Returns:
            Tuple[str, Dict]: (contenu rempli, variables manquantes)
        """
        try:
            logger.info(f"Début du remplissage du modèle: {template_path}")
            
            # 1. Validation initiale
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"Le modèle n'existe pas: {template_path}")
            
            if not client_info or not isinstance(client_info, dict):
                raise ValueError("Les informations client sont invalides")
            
            # 2. Lecture du modèle avec gestion d'encodage
            content = None
            encoding_errors = []
            
            for enc in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    with open(template_path, 'r', encoding=enc) as f:
                        content = f.read()
                        logger.debug(f"Lecture réussie avec l'encodage: {enc}")
                        break
                except UnicodeDecodeError as e:
                    encoding_errors.append(f"Erreur avec {enc}: {str(e)}")
                    continue
                    
            if content is None:
                raise UnicodeError(f"Impossible de lire le fichier. Erreurs: {', '.join(encoding_errors)}")
            
            # 3. Analyse du modèle
            logger.debug("Analyse du modèle...")
            variables = self._analyze_template_content(content)
            
            # 4. Validation des variables requises
            missing_vars = {}
            invalid_vars = []
            
            for var_name, var_info in variables.items():
                if var_info.get('required', False):
                    if var_name not in client_info:
                        missing_vars[var_name] = var_info
                    else:
                        # Valider la valeur selon le type
                        value = client_info[var_name]
                        validation_rules = self._get_validation_rules(var_info.get('type', 'text'))
                        if not self._validate_variable(value, validation_rules):
                            invalid_vars.append(var_name)
            
            if invalid_vars:
                logger.warning(f"Variables invalides: {invalid_vars}")
                raise ValueError(f"Les variables suivantes sont invalides: {', '.join(invalid_vars)}")
            
            # 5. Remplissage du modèle
            if not self.fallback_mode and len(missing_vars) == 0:
                try:
                    logger.info("Utilisation de l'IA pour le remplissage...")
                    filled_content = self._fill_with_ai(content, client_info)
                except Exception as e:
                    logger.error(f"Erreur lors du remplissage IA: {e}")
                    filled_content = self._fill_template_basic(content, client_info)
                else:
                    logger.info("Utilisation du remplissage basique...")
                    filled_content = self._fill_template_basic(content, client_info)
            
            return filled_content, missing_vars
            
        except Exception as e:
            logger.error(f"Erreur lors du remplissage du modèle: {e}")
            return str(e), {
                    "error": {
                        "type": "text",
                    "description": "Erreur de traitement",
                        "required": True,
                    "current_value": str(e)
                }
            }
    
    def _validate_variable(self, value: Any, rules: Dict[str, Any]) -> bool:
        """
        Valide une variable selon les règles définies
        
        Args:
            value: Valeur à valider
            rules: Règles de validation
            
        Returns:
            bool: True si la valeur est valide
        """
        try:
            if 'regex' in rules:
                if not re.match(rules['regex'], str(value)):
                    return False
            
            if 'min_length' in rules and isinstance(value, (str, list)):
                if len(value) < rules['min_length']:
                    return False
            
            if 'max_length' in rules and isinstance(value, (str, list)):
                if len(value) > rules['max_length']:
                    return False
            
            if 'min_value' in rules and isinstance(value, (int, float)):
                if value < rules['min_value']:
                    return False
            
            if 'decimal_places' in rules and isinstance(value, (int, float)):
                str_value = str(value)
                if '.' in str_value:
                    if len(str_value.split('.')[1]) > rules['decimal_places']:
                        return False
            
            return True
            
        except Exception as e:
            logger.error(f"Erreur lors de la validation: {e}")
            return False

    def generate_final_document(self, template_path: str, all_variables: Dict[str, str]) -> str:
        """
        Génère le document final avec toutes les variables
        
        Args:
            template_path: Chemin vers le modèle
            all_variables: Toutes les variables (client + complétées)
            
        Returns:
            Texte final du document
        """
        try:
            # Début du processus - Traçage détaillé
            logger.info(f"=== DÉBUT GÉNÉRATION DOCUMENT ({template_path}) ===")
            logger.info(f"Variables disponibles: {list(all_variables.keys())}")
            
            # Lire le contenu du modèle
            content = self._read_file_safely(template_path)
            if not content:
                logger.error(f"Erreur: Le modèle '{template_path}' est vide ou n'a pas pu être lu")
                return "ERREUR: Le modèle est vide ou n'a pas pu être lu."
            
            # Afficher un aperçu du contenu original
            content_preview = content[:100] + "..." if len(content) > 100 else content
            logger.debug(f"Contenu du template (aperçu): {content_preview}")
            logger.debug(f"Longueur du contenu: {len(content)} caractères")
            
            # Utiliser directement replace_variables qui est optimisé pour préserver exactement la structure
            logger.info(f"Utilisation de replace_variables pour remplacer {len(all_variables)} variables")
            final_document = self.replace_variables(content, all_variables)
            
            # Vérification des remplacements effectués
            final_preview = final_document[:100] + "..." if len(final_document) > 100 else final_document
            logger.debug(f"Document final (aperçu): {final_preview}")
            logger.debug(f"Longueur du document final: {len(final_document)} caractères")
            
            # Vérifier que le document final n'est pas vide
            if not final_document or final_document.strip() == "":
                logger.error("ERREUR CRITIQUE: Le document généré est vide!")
                return "ERREUR: Le document généré est vide. Veuillez vérifier le template."
            
            # Rechercher des variables non remplacées
            remaining_vars = []
            for pattern in [r'\{[a-zA-Z0-9_]+\}', r'\[[a-zA-Z0-9_]+\]', r'<[a-zA-Z0-9_]+>', r'\$[a-zA-Z0-9_]+\$']:
                matches = re.findall(pattern, final_document)
                if matches:
                    remaining_vars.extend(matches)
            
            if remaining_vars:
                logger.warning(f"Variables non remplacées: {remaining_vars}")
                
                # Deuxième passe avec personalize_document comme fallback
                logger.info("Deuxième passe avec personalize_document")
                final_document = self.personalize_document(final_document, all_variables)
                
                # Vérifier à nouveau pour les variables non remplacées
                remaining_vars = []
                for pattern in [r'\{[a-zA-Z0-9_]+\}', r'\[[a-zA-Z0-9_]+\]', r'<[a-zA-Z0-9_]+>', r'\$[a-zA-Z0-9_]+\$']:
                    matches = re.findall(pattern, final_document)
                    if matches:
                        remaining_vars.extend(matches)
                
                if remaining_vars:
                    logger.warning(f"Variables toujours non remplacées après deuxième passe: {remaining_vars}")
            
            # Fin du processus
            logger.info(f"=== FIN GÉNÉRATION DOCUMENT ({len(final_document)} caractères) ===")
            
            # Retourner le document final
            return final_document
            
        except Exception as e:
            logger.error(f"ERREUR lors de la génération du document: {str(e)}")
            traceback.print_exc()
            
            # En cas d'erreur, essayer une méthode très basique
            try:
                logger.info("Tentative de récupération avec méthode ultra-basique")
                content = self._read_file_safely(template_path)
                
                if not content:
                    return "ERREUR: Impossible de lire le template."
                    
                simple_doc = content
                
                # Remplacement basique au format {var}
                for var_name, value in all_variables.items():
                    value_str = str(value) if value is not None else ""
                    pattern = f"{{{var_name}}}"
                    simple_doc = simple_doc.replace(pattern, value_str)
                
                logger.info(f"Document récupéré avec méthode basique ({len(simple_doc)} caractères)")
                return simple_doc
                
            except Exception as recovery_error:
                logger.error(f"ÉCHEC TOTAL de la génération du document: {str(recovery_error)}")
                return f"ERREUR CRITIQUE: Impossible de générer le document. Détails: {str(e)}"
    
    def _complete_document_with_ai(self, content: str, variables: Dict[str, str], remaining_vars: List[str]) -> str:
        """
        Utilise l'IA pour compléter intelligemment le document avec toutes les variables
        
        Args:
            content: Contenu du document partiellement rempli
            variables: Toutes les variables disponibles
            remaining_vars: Liste des variables non remplacées
            
        Returns:
            Document complété par l'IA
        """
        # Formater pour plus de lisibilité
        remaining_formatted = ', '.join(remaining_vars[:10])
        if len(remaining_vars) > 10:
            remaining_formatted += f" et {len(remaining_vars) - 10} autres"
        
        # Préparation du prompt
        prompt = f"""Vous êtes un expert en génération de documents juridiques et administratifs.

DOCUMENT ACTUEL:
```
{content}
```

VARIABLES DISPONIBLES:
```
{json.dumps(variables, indent=2, ensure_ascii=False)}
```

VARIABLES NON REMPLACÉES: {remaining_formatted}

TÂCHE:
1. Analysez le document et identifiez toutes les variables non remplacées
2. Remplacez chaque variable par sa valeur correspondante dans le dictionnaire
3. Si une variable n'a pas de valeur correspondante, utilisez une valeur standard appropriée
4. Assurez-vous que le texte final est cohérent et complet
5. Conservez EXACTEMENT la mise en forme du document (espaces, retours à la ligne, etc.)

IMPORTANT:
- Ne retournez que le document complété, sans explication
- Respectez la structure exacte du document original
- Utilisez toujours les informations fournies dans le dictionnaire quand elles existent

DOCUMENT COMPLÉTÉ:
"""
        
        try:
            # Appel à Ollama avec un timeout plus long pour garantir l'achèvement
            response = self._call_ollama(prompt, timeout=180)
            
            # Vérifier si la réponse contient du texte
            if not response or len(response) < 50:
                logger.warning("Réponse de l'IA trop courte ou vide")
                return content
                
            return response
            
        except Exception as e:
            logger.error(f"Échec de la complétion IA: {e}")
            return content
    
    def _verify_document_quality(self, final_doc: str, original_doc: str, variables: Dict[str, str]) -> bool:
        """
        Vérifie la qualité du document généré pour s'assurer qu'il est complet
        
        Args:
            final_doc: Document généré
            original_doc: Document original
            variables: Variables utilisées
            
        Returns:
            True si le document est de bonne qualité, False sinon
        """
        # 1. Vérifier que le document n'est pas vide
        if not final_doc or len(final_doc) < 50:
            logger.warning("Document final trop court")
            return False
            
        # 2. Vérifier que le document a une taille raisonnable par rapport à l'original
        if len(final_doc) < len(original_doc) * 0.7:
            logger.warning(f"Document final trop court: {len(final_doc)} vs {len(original_doc)} caractères")
            return False
            
        # 3. Vérifier que les variables importantes sont présentes dans le document final
        important_vars = ['name', 'date', 'montant', 'reference', 'adresse', 'telephone']
        for var in important_vars:
            if var in variables and variables[var]:
                if variables[var] not in final_doc:
                    logger.warning(f"Variable importante '{var}' absente du document final")
                    return False
        
        # 4. Vérifier qu'il ne reste pas de motifs de variables
        for pattern in [r'\{[a-zA-Z0-9_]+\}', r'\[[a-zA-Z0-9_]+\]', r'<[a-zA-Z0-9_]+>', r'\$[a-zA-Z0-9_]+\$', r'«[a-zA-Z0-9_]+»']:
            if re.search(pattern, final_doc):
                logger.warning(f"Variables non remplacées trouvées dans le document final")
                return False
                
        return True

    def personalize_document(self, content: str, variables: Dict[str, Any]) -> str:
        """
        Personnalise un document en remplaçant toutes les variables par leurs valeurs
        Préserve exactement la structure originale du document
        
        Args:
            content: Contenu du document
            variables: Dictionnaire des variables à remplacer
            
        Returns:
            Document personnalisé
        """
        try:
            # Début du processus de personnalisation
            logger.info(f"Début de personnalisation du document avec {len(variables)} variables")
            
            # Afficher le contenu original pour le débogage
            logger.debug(f"Contenu original (premiers 200 caractères): {content[:200]}...")
            logger.debug(f"Taille du contenu: {len(content)} caractères")
            
            # Extraire les valeurs simples du dictionnaire variables
            values_dict = {}
            for key, value in variables.items():
                # Si c'est un dict (comme {"valeur": "xyz", "type": "..."}), extraire la valeur
                if isinstance(value, dict) and "valeur" in value:
                    values_dict[key] = str(value["valeur"]) if value["valeur"] is not None else ""
                else:
                    values_dict[key] = str(value) if value is not None else ""
            
            # Afficher les variables pour le débogage
            logger.debug(f"Variables à remplacer: {values_dict}")
            
            # Copier le contenu original pour le processus de remplacement
            result = content
            
            # Premier remplacement : supporter différents formats de variables courants
            for var_name, var_value in values_dict.items():
                # Convertir la valeur en string si nécessaire
                value_str = str(var_value) if var_value is not None else ""
                
                # Liste des formats de variables à remplacer
                formats = [
                    f"{{{var_name}}}",       # {variable}
                    f"[{var_name}]",         # [variable]
                    f"<{var_name}>",         # <variable>
                    f"${{{var_name}}}",      # ${variable}
                    f"${var_name}$",         # $variable$
                    f"«{var_name}»",         # «variable»
                    f"%{var_name}%",         # %variable%
                    f"[[{var_name}]]",       # [[variable]]
                    f"{{{{{var_name}}}}}"    # {{variable}}
                ]
                
                # Faire le remplacement pour chaque format
                for fmt in formats:
                    if fmt in result:
                        original_length = len(result)
                        result = result.replace(fmt, value_str)
                        new_length = len(result)
                        
                        # Afficher le résultat du remplacement
                        replacements = (original_length - new_length) // (len(fmt) - len(value_str)) if len(fmt) != len(value_str) else 0
                        if replacements > 0:
                            logger.debug(f"Remplacé '{fmt}' par '{value_str}' ({replacements} fois)")
                
                # Essayer aussi avec des variations de casse
                var_name_lower = var_name.lower()
                var_name_upper = var_name.upper()
                var_name_title = var_name.title()
                
                for name_variant in [var_name_lower, var_name_upper, var_name_title]:
                    if name_variant != var_name:
                        for fmt_template in ["{}", "[]", "<>", "${}", "${}$", "«{}»", "%{}%", "[[{}]]", "{{{}}}"]:
                            fmt = fmt_template.format(name_variant)
                            if fmt in result:
                                result = result.replace(fmt, value_str)
                                logger.debug(f"Remplacé (variation de casse) '{fmt}' par '{value_str}'")
            
            # Afficher les résultats du processus de remplacement
            logger.debug(f"Contenu final (premiers 200 caractères): {result[:200]}...")
            logger.info(f"Personnalisation terminée. Taille du document: {len(result)} caractères")
            
            # Vérifier si des variables n'ont pas été remplacées
            unprocessed_vars = []
            for pattern in [r'\{[a-zA-Z0-9_]+\}', r'\[[a-zA-Z0-9_]+\]', r'<[a-zA-Z0-9_]+>', r'\$[a-zA-Z0-9_]+\$']:
                matches = re.findall(pattern, result)
                if matches:
                    unprocessed_vars.extend(matches)
                    
            if unprocessed_vars:
                logger.warning(f"Variables non traitées: {unprocessed_vars}")
            
            return result
            
        except Exception as e:
            # En cas d'erreur, enregistrer le détail et retourner le document original
            logger.error(f"Erreur lors de la personnalisation du document: {str(e)}")
            return content

    def _clean_json_response(self, response: str) -> str:
        """
        Nettoie la réponse pour extraire uniquement le JSON valide
        """
        try:
            # Trouver le premier caractère { et le dernier }
            start = response.find('{')
            end = response.rfind('}') + 1
            if start >= 0 and end > start:
                return response[start:end]
            return response
        except Exception:
            return response

    def _call_ollama(self, prompt: str, timeout: int = 10, max_retries: int = 1, fast_mode: bool = True) -> str:
        """
        Appelle l'API Ollama avec optimisation pour la détection de variables
        
        Args:
            prompt: Prompt à envoyer à l'API
            timeout: Timeout en secondes (par défaut 10s)
            max_retries: Nombre maximum de tentatives en cas d'échec (par défaut 1)
            fast_mode: Utiliser le mode rapide avec moins de tokens (par défaut True)
            
        Returns:
            Réponse de l'API
        """
        try:
            # Limiter et optimiser le prompt
            if len(prompt) > 3000:
                prompt = prompt[:3000]
                
            # Options optimisées à partir du fichier de configuration
            config = self._load_config()
            options = config.get("options", {})
            
            # En mode rapide, limiter le nombre de tokens
            if fast_mode:
                options["num_predict"] = min(options.get("num_predict", 800), 300)
                
            # Récupérer modèle et URL depuis la configuration
            api_url = config.get("api_url", "http://localhost:11434/api/generate")
            model = config.get("model", "llama3:latest")
            
            # Ajouter un délai de sécurité pour le timeout réseau
            request_timeout = min(timeout + 3, 10)  # Timeout maximum de 10 secondes et 3 secondes de plus que le timeout API
            
            # Tentatives avec délai minimal entre chaque essai
            for attempt in range(max_retries + 1):
                try:
                    if attempt > 0:
                        logger.info(f"Nouvelle tentative ({attempt}/{max_retries})")
                    
                    # Appel optimisé à l'API sans stream pour maximiser la vitesse
                    response = requests.post(
                        api_url,
                        json={"model": model, "prompt": prompt, "stream": False, "options": options},
                        timeout=request_timeout
                    ).json()
                    
                    if "response" in response:
                        # Extraire la réponse
                        result = response["response"].strip()
                        return result
                        
                    return ""
                    
                except Exception as e:
                    logger.error(f"Erreur lors de l'appel à Ollama: {str(e)}")
                    if attempt < max_retries:
                        time.sleep(1)  # Délai minimal
            
            logger.error(f"Échec après {max_retries} tentatives")
            return ""
            
        except Exception as e:
            logger.error(f"Erreur Ollama (globale): {str(e)}")
            return ""

    def _analyze_document_by_sections(self, content: str) -> Dict[str, Any]:
        """
        Analyse un document en le divisant en sections pour un traitement plus efficace
        
        Args:
            content: Contenu du document à analyser
            
        Returns:
            Dict contenant les informations extraites
        """
        try:
            logger.info("Analyse du document par sections")
            
            # Évaluer la complexité globale du document
            complexity = self._estimate_document_complexity(content)
            logger.info(f"Complexité du document estimée: {complexity}")
            
            # Diviser le document en sections logiques
            sections = self._split_document_into_sections(content)
            logger.info(f"Document divisé en {len(sections)} sections")
            
            # Limiter le nombre de sections analysées pour accélérer le processus
            max_sections = 2 if complexity == "complex" else 3
            sections_to_analyze = sections[:max_sections]
            
            results = []
            retries_left = 1  # Réduit le nombre de retries pour accélérer
            
            # Analyser seulement les sections les plus pertinentes
            # Section 1: Souvent l'en-tête ou introduction
            if len(sections) > 0:
                # Réduire davantage la taille pour éviter les timeouts
                section = sections[0][:1000]  # 1000 caractères au lieu de 2000
                
                # Vérifier d'abord s'il y a des variables évidentes avec regex
                basic_vars = self._extract_basic_variables_from_text(section)
                if basic_vars:
                    logger.info(f"Variables extraites de l'en-tête par regex: {len(basic_vars)}")
                    results.append({"variables": basic_vars})
                else:
                    # Sinon essayer avec l'analyse IA
                    result = self._analyze_section_simplified(section)
                    if result and "variables" in result and result["variables"]:
                        logger.info(f"En-tête/Introduction: {len(result['variables'])} variables trouvées")
                        results.append(result)
            
            # Section 2: Souvent le corps du document 
            if len(sections) > 2 and len(results) < 2:  # N'analyser que si on n'a pas assez de résultats
                section = sections[2][:1000]  # 1000 caractères au lieu de 2000
                
                # Vérifier d'abord s'il y a des variables évidentes avec regex
                basic_vars = self._extract_basic_variables_from_text(section)
                if basic_vars:
                    logger.info(f"Variables extraites du corps par regex: {len(basic_vars)}")
                    results.append({"variables": basic_vars})
                else:
                    # Sinon essayer avec l'analyse IA
                    result = self._analyze_section_simplified(section)
                    if result and "variables" in result and result["variables"]:
                        logger.info(f"Corps du document: {len(result['variables'])} variables trouvées")
                        results.append(result)
            
            # Fusionner les résultats
            merged_result = self._merge_results(results)
            
            # Fallback rapide si pas assez de variables
            if len(merged_result.get("variables", {})) < 3:
                logger.warning("Peu de variables trouvées, extraction directe")
                basic_vars = self._extract_basic_variables_from_text(content[:5000])
                if basic_vars:
                    merged_result["variables"] = merged_result.get("variables", {})
                    merged_result["variables"].update(basic_vars)
            
            # Ajouter des métadonnées
            merged_result["_meta"] = {
                "sections_count": len(sections),
                "sections_analyzed": min(max_sections, len(sections)),
                "complexity": complexity,
                "extraction_method": "sections",
            }
            
            return merged_result
            
        except Exception as e:
            logger.error(f"Erreur lors de l'analyse par sections: {str(e)}")
            # Fallback vers une méthode plus simple
            variables = self._extract_basic_variables_from_text(content[:3000])
            return {"variables": variables, "_meta": {"extraction_method": "fallback_basic"}}

    def _analyze_section_simplified(self, content: str) -> Dict[str, Any]:
        """
        Analyse une section de document de manière simplifiée pour en extraire les variables
        
        Args:
            content: Section de texte à analyser
            
        Returns:
            Dict contenant les variables extraites
        """
        try:
            # Estimer la complexité de la section pour ajuster le timeout
            complexity = self._estimate_document_complexity(content)
            timeout = 10 if complexity == "complex" else 7  # Réduire les timeouts pour éviter les erreurs
            
            # Optimisation: réduire le contenu pour cibler l'essentiel
            if len(content) > 1000:
                content = content[:1000]  # Réduire à 1000 caractères maximum
            
            # Prompt en format texte simple au lieu de JSON
            prompt = f"""
Liste les variables et leurs valeurs dans ce document, une par ligne:
VARIABLE: valeur

Exemple:
NOM: Jean Dupont
DATE: 12/03/2024

Texte:
{content}
"""
            # Appel à l'IA
            response = self._call_ollama(prompt, timeout=timeout, fast_mode=True)
            if not response:
                # Utiliser extraction basique si pas de réponse
                variables = self._extract_basic_variables_from_text(content)
                return {"variables": variables}
            
            # Parser la réponse en format texte simple
            variables = {}
            lines = response.strip().split('\n')
            for line in lines:
                if ':' in line:
                    parts = line.split(':', 1)
                    if len(parts) == 2:
                        key = parts[0].strip().lower()
                        value = parts[1].strip()
                        if key and value:
                            variables[key] = value
            
            # Si aucune variable n'est trouvée, essayer l'extraction basique
            if not variables:
                variables = self._extract_basic_variables_from_text(content)
            
            return {"variables": variables}
                
        except Exception as e:
            logger.error(f"Erreur globale dans _analyze_section_simplified: {str(e)}")
            # En cas d'échec, utiliser l'extraction basique par regex
            variables = self._extract_basic_variables_from_text(content)
            return {"variables": variables}
    
    def _extract_basic_variables_from_text(self, text: str) -> Dict[str, Any]:
        """
        Extrait des variables basiques d'un texte sans utiliser l'IA
        
        Args:
            text: Texte brut dont on veut extraire les variables
            
        Returns:
            Dictionnaire des variables extraites
        """
        variables = {}
        
        # Normalisation du texte
        normalized_text = text.replace('\r\n', '\n').replace('\t', ' ')
        
        # Définition des patterns de reconnaissance
        patterns = {
            # Contact
            'email': r'\b([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\b',
            'telephone': r'(?:\+?\d{1,3}[.\s-]?)?(?:\(?\d{2,3}\)?[.\s-]?)?\d{2,4}[.\s-]?\d{2,4}[.\s-]?\d{2,4}',
            
            # Identifiants et références
            'reference': r'(?:ref|réf|référence|n°|numéro)\s*:?\s*([A-Z0-9-_/]+)',
            
            # Données financières
            'montant': r'\b\d+(?:[,.]\d{1,2})?\s*(?:€|USD|£|\$|EUR)?\b',
            
            # Dates
            'date': r'\b\d{1,2}[/.-]\d{1,2}[/.-]\d{2,4}\b|\b\d{4}[-]\d{1,2}[-]\d{1,2}\b',
            
            # Sites web
            'site_web': r'\b(?:https?://|www\.)[a-zA-Z0-9.-]+\.[a-zA-Z]{2,6}(?:/[^\s]*)?'
        }
        
        # Extraction des données
        for var_name, pattern in patterns.items():
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                if isinstance(matches[0], tuple):
                    for match in matches[0]:
                        if match:
                            variables[var_name] = match.strip()
                            break
                else:
                    variables[var_name] = matches[0].strip()
        
        # Nettoyage des valeurs
        for key in variables:
            if isinstance(variables[key], str):
                variables[key] = re.sub(r'^\W+|\W+$', '', variables[key]).strip()
                variables[key] = re.sub(r'\s+', ' ', variables[key])
        
        # Reformatage du résultat
        result = {}
        for key, value in variables.items():
            var_type = self._guess_variable_type(key, normalized_text)
            result[key] = {
                "type": var_type,
                "description": self._get_variable_description(key),
                "valeur": value
            }
        
        return result

    def _merge_results(self, results: List[Dict]) -> Dict[str, Any]:
        merged = {}
        for result in results:
            for key, value in result.items():
                if key not in merged:
                    merged[key] = value
        return merged

    def _analyze_fallback(self, content: str) -> Dict[str, Any]:
        """
        Méthode de secours pour l'analyse de document lorsque l'analyse principale échoue
        
        Args:
            content: Contenu textuel du document à analyser
            
        Returns:
            Dict contenant les informations extraites
        """
        logger.info("Utilisation de la méthode de secours pour l'analyse")
        try:
            # 1. Détection préliminaire des variables évidentes (templates)
            template_variables = self._detect_template_variables(content)
            
            # 2. Utiliser l'IA avec un timeout court
            document_complexity = self._estimate_document_complexity(content)
            timeout = 20 if document_complexity == "complex" else 15
            
            # Prompt en format texte simple pour plus de rapidité
            prompt = f"""
Analyse ce document et identifie les variables personnalisables.
Liste chaque variable avec sa valeur sur une ligne séparée, au format VARIABLE: VALEUR.

Exemple:
NOM: Jean Dupont
ADRESSE: 15 rue des Fleurs
CODE_POSTAL: 75001
VILLE: Paris
MONTANT: 150,50€
DATE: 12/03/2024

Document:
---
{content[:3000]}
---
"""
            # Essayer avec timeout adapté
            response = self._call_ollama(prompt, timeout=timeout, max_retries=2)
            
            if not response:
                # Si pas de réponse, retourner les variables du template et l'extraction basique
                logger.warning("Pas de réponse pour l'analyse de secours, utilisation des extractions basiques")
                basic_variables = self._extract_basic_variables_from_text(content[:5000])
                
                # Fusionner avec les variables de template
                all_variables = {**template_variables, **basic_variables}
                return {"variables": all_variables}
            
            # Parser la réponse en format texte
            extracted_variables = {}
            for line in response.strip().split('\n'):
                if ':' in line:
                    parts = line.split(':', 1)
                    if len(parts) == 2:
                        key = parts[0].strip().lower()
                        value = parts[1].strip()
                        if key and value and len(key) > 1 and len(value) > 1:
                            extracted_variables[key] = value
            
            # Si aucune variable n'est extraite de la réponse, revenir aux méthodes basiques
            if not extracted_variables:
                logger.warning("Aucune variable extraite de la réponse, recours aux méthodes basiques")
                basic_variables = self._extract_basic_variables_from_text(content[:5000])
                extracted_variables = basic_variables
                
            # Fusionner les variables du template avec celles extraites
            all_variables = {**template_variables, **extracted_variables}
            
            return {"variables": all_variables}
            
        except Exception as e:
            logger.error(f"Erreur lors de l'analyse de secours: {str(e)}")
            # En dernier recours, utiliser uniquement l'extraction basique
            basic_variables = self._extract_basic_variables_from_text(content[:5000])
            return {"variables": basic_variables}

    def _estimate_document_complexity(self, content: str) -> str:
        """
        Estime la complexité d'un document pour ajuster les paramètres d'analyse
        
        Args:
            content: Contenu du document
            
        Returns:
            str: "simple", "medium" ou "complex"
        """
        # Calculer des indicateurs de complexité
        char_count = len(content)
        line_count = content.count('\n') + 1
        word_count = len(content.split())
        
        # Indicateurs de structure et contenu complexe
        table_indicators = content.count('|') + content.count(';') + content.count('\t')
        number_count = len(re.findall(r'\d+', content))
        special_char_count = len(re.findall(r'[€$£%&*()[\]{}]', content))
        
        # Calculer un score de complexité
        complexity_score = 0
        
        # Longueur du document
        if char_count > 10000:
            complexity_score += 3
        elif char_count > 5000:
            complexity_score += 2
        elif char_count > 2000:
            complexity_score += 1
            
        # Densité d'information
        if word_count > 0:
            chars_per_word = char_count / word_count
            if chars_per_word > 8:  # Mots longs/complexes
                complexity_score += 2
            elif chars_per_word > 6:
                complexity_score += 1
                
        # Structure
        if table_indicators > 50:  # Document très structuré (tables, CSV)
            complexity_score += 3
        elif table_indicators > 20:
            complexity_score += 2
            
        # Données numériques
        if number_count > 100:  # Beaucoup de chiffres (factures, données financières)
            complexity_score += 2
        elif number_count > 50:
            complexity_score += 1
            
        # Caractères spéciaux
        if special_char_count > 100:
            complexity_score += 1
            
        # Variété lexicale (estimation simple)
        unique_words = len(set(re.findall(r'\b\w+\b', content.lower())))
        if unique_words > 300:
            complexity_score += 2
        elif unique_words > 150:
            complexity_score += 1
            
        # Déterminer la catégorie finale
        if complexity_score >= 6:
            return "complex"
        elif complexity_score >= 3:
            return "medium"
        else:
            return "simple"

    def _detect_template_variables(self, content: str) -> Dict[str, str]:
        """
        Détecte les variables évidentes de type template dans le contenu
        
        Args:
            content: Contenu textuel à analyser
            
        Returns:
            Dict contenant les variables de template détectées
        """
        logger.info("Détection des variables de template")
        variables = {}
        
        # Modèles de variables de template
        template_patterns = [
            r'\{\{([a-zA-Z0-9_]+)\}\}',  # Format {{variable}}
            r'\{([a-zA-Z0-9_]+)\}',      # Format {variable}
            r'\[([a-zA-Z0-9_]+)\]',      # Format [variable]
            r'<<([a-zA-Z0-9_]+)>>',      # Format <<variable>>
            r'<([a-zA-Z0-9_]+)>',        # Format <variable>
            r'\$([a-zA-Z0-9_]+)\$',      # Format $variable$
            r'%([a-zA-Z0-9_]+)%',        # Format %variable%
            r'«([a-zA-Z0-9_]+)»'         # Format «variable»
        ]
        
        for pattern in template_patterns:
            matches = re.findall(pattern, content)
            for var_name in matches:
                var_name = var_name.strip()
                # Ignorer les noms de variables courts ou communs
                if len(var_name) > 2 and var_name.lower() not in ['if', 'for', 'end', 'else', 'var', 'let']:
                    variables[var_name] = f"<{var_name}>"
        
        logger.info(f"Détection de {len(variables)} variables de template")
        return variables
    
    def _extract_variables_with_scoring(self, content: str) -> Tuple[Dict[str, str], Dict[str, float]]:
        """
        Extrait les variables du texte en utilisant des regex et attribue un score de confiance
        
        Args:
            content: Contenu textuel à analyser
            
        Returns:
            Tuple[Dict[str, str], Dict[str, float]]: Variables extraites et leurs scores de confiance
        """
        logger.info("Extraction de variables par regex avec scoring")
        variables = {}
        scores = {}  # Scores de confiance (0.0 à 1.0)
        content_lower = content.lower()
        
        # 1. INFORMATIONS PERSONNELLES
        
        # Civilités (M., Mme, Dr, etc.)
        civility_patterns = [
            (r'\b(M\.|Monsieur|Mme\.?|Madame|Mlle\.?|Mademoiselle|Dr\.?|Docteur|Pr\.?|Professeur)\s+([A-Z][a-zÀ-ÿ]+(?:\s+[A-Z][a-zÀ-ÿ]+)*)', 0.8),
            (r'\b([A-Z][a-zÀ-ÿ]+(?:\s+[A-Z][a-zÀ-ÿ]+)*)\s+(M\.|Monsieur|Mme\.?|Madame|Mlle\.?|Mademoiselle|Dr\.?|Docteur|Pr\.?|Professeur)', 0.7)
        ]
        
        for pattern, base_score in civility_patterns:
            matches = re.findall(pattern, content)
            if matches:
                civility_map = {
                    'monsieur': 'M.', 'm.': 'M.',
                    'madame': 'Mme', 'mme': 'Mme', 'mme.': 'Mme',
                    'mademoiselle': 'Mlle', 'mlle': 'Mlle', 'mlle.': 'Mlle',
                    'docteur': 'Dr', 'dr': 'Dr', 'dr.': 'Dr',
                    'professeur': 'Pr', 'pr': 'Pr', 'pr.': 'Pr'
                }
                
                if len(matches[0]) >= 2:
                    # Format: (civilité, nom)
                    civ, nom = matches[0]
                    civ_lower = civ.lower().replace('.', '')
                    if civ_lower in civility_map:
                        variables['civilite'] = civility_map[civ_lower]
                        scores['civilite'] = base_score + 0.1
                    else:
                        variables['civilite'] = civ
                        scores['civilite'] = base_score
                    variables['nom'] = nom.strip()
                    scores['nom'] = base_score
                break
                
        # Noms complets (avec ou sans civilité)
        if 'nom' not in variables:
            name_patterns = [
                (r'(?:nom|client)(?:\s*:|\s*du\s*client\s*:|\s*)\s*([A-Z][a-zÀ-ÿ]+(?:[\s-][A-Z][a-zÀ-ÿ]+)*)', 0.9),
                (r'\b([A-Z][A-Z\s-]+)\b', 0.6),  # Noms en majuscules (moins fiable)
                (r'(?:je soussigné\(e\)|signataire)\s*:?\s*([A-Z][a-zÀ-ÿ]+(?:[\s-][A-Z][a-zÀ-ÿ]+)*)', 0.85)
            ]
            
            for pattern, base_score in name_patterns:
                matches = re.findall(pattern, content, re.IGNORECASE)
                if matches:
                    for match in matches:
                        name = match.strip()
                        # Vérifier que ce n'est pas un terme commun
                        common_terms = ['FACTURE', 'DEVIS', 'CONTRAT', 'DOCUMENT', 'CLIENT']
                        if name.upper() not in common_terms and len(name) > 3:
                            variables['nom'] = name
                            # Ajuster le score selon le contexte et la qualité
                            context_score = 0.0
                            if "nom" in content_lower[:content_lower.find(name.lower())]:
                                context_score += 0.1
                            if name.isupper() and pattern != r'\b([A-Z][A-Z\s-]+)\b':
                                context_score -= 0.1  # Pénalité pour les noms tout en majuscules sauf si c'est le pattern dédié
                            
                            scores['nom'] = min(1.0, base_score + context_score)
                            break
                    if 'nom' in variables:
                        break
        
        # Ajouter le reste des sections (prénoms, société, adresse, etc.) avec scoring
        # Pour chaque section, définir des patterns avec un score de base et ajuster 
        # le score final en fonction du contexte
        
        # Par exemple pour les montants:
        amount_patterns = [
            (r'(?:montant|total|prix|somme)(?:\s*:|\s*)\s*([0-9\s]+(?:[,.][0-9]{1,3})?)\s*(?:€|EUR|euros?|USD|\$|£|HT|TTC)', 0.9),
            (r'([0-9\s]+(?:[,.][0-9]{1,3})?\s*(?:€|EUR|euros?|USD|\$|£))', 0.75),
            (r'(?:HT|hors taxes?)(?:\s*:|\s*)\s*([0-9\s]+(?:[,.][0-9]{1,3})?)', 0.85),
            (r'(?:TTC|toutes taxes comprises?)(?:\s*:|\s*)\s*([0-9\s]+(?:[,.][0-9]{1,3})?)', 0.85),
            (r'(?:TVA|taux de TVA)(?:\s*:|\s*)\s*([0-9\s]+(?:[,.][0-9]{1,3})?\s*\%?)', 0.85)
        ]
        
        for pattern, base_score in amount_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            if matches:
                amount_value = matches[0].strip() if isinstance(matches[0], str) else matches[0]
                
                # Déterminer le type de montant en fonction du pattern
                var_name = ""
                if 'montant' in pattern.lower() or 'total' in pattern.lower() or 'prix' in pattern.lower() or 'somme' in pattern.lower():
                    var_name = 'montant'
                elif 'HT' in pattern or 'hors taxes' in pattern.lower():
                    var_name = 'montant_ht'
                elif 'TTC' in pattern or 'toutes taxes' in pattern.lower():
                    var_name = 'montant_ttc'
                elif 'TVA' in pattern:
                    var_name = 'tva'
                
                if var_name:
                    variables[var_name] = amount_value
                    
                    # Ajuster le score selon le format et la qualité
                    context_score = 0.0
                    # Un montant bien formaté avec devise augmente la confiance
                    if re.search(r'[0-9]+[,.][0-9]{2}\s*(?:€|EUR)', amount_value):
                        context_score += 0.1
                    # Des caractères non numériques inattendus réduisent la confiance
                    if re.search(r'[^0-9\s,.€$£HTCVAEURG]', amount_value):
                        context_score -= 0.2
                        
                    scores[var_name] = min(1.0, base_score + context_score)
        
        # Adaptation similaire pour les autres types de variables
        # Par exemple, ici pour dates, emails, téléphones, etc.
        
        # Code complet du _extract_variables_with_regex mais avec des scores
        
        # Le reste de la méthode _extract_variables_with_regex à adapter avec scores
        
        # Reproduire la logique pour les autres parties (téléphones, emails, etc.)
        # en suivant la même structure: patterns avec scores de base, 
        # puis ajustement selon le contexte

        logger.info(f"Extraction par regex terminée, {len(variables)} variables extraites avec scoring")
        return variables, scores

    def _guess_variable_type(self, key: str, content: str = None) -> str:
        """
        Devine le type d'une variable en fonction de son nom et contenu
        
        Args:
            key: Nom de la variable
            content: Contenu pour analyse contextuelle (optionnel)
            
        Returns:
            Type de la variable (text, number, date, etc.)
        """
        # Mapping des clés vers les types
        type_mapping = {
            # Dates
            'date': 'date',
            'date_emission': 'date',
            'date_echeance': 'date',
            'date_livraison': 'date',
            
            # Montants et nombres
            'montant': 'number',
            'montant_ht': 'number',
            'montant_ttc': 'number',
            'tva': 'number',
            
            # Contacts
            'telephone': 'phone',
            'telephone_mobile': 'phone',
            'telephone_fixe': 'phone',
            'email': 'email',
            
            # Adresses
            'adresse': 'address',
            'code_postal': 'postal_code',
            'ville': 'city',
            'pays': 'country',
            
            # Identifiants
            'siret': 'id',
            'siren': 'id',
            'tva_intra': 'id',
            'reference': 'id',
            'numero_facture': 'id',
            'numero_client': 'id'
        }
        
        # Analyse du contenu de la valeur si disponible
        if key not in type_mapping and content:
            if re.search(r'\b\d{1,2}[/.-]\d{1,2}[/.-]\d{2,4}\b', content):
                return 'date'
            elif re.search(r'\b\d+(?:[,.]\d{1,3})?\s*(?:€|EUR|USD|\$|£)\b', content):
                return 'number'
            elif re.search(r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b', content):
                return 'email'
            elif re.search(r'\b(?:\+?\d{1,3}[.\s-]?)?(?:\(?\d{2,3}\)?[.\s-]?)?\d{2,4}[.\s-]?\d{2,4}[.\s-]?\d{2,4}\b', content):
                return 'phone'
        
        # Retour du type ou par défaut "text"
        return type_mapping.get(key, 'text')

    def _get_variable_type(self, key: str) -> str:
        """
        Détermine le type d'une variable en fonction de sa clé
        
        Args:
            key: Nom de la variable
            
        Returns:
            Type de la variable
        """
        type_mapping = {
            'nom': 'identite', 'prenom': 'identite', 'civilite': 'identite',
            'societe': 'organisation', 'entreprise': 'organisation', 
            'adresse': 'coordonnees', 'ville': 'coordonnees', 'code_postal': 'coordonnees', 'pays': 'coordonnees',
            'email': 'contact', 'telephone': 'contact', 'telephone_mobile': 'contact', 'telephone_fixe': 'contact',
            'date': 'temporel', 'date_emission': 'temporel', 'date_echeance': 'temporel', 'date_livraison': 'temporel',
            'montant': 'financier', 'montant_ht': 'financier', 'montant_ttc': 'financier', 'tva': 'financier',
            'reference': 'identifiant', 'numero_facture': 'identifiant', 'numero_client': 'identifiant',
            'type_document': 'categorie', 'objet': 'contenu', 'description': 'contenu'
        }
        return type_mapping.get(key, 'autre')
    
    def _get_variable_description(self, key: str) -> str:
        """
        Fournit une description pour une variable en fonction de sa clé
        
        Args:
            key: Nom de la variable
            
        Returns:
            Description de la variable
        """
        description_mapping = {
            'nom': 'Nom de famille ou nom complet',
            'prenom': 'Prénom de la personne',
            'civilite': 'Civilité (M., Mme, etc.)',
            'societe': 'Nom de la société ou organisation',
            'entreprise': 'Nom de l\'entreprise',
            'adresse': 'Adresse postale',
            'ville': 'Ville de l\'adresse',
            'code_postal': 'Code postal',
            'pays': 'Pays de l\'adresse',
            'email': 'Adresse email principale',
            'telephone': 'Numéro de téléphone',
            'telephone_mobile': 'Numéro de téléphone mobile',
            'telephone_fixe': 'Numéro de téléphone fixe',
            'date': 'Date générique mentionnée dans le document',
            'date_emission': 'Date d\'émission du document',
            'date_echeance': 'Date d\'échéance de paiement',
            'date_livraison': 'Date de livraison prévue',
            'montant': 'Montant global',
            'montant_ht': 'Montant hors taxes',
            'montant_ttc': 'Montant toutes taxes comprises',
            'tva': 'Montant ou taux de TVA',
            'reference': 'Référence du document',
            'numero_facture': 'Numéro de facture',
            'numero_client': 'Numéro ou identifiant client',
            'type_document': 'Type ou catégorie du document',
            'objet': 'Objet ou sujet du document',
            'description': 'Description ou détails'
        }
        return description_mapping.get(key, f'Valeur de {key}')

    def diagnose_analysis_issues(self, file_path: str) -> Dict[str, Any]:
        """
        Diagnostique les problèmes potentiels dans l'analyse d'un document
        
        Args:
            file_path: Chemin vers le document à analyser
            
        Returns:
            Dict contenant les diagnostics et informations sur les problèmes
        """
        logger.info(f"Diagnostic de l'analyse pour le fichier: {file_path}")
        
        diagnostics = {
            "file_path": file_path,
            "file_exists": os.path.exists(file_path),
            "file_size": os.path.getsize(file_path) if os.path.exists(file_path) else 0,
            "file_extension": os.path.splitext(file_path)[1].lower() if os.path.exists(file_path) else "",
            "text_extraction": None,
            "text_length": 0,
            "llama3_connection": False,
            "llama3_response": None,
            "issues": [],
            "recommendations": []
        }
        
        # Vérification du fichier
        if not diagnostics["file_exists"]:
            diagnostics["issues"].append("Le fichier n'existe pas")
            diagnostics["recommendations"].append("Vérifiez le chemin du fichier")
            return diagnostics
            
        if diagnostics["file_size"] == 0:
            diagnostics["issues"].append("Le fichier est vide")
            diagnostics["recommendations"].append("Vérifiez que le fichier contient des données")
            return diagnostics
            
        # Extraction de texte
        try:
            content = self.convert_to_text(file_path)
            diagnostics["text_extraction"] = "success"
            diagnostics["text_length"] = len(content)
            
            # Échantillon du texte
            diagnostics["text_sample"] = content[:200] + "..." if len(content) > 200 else content
            
            if diagnostics["text_length"] == 0:
                diagnostics["issues"].append("Aucun texte n'a pu être extrait du document")
                diagnostics["recommendations"].append("Vérifiez que le document contient du texte et pas seulement des images")
                
        except Exception as e:
            diagnostics["text_extraction"] = "failed"
            diagnostics["issues"].append(f"Erreur lors de l'extraction du texte: {str(e)}")
            diagnostics["recommendations"].append("Vérifiez le format du document et réessayez avec un autre format si possible")
            
        # Test de la connexion à Phi-3
        try:
            test_prompt = "Renvoie simplement le mot 'OK'"
            response = self._call_ollama(test_prompt, timeout=3)
            
            diagnostics["llama3_connection"] = "OK" in response
            diagnostics["llama3_response"] = response[:50] if response else "Pas de réponse"
            
            if not diagnostics["llama3_connection"]:
                diagnostics["issues"].append("Problème de connexion avec le modèle Llama3")
                diagnostics["recommendations"].append("Vérifiez que le serveur Ollama est en cours d'exécution et que le modèle llama3 est installé")
                
        except Exception as e:
            diagnostics["issues"].append(f"Erreur lors du test de connexion à Llama3: {str(e)}")
            diagnostics["recommendations"].append("Vérifiez la configuration du serveur Ollama et assurez-vous qu'il fonctionne correctement")
            
        # Analyse conditionnelle du contenu
        if diagnostics["text_extraction"] == "success" and diagnostics["text_length"] > 0:
            # Vérification pour les formats structurés/semi-structurés
            if diagnostics["file_extension"] in [".pdf", ".docx", ".odt"]:
                if content.count("\n") < 5:
                    diagnostics["issues"].append("Le document semble manquer de structure (peu de sauts de ligne)")
                    diagnostics["recommendations"].append("Vérifiez que la mise en forme du document est préservée")
            
            # Vérification des caractères spéciaux/problèmes d'encodage
            strange_chars_count = len(re.findall(r'[^\x00-\x7F\u00C0-\u00FF]', content))
            if strange_chars_count > len(content) * 0.1:  # Plus de 10% de caractères étranges
                diagnostics["issues"].append("Le document contient beaucoup de caractères non standard")
                diagnostics["recommendations"].append("Vérifiez l'encodage du document")
            
        # Ajout d'un statut global
        if not diagnostics["issues"]:
            diagnostics["status"] = "ok"
            diagnostics["recommendations"].append("Aucun problème détecté, l'analyse devrait fonctionner correctement")
        else:
            diagnostics["status"] = "issues_detected"
            
        logger.info(f"Diagnostic terminé, statut: {diagnostics['status']}")
        return diagnostics

    def _split_document_into_sections(self, content: str) -> List[str]:
        """
        Divise un document en sections logiques pour analyse
        
        Args:
            content: Contenu du document à diviser
            
        Returns:
            Liste des sections du document
        """
        try:
            logger.info("Division du document en sections pour analyse")
            
            # Si le contenu est très court, le traiter comme une seule section
            if len(content) < 500:
                return [content]
                
            # Tentative de division basée sur les titres et les sauts de ligne
            sections = []
            
            # Identifier les marqueurs de séparation de sections
            section_markers = [
                # Titres numérotés (1. Introduction, 1.1 Contexte, etc.)
                r'\n\s*(?:\d+\.)+\s*[A-Z][^\n]{3,}',
                # Titres en majuscules
                r'\n\s*[A-Z][A-Z\s]{5,}[A-Z]\s*\n',
                # Titres avec formatage spécial (encadrés par des astérisques, soulignés, etc.)
                r'\n\s*[\*\-=_]{2,}[^\n]{3,}[\*\-=_]{2,}',
                # Séparateurs horizontaux
                r'\n\s*[\-=_\*]{3,}\s*\n',
                # Doubles sauts de ligne suivis d'un mot en majuscule
                r'\n\s*\n\s*[A-Z][^\n]{3,}',
                # Mots-clés spécifiques qui indiquent souvent un changement de section
                r'\n(?:ARTICLE|ANNEXE|SECTION|CHAPITRE|PARTIE)[^\n]*'
            ]
            
            # Combiner tous les marqueurs en une seule expression régulière
            combined_pattern = '|'.join(section_markers)
            
            # Trouver tous les points de séparation
            split_points = []
            for m in re.finditer(combined_pattern, content):
                split_points.append(m.start())
                
            # Ajouter le début et la fin du document
            split_points = [0] + split_points + [len(content)]
            
            # Créer les sections
            for i in range(len(split_points) - 1):
                start = split_points[i]
                end = split_points[i + 1]
                section_content = content[start:end].strip()
                
                # Ne pas ajouter de sections vides
                if section_content:
                    sections.append(section_content)
                    
            # Si aucune section n'a été identifiée, diviser le document en paragraphes
            if len(sections) <= 1:
                paragraphs = re.split(r'\n\s*\n', content)
                
                # Regrouper les paragraphes en sections d'environ 500 caractères
                current_section = ""
                for para in paragraphs:
                    if para.strip():
                        if len(current_section) + len(para) > 500:
                            if current_section:
                                sections.append(current_section)
                                current_section = para
                            else:
                                sections.append(para)
                        else:
                            current_section += "\n\n" + para if current_section else para
                
                # Ajouter la dernière section
                if current_section:
                    sections.append(current_section)
            
            # Si le document est très long mais qu'on n'a pas trouvé beaucoup de sections,
            # diviser en morceaux de taille similaire
            if len(content) > 2000 and len(sections) < 3:
                # Réinitialiser les sections
                sections = []
                
                # Diviser en sections de 500 caractères environ
                section_size = 500
                for i in range(0, len(content), section_size):
                    section = content[i:i + section_size]
                    if section.strip():
                        sections.append(section)
                        
            logger.info(f"Document divisé en {len(sections)} sections")
            return sections
                
        except Exception as e:
            logger.error(f"Erreur lors de la division du document en sections: {e}")
            # En cas d'erreur, retourner le document entier comme une seule section
            return [content]

    def _fallback_personalization(self, content: str, variables: Dict[str, Any]) -> str:
        """
        Méthode de personnalisation de secours pour remplacer les variables par leurs valeurs
        tout en préservant la structure exacte du document.
        
        Args:
            content: Contenu du document à personnaliser
            variables: Dictionnaire des variables à remplacer
            
        Returns:
            Document personnalisé avec sa structure originale préservée
        """
        logger.info("Utilisation de la méthode de personnalisation de secours")
        try:
            result = content
            
            # Préparer les variables pour le remplacement
            replace_dict = {}
            for var_name, var_info in variables.items():
                value = ""
                
                # Extraire la valeur selon le format (dict ou valeur simple)
                if isinstance(var_info, dict):
                    if "valeur" in var_info:
                        value = var_info["valeur"]
                    elif "current_value" in var_info:
                        value = var_info["current_value"]
                else:
                    value = str(var_info)
                    
                if value:
                    replace_dict[var_name] = value
            
            # Première passe: utiliser replace_variables pour un remplacement standard
            result = self.replace_variables(result, replace_dict)
            
            # Deuxième passe: rechercher des variables qui pourraient avoir été manquées
            var_pattern_formats = [
                r'\{([a-zA-Z0-9_]+)\}',        # {variable}
                r'\{\{([a-zA-Z0-9_]+)\}\}',    # {{variable}}
                r'\[([a-zA-Z0-9_]+)\]',        # [variable]
                r'\[\[([a-zA-Z0-9_]+)\]\]',    # [[variable]]
                r'<([a-zA-Z0-9_]+)>',          # <variable>
                r'\$\{([a-zA-Z0-9_]+)\}',      # ${variable}
                r'%([a-zA-Z0-9_]+)%',          # %variable%
                r'\$([a-zA-Z0-9_]+)\$',        # $variable$
                r'«([a-zA-Z0-9_]+)»'           # «variable»
            ]
            
            for pattern_format in var_pattern_formats:
                matches = re.findall(pattern_format, result)
                for var_name in matches:
                    var_name_lower = var_name.lower()
                    
                    # Chercher si on a cette variable exacte ou sous une autre casse
                    value = None
                    for key, val in replace_dict.items():
                        if key.lower() == var_name_lower:
                            value = val
                            break
                    
                    # Si on a trouvé une valeur, remplacer cette occurrence spécifique
                    if value:
                        pattern_with_var = pattern_format.replace('([a-zA-Z0-9_]+)', re.escape(var_name))
                        result = re.sub(pattern_with_var, value, result)
            
            # Troisième passe: essayer des correspondances approximatives pour des variables connues
            common_mappings = {
                "nom_complet": ["nom", "prenom"],
                "adresse_complete": ["adresse", "code_postal", "ville"],
                "contact": ["telephone", "email"],
                "identifiant": ["reference", "numero_client"]
            }
            
            for composite_var, source_vars in common_mappings.items():
                # Chercher si le motif composite existe dans le document
                for pattern_format in var_pattern_formats:
                    pattern_with_var = pattern_format.replace('([a-zA-Z0-9_]+)', re.escape(composite_var))
                    if re.search(pattern_with_var, result):
                        # Construire une valeur composée à partir des variables sources
                        composite_value = ""
                        for source_var in source_vars:
                            for key, val in replace_dict.items():
                                if key.lower() == source_var:
                                    if composite_value:
                                        composite_value += " "
                                    composite_value += val
                                    break
                        
                        # Si on a construit une valeur composée, remplacer
                        if composite_value:
                            result = re.sub(pattern_with_var, composite_value, result)
            
            logger.info("Personnalisation de secours terminée")
            return result
        
        except Exception as e:
            logger.error(f"Erreur lors de la personnalisation de secours: {e}")
            # En dernier recours, retourner le document original
            return content

    def extract_template_variables(self, content: str) -> List[str]:
        """
        Extrait uniquement les noms des variables à remplacer dans un template
        
        Args:
            content: Contenu du document
        
        Returns:
            Liste des noms de variables
        """
        logger.info("Extraction des variables de template")
        
        # Réduire le contenu pour l'analyse
        if len(content) > 3000:
            content = content[:3000]
        
        # Extraction par IA avec prompt optimisé
        prompt = f"""Identifie les champs qui peuvent être remplacés dans ce template de document.
Cherche uniquement les éléments qui varieraient si ce document était utilisé pour une autre personne.

IMPORTANT: Je veux SEULEMENT les noms des champs (comme "nom", "adresse"), PAS leurs valeurs actuelles.

Liste uniquement les noms des champs à remplacer (par exemple: nom, prénom, email, téléphone, date, etc.).

DOCUMENT:
{content}
"""
        
        try:
            response = self._call_ollama(prompt, timeout=15)
            if response:
                # Extraire les variables à partir de la réponse
                # Rechercher les motifs comme "nom", "adresse", "email", etc.
                variables = re.findall(r'[*-]?\s*([a-zA-Z_éèêëàâäôöûüçÉÈÊËÀÂÄÔÖÛÜÇ]+(?:\s+[a-zA-Z_éèêëàâäôöûüçÉÈÊËÀÂÄÔÖÛÜÇ]+)*)', response)
                # Nettoyer les résultats
                variables = [v.strip().lower() for v in variables if len(v.strip()) > 2]
                
                # Filtrer les mots communs qui ne sont pas des variables
                common_words = ["liste", "exemple", "variables", "champs", "document", "template", "important"]
                variables = [v for v in variables if v not in common_words]
                
                if variables:
                    logger.info(f"Variables extraites par IA: {', '.join(variables)}")
                    return variables
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction des variables par IA: {str(e)}")
        
        # Fallback: extraction par regex des patterns courants
        return self._extract_variables_by_regex(content)

    def _extract_variables_by_regex(self, content: str) -> List[str]:
        """
        Extraction de variables par regex comme méthode de secours
        
        Args:
            content: Contenu du document
        
        Returns:
            Liste des noms de variables
        """
        logger.info("Utilisation de regex pour extraire les variables de template")
        
        # Liste pour stocker les variables trouvées
        variables = set()
        
        # Chercher les patterns de variables courants dans les templates
        patterns = [
            r'\{\{([a-zA-Z_éèêëàâäôöûüçÉÈÊËÀÂÄÔÖÛÜÇ]+(?:_[a-zA-Z0-9_éèêëàâäôöûüçÉÈÊËÀÂÄÔÖÛÜÇ]+)*)\}\}',  # {{nom_variable}}
            r'\[([a-zA-Z_éèêëàâäôöûüçÉÈÊËÀÂÄÔÖÛÜÇ]+(?:_[a-zA-Z0-9_éèêëàâäôöûüçÉÈÊËÀÂÄÔÖÛÜÇ]+)*)\]',      # [nom_variable]
            r'\$\{([a-zA-Z_éèêëàâäôöûüçÉÈÊËÀÂÄÔÖÛÜÇ]+(?:_[a-zA-Z0-9_éèêëàâäôöûüçÉÈÊËÀÂÄÔÖÛÜÇ]+)*)\}',    # ${nom_variable}
            r'<([a-zA-Z_éèêëàâäôöûüçÉÈÊËÀÂÄÔÖÛÜÇ]+(?:_[a-zA-Z0-9_éèêëàâäôöûüçÉÈÊËÀÂÄÔÖÛÜÇ]+)*)>',        # <nom_variable>
            r'%([a-zA-Z_éèêëàâäôöûüçÉÈÊËÀÂÄÔÖÛÜÇ]+(?:_[a-zA-Z0-9_éèêëàâäôöûüçÉÈÊËÀÂÄÔÖÛÜÇ]+)*)%'         # %nom_variable%
        ]
        
        # Trouver les correspondances pour chaque pattern
        for pattern in patterns:
            matches = re.findall(pattern, content)
            for match in matches:
                variables.add(match.lower())
        
        # Convertir le set en liste
        var_list = list(variables)
        
        if var_list:
            logger.info(f"Variables extraites par regex: {', '.join(var_list)}")
        
        return var_list

    def replace_variables(self, template: str, values: Dict[str, str]) -> str:
        """
        Remplace les variables dans un template par des valeurs
        Méthode simplifiée pour éviter les problèmes de remplacement
        
        Args:
            template: Template du document
            values: Dictionnaire des valeurs pour chaque variable
            
        Returns:
            Document personnalisé
        """
        logger.info(f"Remplacement direct des variables ({len(values)} variables)")
        
        # Vérification du template
        if not template:
            logger.error("Template vide, impossible de remplacer les variables")
            return ""
            
        # Utiliser le template original
        result = template
        
        # Nombre de remplacements effectués
        replacements_count = 0
        
        # Pour chaque variable, faire un remplacement direct
        for var_name, value in values.items():
            # S'assurer que la valeur est une chaîne
            value_str = str(value) if value is not None else ""
            
            # Format 1: {variable}
            original_count = result.count(f"{{{var_name}}}")
            result = result.replace(f"{{{var_name}}}", value_str)
            count1 = original_count - result.count(f"{{{var_name}}}")
            
            # Format 2: {{variable}}
            original_count = result.count(f"{{{{{var_name}}}}}")
            result = result.replace(f"{{{{{var_name}}}}}", value_str)
            count2 = original_count - result.count(f"{{{{{var_name}}}}}")
            
            # Compter le nombre total de remplacements pour cette variable
            total_count = count1 + count2
            if total_count > 0:
                replacements_count += total_count
                logger.debug(f"Variable '{var_name}' remplacée {total_count} fois")
        
        # Log du résultat
        logger.info(f"Remplacement terminé: {replacements_count} remplacements effectués")
        
        # Vérification du résultat
        if not result or len(result.strip()) == 0:
            logger.error("ERREUR: Document final vide après remplacement des variables!")
        elif len(result) < 10:
            logger.warning(f"Document final très court après remplacement: '{result}'")
        
        return result

    def process_template(self, template_content: str, values: Dict[str, str] = None) -> Dict[str, Any]:
        """
        Traite un template pour extraire ses variables et optionnellement les remplacer
        
        Args:
            template_content: Contenu du template
            values: Valeurs de remplacement (optionnel)
            
        Returns:
            Dictionnaire contenant les variables détectées et le document personnalisé
        """
        logger.info("Traitement d'un template de document")
        
        # 1. Extraire les variables du template
        variables = self.extract_template_variables(template_content)
        
        result = {
            "variables": variables,
            "template": template_content,
            "personalized": None
        }
        
        # 2. Si des valeurs sont fournies, remplacer les variables
        if values and variables:
            # Filtrer les valeurs pour ne conserver que celles qui correspondent aux variables
            valid_values = {}
            for var in variables:
                # Chercher la variable dans les valeurs (insensible à la casse)
                for key, value in values.items():
                    if key.lower() == var.lower():
                        valid_values[var] = value
                        break
            
            # Remplacer les variables par les valeurs
            if valid_values:
                personalized_doc = self.replace_variables(template_content, valid_values)
                result["personalized"] = personalized_doc
                result["replaced_variables"] = list(valid_values.keys())
        
        return result

    def test_document_generation(self, template_content: str = None, variables: Dict[str, str] = None) -> str:
        """
        Fonction de test pour vérifier directement la génération de document
        
        Args:
            template_content: Contenu du template (optionnel, un template simple sera utilisé sinon)
            variables: Variables à remplacer (optionnel, des variables de test seront utilisées sinon)
            
        Returns:
            Document généré
        """
        # Template par défaut si aucun n'est fourni
        if not template_content:
            template_content = """
            DOCUMENT DE TEST
            -----------------
            
            Nom: {nom}
            Prénom: {prenom}
            Email: {email}
            Date: {date}
            
            Ce document est adressé à {nom} {prenom}.
            Vous pouvez nous contacter à l'adresse {email}.
            
            Document généré le {date}.
            """
            
        # Variables par défaut si aucune n'est fournie
        if not variables:
            variables = {
                "nom": "Dupont",
                "prenom": "Jean",
                "email": "jean.dupont@example.com",
                "date": "22/03/2025"
            }
            
        # Afficher les informations de test
        print("\n===== TEST DE GÉNÉRATION DE DOCUMENT =====")
        print(f"Template utilisé ({len(template_content)} caractères):")
        print("-" * 40)
        print(template_content)
        print("-" * 40)
        print(f"Variables à remplacer: {variables}")
        
        # Essayer les deux méthodes de remplacement
        try:
            # Méthode 1: replace_variables
            result1 = self.replace_variables(template_content, variables)
            print("\nRésultat avec replace_variables:")
            print("-" * 40)
            print(result1)
            print("-" * 40)
            print(f"Longueur du document généré: {len(result1)} caractères")
            
            # Méthode 2: personalize_document
            result2 = self.personalize_document(template_content, variables)
            print("\nRésultat avec personalize_document:")
            print("-" * 40)
            print(result2)
            print("-" * 40)
            print(f"Longueur du document généré: {len(result2)} caractères")
            
            # Vérifier si les deux méthodes donnent le même résultat
            if result1 == result2:
                print("\n✅ Les deux méthodes donnent le même résultat!")
            else:
                print("\n❌ Les deux méthodes donnent des résultats différents!")
                
            return result1
            
        except Exception as e:
            print(f"\n❌ ERREUR lors du test: {str(e)}")
            traceback.print_exc()
            return f"ERREUR: {str(e)}"

# Test direct si le script est exécuté directement
if __name__ == "__main__":
    import logging
    
    # Configuration du logging pour le test
    logging.basicConfig(
        level=logging.DEBUG,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    print("===== TEST DU GÉNÉRATEUR DE DOCUMENTS =====")
    
    # Créer une instance du processeur de documents
    processor = AIDocumentProcessor()
    
    # Cas de test 1 : Template simple
    print("\n----- TEST 1: TEMPLATE SIMPLE -----")
    template_simple = """
    DOCUMENT TEST SIMPLE
    --------------------
    
    Nom: {nom}
    Email: {email}
    Date: {date}
    
    Bonjour {nom},
    
    Merci de nous avoir contacté à l'adresse {email}.
    
    Date: {date}
    """
    
    variables_simple = {
        "nom": "Dupont",
        "email": "dupont@example.com",
        "date": "22/03/2025"
    }
    
    result_simple = processor.test_document_generation(template_simple, variables_simple)
    
    # Cas de test 2 : Formats de variables variés
    print("\n----- TEST 2: FORMATS DE VARIABLES VARIÉS -----")
    template_formats = """
    DOCUMENT MULTI-FORMATS
    ---------------------
    
    Format 1: {nom}
    Format 2: {{nom}}
    Format 3: [nom]
    Format 4: <nom>
    Format 5: $nom$
    Format 6: «nom»
    
    Email format 1: {email}
    Email format 2: [email]
    """
    
    variables_formats = {
        "nom": "Dupont",
        "email": "dupont@example.com"
    }
    
    result_formats = processor.test_document_generation(template_formats, variables_formats)
    
    # Cas de test 3 : Template avec variables non définies
    print("\n----- TEST 3: VARIABLES NON DÉFINIES -----")
    template_undefined = """
    DOCUMENT AVEC VARIABLES MANQUANTES
    --------------------------------
    
    Nom: {nom}
    Email: {email}
    Téléphone: {telephone}
    Adresse: {adresse}
    """
    
    variables_incomplete = {
        "nom": "Dupont",
        "email": "dupont@example.com"
        # telephone et adresse manquants
    }
    
    result_undefined = processor.test_document_generation(template_undefined, variables_incomplete)
    
    print("\n===== TESTS TERMINÉS =====")
