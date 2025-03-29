import sys
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import logging
import traceback
import tempfile

# Configuration du logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    stream=sys.stdout
)
logger = logging.getLogger(__name__)

def ensure_temp_directory():
    """S'assure que le dossier temp existe et est accessible en écriture"""
    temp_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'temp')
    try:
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)
            logger.info(f"Dossier temp créé: {temp_dir}")
        # Vérifier les permissions d'écriture
        test_file = os.path.join(temp_dir, 'test.txt')
        with open(test_file, 'w') as f:
            f.write('test')
        os.remove(test_file)
        logger.info("Permissions du dossier temp vérifiées avec succès")
        return temp_dir
    except Exception as e:
        logger.error(f"Erreur lors de la vérification du dossier temp: {e}")
        raise

def convert_html_to_docx(html_path, output_path):
    """
    Convertit un fichier HTML en DOCX
    
    Args:
        html_path: Chemin du fichier HTML
        output_path: Chemin du fichier DOCX de sortie
    """
    try:
        logger.info(f"Début de la conversion de {html_path} vers {output_path}")

        # Vérifier que le fichier HTML existe
        if not os.path.exists(html_path):
            logger.error(f"Le fichier HTML n'existe pas: {html_path}")
            return False

        # Vérifier que le dossier de sortie existe
        output_dir = os.path.dirname(output_path)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            logger.info(f"Dossier de sortie créé: {output_dir}")

        # Lire le fichier HTML
        logger.info("Lecture du fichier HTML...")
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        logger.info(f"Taille du contenu HTML: {len(html_content)} caractères")

        # Vérifier que le contenu HTML est valide
        if not html_content.strip():
            logger.error("Le contenu HTML est vide")
            return False

        # Parser le HTML
        logger.info("Parsing du HTML...")
        soup = BeautifulSoup(html_content, 'html.parser')

        # Créer un nouveau document Word
        logger.info("Création du document Word...")
        doc = Document()

        # Définir les marges
        logger.info("Configuration des marges...")
        for section in doc.sections:
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            section.top_margin = Inches(1.2)
            section.bottom_margin = Inches(1.2)

        # Définir la police par défaut
        logger.info("Configuration de la police par défaut...")
        doc_style = doc.styles['Normal']
        doc_style.font.name = 'Arial'
        doc_style.font.size = Pt(11)

        # Parcourir les éléments HTML
        logger.info("Traitement des éléments HTML...")
        element_count = 0
        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'ul', 'ol', 'li', 'strong', 'em', 'u', 's', 'blockquote', 'img']):
            element_count += 1
            try:
                if element.name == 'p':
                    p = doc.add_paragraph()
                    p.add_run(element.get_text())
                elif element.name in ['h1', 'h2', 'h3']:
                    level = int(element.name[1])
                    p = doc.add_heading(element.get_text(), level=level)
                elif element.name == 'ul':
                    for li in element.find_all('li'):
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(li.get_text())
                elif element.name == 'ol':
                    for li in element.find_all('li'):
                        p = doc.add_paragraph(style='List Number')
                        p.add_run(li.get_text())
                elif element.name == 'blockquote':
                    p = doc.add_paragraph(style='Quote')
                    p.add_run(element.get_text())
                elif element.name == 'img':
                    try:
                        img_src = element.get('src')
                        if img_src:
                            logger.info(f"Traitement de l'image: {img_src[:50]}...")
                            # Gérer les images en base64
                            if img_src.startswith('data:image'):
                                import base64
                                # Extraire les données de l'image
                                header, encoded = img_src.split(",", 1)
                                data = base64.b64decode(encoded)
                                # Créer un fichier temporaire
                                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                                    tmp.write(data)
                                    img_src = tmp.name
                                logger.info("Image base64 convertie en fichier temporaire")
                            
                            # Ajouter l'image au document
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.add_run().add_picture(img_src, width=Inches(6.0))
                            
                            # Nettoyer le fichier temporaire si c'était une image base64
                            if img_src.startswith(tempfile.gettempdir()):
                                os.unlink(img_src)
                                logger.info("Fichier temporaire d'image supprimé")
                    except Exception as img_error:
                        logger.error(f"Erreur lors de l'ajout de l'image: {img_error}")
                        logger.error(traceback.format_exc())
            except Exception as element_error:
                logger.error(f"Erreur lors du traitement de l'élément {element.name}: {element_error}")
                logger.error(traceback.format_exc())

        logger.info(f"Nombre total d'éléments traités: {element_count}")

        # Sauvegarder le document
        logger.info("Sauvegarde du document...")
        doc.save(output_path)
        logger.info(f"Document DOCX généré avec succès: {output_path}")
        return True

    except Exception as e:
        logger.error(f"Erreur lors de la conversion HTML vers DOCX: {e}")
        logger.error(traceback.format_exc())
        return False

if __name__ == "__main__":
    try:
        # Vérifier le dossier temp
        ensure_temp_directory()

        if len(sys.argv) != 3:
            print("Usage: python convert_to_docx.py <input_html> <output_docx>")
            sys.exit(1)

        html_path = sys.argv[1]
        output_path = sys.argv[2]

        if not os.path.exists(html_path):
            print(f"Le fichier HTML n'existe pas: {html_path}")
            sys.exit(1)

        success = convert_html_to_docx(html_path, output_path)
        sys.exit(0 if success else 1)
    except Exception as e:
        logger.error(f"Erreur fatale: {e}")
        logger.error(traceback.format_exc())
        sys.exit(1) 