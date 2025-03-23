"""
Module contenant le modèle principal de l'application Vynal Docs Automator.
Ce modèle gère les données des clients, des modèles de documents et des documents générés.
"""

import os
import json
import shutil
import uuid
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple
from functools import lru_cache
from utils.cache_manager import CacheManager
from utils.lazy_loader import LazyModelLoader, ModelType, ModelMetadata

logger = logging.getLogger("VynalDocsAutomator.DocumentModel")

class AppModel:
    """
    Modèle principal de l'application qui gère toutes les données.
    
    Attributs:
        data_dir (str): Chemin vers le répertoire de données
        clients (Dict): Dictionnaire des clients
        templates (Dict): Dictionnaire des modèles de documents
        documents (Dict): Dictionnaire des documents générés
        recent_activities (List): Liste des activités récentes
        config (Dict): Configuration de l'application
        template_folders (Dict): Dictionnaire des dossiers de modèles
    """
    
    def __init__(self, data_dir: str = "./data"):
        """
        Initialise le modèle principal de l'application.
        
        Args:
            data_dir (str): Chemin vers le répertoire de données
        """
        self.data_dir = data_dir
        self._create_directories()
        
        # Initialisation des données avec chargement paresseux
        self._clients = None
        self._templates = None
        self._documents = None
        self._recent_activities = None
        self._config = None
        
        # Initialisation du cache
        self.cache_manager = CacheManager()
        
        # Cache pour les documents fréquemment consultés
        self._document_cache = {}
        self._document_cache_size = 100
        self._document_access_patterns = {}
        
        # Initialisation des dossiers de modèles
        self.template_folders = {
            "juridique": "Juridique",
            "contrats": "Contrats",
            "commercial": "Commercial",
            "administratif": "Administratif",
            "ressources_humaines": "Ressources Humaines",
            "fiscal": "Fiscal",
            "correspondance": "Correspondance",
            "bancaire": "Bancaire",
            "corporate": "Corporate",
            "immobilier": "Immobilier"
        }
        
        # Initialisation avec chargement différé
        self.lazy_loader = LazyModelLoader(
            cache_ttl=self.config.get('cache_ttl', 3600),
            max_cache_size=self.config.get('max_cache_size', 1000)
        )
        
        # Enregistrement des callbacks
        self.lazy_loader.register_callbacks(
            ModelType.DOCUMENT,
            self._load_document_by_id,
            self._load_document_metadata
        )
    
    @property
    def clients(self):
        if self._clients is None:
            self._clients = self._load_clients()
        return self._clients
    
    @property
    def templates(self):
        if self._templates is None:
            self._templates = self._load_templates()
        return self._templates
    
    @property
    def documents(self):
        if self._documents is None:
            self._documents = self._load_documents()
        return self._documents
    
    @property
    def recent_activities(self):
        if self._recent_activities is None:
            self._recent_activities = self._load_activities()
        return self._recent_activities
    
    @property
    def config(self):
        if self._config is None:
            self._config = self._load_config()
        return self._config
    
    @lru_cache(maxsize=100)
    def get_document(self, document_id: str) -> Optional[Dict]:
        """Récupère un document avec mise en cache LRU."""
        return self.documents.get(document_id)
    
    @lru_cache(maxsize=50)
    def get_template(self, template_id: str) -> Optional[Dict]:
        """Récupère un template avec mise en cache LRU."""
        return self.templates.get(template_id)
    
    @lru_cache(maxsize=100)
    def get_client(self, client_id: str) -> Optional[Dict]:
        """Récupère un client avec mise en cache LRU."""
        return self.clients.get(client_id)
    
    def _load_clients(self) -> Dict:
        """Charge les clients depuis le fichier JSON avec mise en cache."""
        clients_path = os.path.join(self.data_dir, "clients.json")
        if os.path.exists(clients_path):
            with open(clients_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}
    
    def _load_templates(self) -> Dict:
        """Charge les templates depuis le fichier JSON avec mise en cache."""
        templates_path = os.path.join(self.data_dir, "templates.json")
        if os.path.exists(templates_path):
            with open(templates_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}
    
    def _load_documents(self) -> Dict:
        """Charge les documents depuis le fichier JSON avec mise en cache."""
        documents_path = os.path.join(self.data_dir, "documents.json")
        if os.path.exists(documents_path):
            with open(documents_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}
    
    def _load_activities(self) -> List:
        """Charge les activités depuis le fichier JSON avec mise en cache."""
        activities_path = os.path.join(self.data_dir, "activities.json")
        if os.path.exists(activities_path):
            with open(activities_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return []
    
    def _load_config(self) -> Dict:
        """Charge la configuration depuis le fichier JSON avec mise en cache."""
        config_path = os.path.join(self.data_dir, "config.json")
        if os.path.exists(config_path):
            with open(config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return self._get_default_config()
    
    def _get_default_config(self) -> Dict:
        """Retourne la configuration par défaut."""
        return {
            "app_name": "Vynal Docs Automator",
            "theme": "system",
            "language": "fr",
            "backup_frequency": "daily",
            "max_recent_activities": 50,
            "document_output_dir": os.path.join(self.data_dir, "documents"),
            "default_format": "pdf"
        }
    
    def save_data(self) -> None:
        """Sauvegarde toutes les données avec gestion des erreurs."""
        try:
            if self._clients is not None:
                self._save_clients()
            if self._templates is not None:
                self._save_templates()
            if self._documents is not None:
                self._save_documents()
            if self._recent_activities is not None:
                self._save_activities()
            if self._config is not None:
                self._save_config()
        except Exception as e:
            print(f"Erreur lors de la sauvegarde des données: {e}")
    
    def _save_clients(self) -> None:
        """Sauvegarde les clients dans le fichier JSON."""
        clients_path = os.path.join(self.data_dir, "clients.json")
        with open(clients_path, 'w', encoding='utf-8') as f:
            json.dump(self._clients, f, indent=4)
    
    def _save_templates(self) -> None:
        """Sauvegarde les templates dans le fichier JSON."""
        templates_path = os.path.join(self.data_dir, "templates.json")
        with open(templates_path, 'w', encoding='utf-8') as f:
            json.dump(self._templates, f, indent=4)
    
    def _save_documents(self) -> None:
        """Sauvegarde les documents dans le fichier JSON."""
        documents_path = os.path.join(self.data_dir, "documents.json")
        with open(documents_path, 'w', encoding='utf-8') as f:
            json.dump(self._documents, f, indent=4)
    
    def _save_activities(self) -> None:
        """Sauvegarde les activités dans le fichier JSON."""
        activities_path = os.path.join(self.data_dir, "activities.json")
        with open(activities_path, 'w', encoding='utf-8') as f:
            json.dump(self._recent_activities, f, indent=4)
    
    def _save_config(self) -> None:
        """Sauvegarde la configuration dans le fichier JSON."""
        config_path = os.path.join(self.data_dir, "config.json")
        with open(config_path, 'w', encoding='utf-8') as f:
            json.dump(self._config, f, indent=4)
    
    def _create_directories(self) -> None:
        """Crée les répertoires nécessaires s'ils n'existent pas."""
        directories = [
            self.data_dir,
            os.path.join(self.data_dir, "clients"),
            os.path.join(self.data_dir, "templates"),
            os.path.join(self.data_dir, "documents"),
            os.path.join(self.data_dir, "backup")
        ]
        
        for directory in directories:
            if not os.path.exists(directory):
                os.makedirs(directory)
    
    def save_config(self) -> None:
        """Sauvegarde la configuration de l'application."""
        config_path = os.path.join(self.data_dir, "config.json")
        with open(config_path, 'w', encoding='utf-8') as f:
            json.dump(self.config, f, indent=4)
    
    def load_data(self) -> None:
        """Charge toutes les données de l'application."""
        self.load_clients()
        self.load_templates()
        self.load_documents()
        self.load_activities()
    
    def save_data(self) -> None:
        """Sauvegarde toutes les données de l'application."""
        self.save_clients()
        self.save_templates()
        self.save_documents()
        self.save_activities()
        self.save_config()
    
    # === Gestion des clients ===
    
    def load_clients(self) -> None:
        """Charge les données des clients depuis le fichier JSON."""
        clients_path = os.path.join(self.data_dir, "clients.json")
        if os.path.exists(clients_path):
            with open(clients_path, 'r', encoding='utf-8') as f:
                self.clients = json.load(f)
    
    def save_clients(self) -> None:
        """Sauvegarde les données des clients dans un fichier JSON."""
        clients_path = os.path.join(self.data_dir, "clients.json")
        with open(clients_path, 'w', encoding='utf-8') as f:
            json.dump(self.clients, f, indent=4)
    
    def add_client(self, client_data: Dict) -> str:
        """
        Ajoute un nouveau client.
        
        Args:
            client_data (Dict): Données du client à ajouter
            
        Returns:
            str: ID du client ajouté
        """
        client_id = str(uuid.uuid4())
        client_data["id"] = client_id
        client_data["created_at"] = datetime.now().isoformat()
        client_data["updated_at"] = datetime.now().isoformat()
        
        self.clients[client_id] = client_data
        self.save_clients()
        
        # Enregistrer l'activité
        self.add_activity(
            action_type="client_added",
            details=f"Client ajouté: {client_data.get('name', 'Sans nom')}",
            related_id=client_id
        )
        
        return client_id
    
    def update_client(self, client_id: str, client_data: Dict) -> bool:
        """
        Met à jour un client existant.
        
        Args:
            client_id (str): ID du client à mettre à jour
            client_data (Dict): Nouvelles données du client
            
        Returns:
            bool: True si la mise à jour est réussie, False sinon
        """
        if client_id not in self.clients:
            return False
        
        # Conserver certaines données existantes
        client_data["id"] = client_id
        client_data["created_at"] = self.clients[client_id]["created_at"]
        client_data["updated_at"] = datetime.now().isoformat()
        
        self.clients[client_id] = client_data
        self.save_clients()
        
        # Enregistrer l'activité
        self.add_activity(
            action_type="client_updated",
            details=f"Client mis à jour: {client_data.get('name', 'Sans nom')}",
            related_id=client_id
        )
        
        return True
    
    def delete_client(self, client_id: str) -> bool:
        """
        Supprime un client.
        
        Args:
            client_id (str): ID du client à supprimer
            
        Returns:
            bool: True si la suppression est réussie, False sinon
        """
        if client_id not in self.clients:
            return False
        
        client_name = self.clients[client_id].get("name", "Sans nom")
        del self.clients[client_id]
        self.save_clients()
        
        # Enregistrer l'activité
        self.add_activity(
            action_type="client_deleted",
            details=f"Client supprimé: {client_name}",
            related_id=client_id
        )
        
        return True
    
    def get_all_clients(self) -> Dict:
        """
        Récupère tous les clients.
        
        Returns:
            Dict: Dictionnaire des clients
        """
        return self.clients
    
    def import_clients_from_csv(self, csv_path: str) -> Tuple[int, List[str]]:
        """
        Importe des clients depuis un fichier CSV.
        
        Args:
            csv_path (str): Chemin vers le fichier CSV
            
        Returns:
            Tuple[int, List[str]]: Nombre de clients importés et liste des erreurs
        """
        import csv
        
        imported_count = 0
        errors = []
        
        try:
            with open(csv_path, 'r', encoding='utf-8') as csvfile:
                reader = csv.DictReader(csvfile)
                
                for row in reader:
                    try:
                        # Vérification des champs obligatoires
                        if not row.get('name'):
                            errors.append(f"Ligne {reader.line_num}: Le nom est obligatoire")
                            continue
                        
                        # Formater les données du client
                        client_data = {
                            "name": row.get('name', ''),
                            "company": row.get('company', ''),
                            "email": row.get('email', ''),
                            "phone": row.get('phone', ''),
                            "address": row.get('address', ''),
                            "notes": row.get('notes', '')
                        }
                        
                        # Ajouter des champs supplémentaires s'ils existent
                        for key, value in row.items():
                            if key not in client_data and value:
                                client_data[key] = value
                        
                        # Ajouter le client
                        self.add_client(client_data)
                        imported_count += 1
                        
                    except Exception as e:
                        errors.append(f"Ligne {reader.line_num}: {str(e)}")
        
        except Exception as e:
            errors.append(f"Erreur lors de l'ouverture du fichier: {str(e)}")
        
        return imported_count, errors
    
    def export_clients_to_csv(self, csv_path: str) -> bool:
        """
        Exporte les clients vers un fichier CSV.
        
        Args:
            csv_path (str): Chemin vers le fichier CSV de destination
            
        Returns:
            bool: True si l'exportation est réussie, False sinon
        """
        import csv
        
        try:
            # Déterminer tous les champs possibles
            all_fields = set()
            for client in self.clients.values():
                all_fields.update(client.keys())
            
            # Exclure certains champs internes
            fields_to_exclude = {'id', 'created_at', 'updated_at'}
            fields = [f for f in all_fields if f not in fields_to_exclude]
            
            # Trier les champs pour une meilleure lisibilité
            fields.sort()
            
            # Placer les champs communs en premier
            common_fields = ['name', 'company', 'email', 'phone', 'address', 'notes']
            for field in reversed(common_fields):
                if field in fields:
                    fields.remove(field)
                    fields.insert(0, field)
            
            with open(csv_path, 'w', encoding='utf-8', newline='') as csvfile:
                writer = csv.DictWriter(csvfile, fieldnames=fields)
                writer.writeheader()
                
                for client in self.clients.values():
                    # Ne pas exporter les champs exclus
                    row = {k: v for k, v in client.items() if k in fields}
                    writer.writerow(row)
            
            return True
        
        except Exception as e:
            print(f"Erreur lors de l'exportation: {str(e)}")
            return False
    
    # === Gestion des modèles de documents ===
    
    def load_templates(self) -> None:
        """Charge les données des modèles depuis le fichier JSON."""
        templates_path = os.path.join(self.data_dir, "templates.json")
        if os.path.exists(templates_path):
            with open(templates_path, 'r', encoding='utf-8') as f:
                self.templates = json.load(f)
    
    def save_templates(self) -> None:
        """Sauvegarde les données des modèles dans un fichier JSON."""
        templates_path = os.path.join(self.data_dir, "templates.json")
        with open(templates_path, 'w', encoding='utf-8') as f:
            json.dump(self.templates, f, indent=4)
    
    def add_template(self, template_data: Dict) -> str:
        """
        Ajoute un nouveau modèle de document.
        
        Args:
            template_data (Dict): Données du modèle à ajouter
            
        Returns:
            str: ID du modèle ajouté
        """
        template_id = str(uuid.uuid4())
        template_data["id"] = template_id
        template_data["created_at"] = datetime.now().isoformat()
        template_data["updated_at"] = datetime.now().isoformat()
        
        # Assurez-vous que les variables requises sont définies
        if "variables" not in template_data:
            template_data["variables"] = []
        
        # Ajouter le dossier par défaut si non spécifié
        if "folder" not in template_data:
            template_data["folder"] = "autre"
        
        # Si un fichier de contenu est fourni, le copier dans le répertoire des modèles
        if "content_file" in template_data and os.path.exists(template_data["content_file"]):
            file_name = f"{template_id}{os.path.splitext(template_data['content_file'])[1]}"
            destination = os.path.join(self.data_dir, "templates", file_name)
            shutil.copy2(template_data["content_file"], destination)
            template_data["content_file"] = file_name
        
        self.templates[template_id] = template_data
        self.save_templates()
        
        # Enregistrer l'activité
        self.add_activity(
            action_type="template_added",
            details=f"Modèle ajouté: {template_data.get('name', 'Sans nom')} dans {self.template_folders.get(template_data['folder'], 'Autre')}",
            related_id=template_id
        )
        
        return template_id
    
    def update_template(self, template_id: str, template_data: Dict) -> bool:
        """
        Met à jour un modèle existant.
        
        Args:
            template_id (str): ID du modèle à mettre à jour
            template_data (Dict): Nouvelles données du modèle
            
        Returns:
            bool: True si la mise à jour est réussie, False sinon
        """
        if template_id not in self.templates:
            return False
        
        # Conserver certaines données existantes
        template_data["id"] = template_id
        template_data["created_at"] = self.templates[template_id]["created_at"]
        template_data["updated_at"] = datetime.now().isoformat()
        
        # Si un nouveau fichier de contenu est fourni
        if "content_file" in template_data and os.path.exists(template_data["content_file"]):
            # Si c'est un chemin, pas juste un nom de fichier
            if os.path.dirname(template_data["content_file"]):
                file_name = f"{template_id}{os.path.splitext(template_data['content_file'])[1]}"
                destination = os.path.join(self.data_dir, "templates", file_name)
                shutil.copy2(template_data["content_file"], destination)
                template_data["content_file"] = file_name
        else:
            # Conserver l'ancien fichier de contenu
            template_data["content_file"] = self.templates[template_id].get("content_file", "")
        
        self.templates[template_id] = template_data
        self.save_templates()
        
        # Enregistrer l'activité
        self.add_activity(
            action_type="template_updated",
            details=f"Modèle mis à jour: {template_data.get('name', 'Sans nom')}",
            related_id=template_id
        )
        
        return True
    
    def delete_template(self, template_id: str) -> bool:
        """
        Supprime un modèle.
        
        Args:
            template_id (str): ID du modèle à supprimer
            
        Returns:
            bool: True si la suppression est réussie, False sinon
        """
        if template_id not in self.templates:
            return False
        
        template_name = self.templates[template_id].get("name", "Sans nom")
        
        # Supprimer le fichier de contenu associé
        content_file = self.templates[template_id].get("content_file", "")
        if content_file:
            file_path = os.path.join(self.data_dir, "templates", content_file)
            if os.path.exists(file_path):
                os.remove(file_path)
        
        del self.templates[template_id]
        self.save_templates()
        
        # Enregistrer l'activité
        self.add_activity(
            action_type="template_deleted",
            details=f"Modèle supprimé: {template_name}",
            related_id=template_id
        )
        
        return True
    
    def get_all_templates(self) -> Dict:
        """
        Récupère tous les modèles.
        
        Returns:
            Dict: Dictionnaire des modèles
        """
        return self.templates
    
    def get_template_content(self, template_id: str) -> Optional[str]:
        """
        Récupère le contenu d'un modèle.
        
        Args:
            template_id (str): ID du modèle
            
        Returns:
            Optional[str]: Contenu du modèle ou None si non trouvé
        """
        template = self.get_template(template_id)
        if not template:
            return None
        
        content_file = template.get("content_file", "")
        if not content_file:
            return template.get("content", "")
        
        file_path = os.path.join(self.data_dir, "templates", content_file)
        if not os.path.exists(file_path):
            return template.get("content", "")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def get_templates_by_folder(self, folder: str) -> List[Dict[str, Any]]:
        """
        Récupère tous les modèles d'un dossier spécifique
        
        Args:
            folder: Identifiant du dossier
            
        Returns:
            list: Liste des modèles dans le dossier
        """
        return [t.copy() for t in self.templates.values() if t.get('folder') == folder]
    
    def get_template_folders(self) -> Dict[str, str]:
        """
        Récupère la liste des dossiers de modèles
        
        Returns:
            dict: Dictionnaire des dossiers {id: nom}
        """
        return self.template_folders.copy()
    
    # === Gestion des documents générés ===
    
    def add_document(self, document_data: Dict, file_path: str) -> str:
        """
        Ajoute un nouveau document avec mise en cache.
        
        Args:
            document_data: Données du document
            file_path: Chemin du fichier
            
        Returns:
            str: ID du document ajouté
        """
        document_id = str(uuid.uuid4())
        document_data["id"] = document_id
        document_data["file_path"] = file_path
        document_data["created_at"] = datetime.now().isoformat()
        document_data["updated_at"] = datetime.now().isoformat()
        
        self.documents[document_id] = document_data
        self.save_documents()
        
        # Mettre en cache
        self.cache_manager.set("documents", document_id, document_data)
        
        # Invalider le cache des documents du client
        if "client_id" in document_data:
            self.cache_manager.invalidate("documents", f"client_docs_{document_data['client_id']}")
        
        # Enregistrer l'activité
        self.add_activity(
            action_type="document_added",
            details=f"Document ajouté: {document_data.get('name', 'Sans nom')}",
            related_id=document_id
        )
        
        return document_id
    
    def delete_document(self, document_id: str) -> bool:
        """
        Supprime un document avec invalidation du cache.
        
        Args:
            document_id: ID du document à supprimer
            
        Returns:
            bool: True si la suppression est réussie, False sinon
        """
        if document_id not in self.documents:
            return False
        
        # Récupérer les informations avant la suppression
        doc = self.documents[document_id]
        client_id = doc.get("client_id")
        
        # Supprimer le fichier physique
        try:
            if os.path.exists(doc["file_path"]):
                os.remove(doc["file_path"])
        except Exception as e:
            logger.error(f"Erreur lors de la suppression du fichier: {e}")
        
        # Supprimer des données
        del self.documents[document_id]
        self.save_documents()
        
        # Invalider le cache
        self.cache_manager.invalidate("documents", document_id)
        if client_id:
            self.cache_manager.invalidate("documents", f"client_docs_{client_id}")
        
        # Enregistrer l'activité
        self.add_activity(
            action_type="document_deleted",
            details=f"Document supprimé: {doc.get('name', 'Sans nom')}",
            related_id=document_id
        )
        
        return True
    
    def get_document_path(self, document_id: str) -> Optional[str]:
        """
        Récupère le chemin du fichier d'un document.
        
        Args:
            document_id (str): ID du document
            
        Returns:
            Optional[str]: Chemin du fichier ou None si non trouvé
        """
        document = self.get_document(document_id)
        if not document:
            return None
        
        file_name = document.get("file_name", "")
        if not file_name:
            return None
        
        return os.path.join(self.data_dir, "documents", file_name)
    
    def get_all_documents(self) -> Dict:
        """
        Récupère tous les documents.
        
        Returns:
            Dict: Dictionnaire des documents
        """
        return self.documents
    
    def get_client_documents(self, client_id: str) -> Dict:
        """
        Récupère tous les documents d'un client avec mise en cache.
        
        Args:
            client_id: ID du client
            
        Returns:
            Dict: Dictionnaire des documents du client
        """
        # Vérifier le cache
        cache_key = f"client_docs_{client_id}"
        cached_docs = self.cache_manager.get("documents", cache_key)
        if cached_docs is not None:
            return cached_docs
            
        # Si non trouvé dans le cache, filtrer les documents
        client_docs = {
            doc_id: doc for doc_id, doc in self.documents.items()
            if doc.get("client_id") == client_id
        }
        
        # Mettre en cache
        self.cache_manager.set("documents", cache_key, client_docs)
        return client_docs
    
    # === Gestion des activités récentes ===
    
    def load_activities(self) -> None:
        """Charge les activités récentes depuis le fichier JSON."""
        activities_path = os.path.join(self.data_dir, "activities.json")
        if os.path.exists(activities_path):
            with open(activities_path, 'r', encoding='utf-8') as f:
                self.recent_activities = json.load(f)
    
    def save_activities(self) -> None:
        """Sauvegarde les activités récentes dans un fichier JSON."""
        activities_path = os.path.join(self.data_dir, "activities.json")
        with open(activities_path, 'w', encoding='utf-8') as f:
            json.dump(self.recent_activities, f, indent=4)
    
    def add_activity(self, action_type: str, details: str, related_id: Optional[str] = None) -> None:
        """
        Ajoute une nouvelle activité.
        
        Args:
            action_type (str): Type d'action
            details (str): Détails de l'activité
            related_id (Optional[str]): ID de l'élément concerné
        """
        activity = {
            "id": str(uuid.uuid4()),
            "timestamp": datetime.now().isoformat(),
            "action_type": action_type,
            "details": details,
            "related_id": related_id
        }
        
        self.recent_activities.insert(0, activity)
        
        # Limiter le nombre d'activités récentes
        max_activities = self.config.get("max_recent_activities", 50)
        if len(self.recent_activities) > max_activities:
            self.recent_activities = self.recent_activities[:max_activities]
        
        self.save_activities()
    
    def get_recent_activities(self, limit: int = None) -> List[Dict]:
        """
        Récupère les activités récentes.
        
        Args:
            limit (int, optional): Nombre maximum d'activités à récupérer
            
        Returns:
            List[Dict]: Liste des activités récentes
        """
        if limit is None:
            limit = self.config.get("max_recent_activities", 50)
        
        return self.recent_activities[:limit]
    
    # === Sauvegarde et restauration ===
    
    def create_backup(self) -> str:
        """
        Crée une sauvegarde complète des données.
        
        Returns:
            str: Nom du fichier de sauvegarde
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_dir = os.path.join(self.data_dir, "backup")
        backup_file = os.path.join(backup_dir, f"backup_{timestamp}.zip")
        
        # Sauvegarder d'abord toutes les données
        self.save_data()
        
        # Créer l'archive ZIP
        import zipfile
        
        with zipfile.ZipFile(backup_file, 'w', zipfile.ZIP_DEFLATED) as zipf:
            # Ajouter les fichiers JSON principaux
            for file_name in ["clients.json", "templates.json", "documents.json", "activities.json", "config.json"]:
                file_path = os.path.join(self.data_dir, file_name)
                if os.path.exists(file_path):
                    zipf.write(file_path, file_name)
            
            # Ajouter les répertoires de fichiers
            for dir_name in ["templates", "documents"]:
                dir_path = os.path.join(self.data_dir, dir_name)
                if os.path.exists(dir_path):
                    for root, _, files in os.walk(dir_path):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arcname = os.path.join(dir_name, file)
                            zipf.write(file_path, arcname)
        
        # Enregistrer l'activité
        self.add_activity(
            action_type="backup_created",
            details=f"Sauvegarde créée: backup_{timestamp}.zip"
        )
        
        return backup_file
    
    def restore_backup(self, backup_file: str) -> bool:
        """
        Restaure une sauvegarde.
        
        Args:
            backup_file (str): Chemin vers le fichier de sauvegarde
            
        Returns:
            bool: True si la restauration est réussie, False sinon
        """
        if not os.path.exists(backup_file):
            return False
        
        import zipfile
        import tempfile
        
        try:
            # Créer un répertoire temporaire pour l'extraction
            temp_dir = tempfile.mkdtemp()
            
            # Extraire l'archive
            with zipfile.ZipFile(backup_file, 'r') as zipf:
                zipf.extractall(temp_dir)
            
            # Restaurer les fichiers JSON principaux
            for file_name in ["clients.json", "templates.json", "documents.json", "activities.json", "config.json"]:
                temp_file = os.path.join(temp_dir, file_name)
                if os.path.exists(temp_file):
                    shutil.copy2(temp_file, os.path.join(self.data_dir, file_name))
            
            # Restaurer les répertoires de fichiers
            for dir_name in ["templates", "documents"]:
                temp_subdir = os.path.join(temp_dir, dir_name)
                target_subdir = os.path.join(self.data_dir, dir_name)
                
                # Vider le répertoire cible
                if os.path.exists(target_subdir):
                    for item in os.listdir(target_subdir):
                        item_path = os.path.join(target_subdir, item)
                        if os.path.isfile(item_path):
                            os.unlink(item_path)
                
                # Copier les fichiers restaurés
                if os.path.exists(temp_subdir):
                    for item in os.listdir(temp_subdir):
                        src_path = os.path.join(temp_subdir, item)
                        if os.path.isfile(src_path):
                            shutil.copy2(src_path, os.path.join(target_subdir, item))
            
            # Recharger les données
            self.load_data()
            
            # Enregistrer l'activité
            self.add_activity(
                action_type="backup_restored",
                details=f"Sauvegarde restaurée: {os.path.basename(backup_file)}"
            )
            
            return True
            
        except Exception as e:
            print(f"Erreur lors de la restauration: {str(e)}")
            return False
        
        finally:
            # Nettoyer le répertoire temporaire
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
    
    def get_backup_list(self) -> List[Dict]:
        """
        Récupère la liste des sauvegardes disponibles.
        
        Returns:
            List[Dict]: Liste des informations sur les sauvegardes
        """
        backup_dir = os.path.join(self.data_dir, "backup")
        backups = []
        
        if os.path.exists(backup_dir):
            for file in os.listdir(backup_dir):
                if file.startswith("backup_") and file.endswith(".zip"):
                    file_path = os.path.join(backup_dir, file)
                    file_stats = os.stat(file_path)
                    
                    # Extraire la date de la sauvegarde du nom de fichier
                    try:
                        date_str = file[7:-4]  # Retirer "backup_" et ".zip"
                        date_obj = datetime.strptime(date_str, "%Y%m%d_%H%M%S")
                        date_formatted = date_obj.strftime("%d/%m/%Y %H:%M:%S")
                    except:
                        date_formatted = "Inconnue"
                    
                    backups.append({
                        "name": file,
                        "path": file_path,
                        "size": file_stats.st_size,
                        "date": date_formatted,
                        "timestamp": file_stats.st_mtime
                    })
        
        # Trier par date, plus récent en premier
        backups.sort(key=lambda x: x["timestamp"], reverse=True)
        return backups
    
    # === Génération de documents ===
    
    def fill_template_variables(self, template_content: str, variables: Dict) -> str:
        """
        Remplace les variables dans un modèle par leurs valeurs.
        
        Args:
            template_content (str): Contenu du modèle avec variables
            variables (Dict): Dictionnaire des variables et leurs valeurs
            
        Returns:
            str: Contenu avec variables remplacées
        """
        filled_content = template_content
        
        # Remplacer les variables de la forme {{variable_name}}
        for var_name, var_value in variables.items():
            placeholder = f"{{{{{var_name}}}}}"
            filled_content = filled_content.replace(placeholder, str(var_value))
        
        return filled_content
    
    def generate_document(self, template_id: str, client_id: str, custom_variables: Dict, 
                         output_format: str = "pdf", output_path: Optional[str] = None) -> Optional[str]:
        """
        Génère un document à partir d'un modèle et des informations d'un client.
        
        Args:
            template_id (str): ID du modèle
            client_id (str): ID du client
            custom_variables (Dict): Variables personnalisées
            output_format (str): Format de sortie ("pdf" ou "docx")
            output_path (Optional[str]): Chemin de sortie personnalisé
            
        Returns:
            Optional[str]: Chemin vers le document généré ou None en cas d'erreur
        """
        # Vérifier que le modèle et le client existent
        template = self.get_template(template_id)
        client = self.get_client(client_id)
        
        if not template or not client:
            return None
        
        # Récupérer le contenu du modèle
        template_content = self.get_template_content(template_id)
        if not template_content:
            return None
        
        # Préparer les variables standard du client
        client_variables = {
            "client_name": client.get("name", ""),
            "client_company": client.get("company", ""),
            "client_email": client.get("email", ""),
            "client_phone": client.get("phone", ""),
            "client_address": client.get("address", ""),
            "current_date": datetime.now().strftime("%d/%m/%Y")
        }
        
        # Ajouter d'autres champs du client comme variables
        for key, value in client.items():
            if key not in ["id", "created_at", "updated_at"] and key not in client_variables:
                client_variables[f"client_{key}"] = value
        
        # Fusionner avec les variables personnalisées
        all_variables = {**client_variables, **custom_variables}
        
        # Remplir le modèle avec les variables
        filled_content = self.fill_template_variables(template_content, all_variables)
        
        # Déterminer le chemin de sortie
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{template.get('name', 'document')}_{timestamp}"
            output_dir = self.config.get("document_output_dir", os.path.join(self.data_dir, "documents"))
            
            if output_format.lower() == "pdf":
                output_path = os.path.join(output_dir, f"{filename}.pdf")
            else:
                output_path = os.path.join(output_dir, f"{filename}.docx")
        
        # Générer le document dans le format demandé
        try:
            if output_format.lower() == "pdf":
                self._generate_pdf(filled_content, output_path, template)
            else:
                self._generate_docx(filled_content, output_path, template)
                
            # Enregistrer le document dans le modèle
            document_data = {
                "name": f"{template.get('name', 'Document')} - {client.get('name', 'Client')}",
                "template_id": template_id,
                "client_id": client_id,
                "format": output_format.lower(),
                "variables": all_variables
            }
            
            self.add_document(document_data, output_path)
            
            return output_path
            
        except Exception as e:
            print(f"Erreur lors de la génération du document: {str(e)}")
            return None
    
    def _generate_pdf(self, content: str, output_path: str, template_data: Dict) -> None:
        """
        Génère un document PDF à partir du contenu.
        
        Args:
            content (str): Contenu à mettre dans le PDF
            output_path (str): Chemin du fichier de sortie
            template_data (Dict): Données du modèle pour le formatage
        """
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.units import cm
        import html
        
        # Configuration du document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=3*cm,  # Marge haute augmentée
            bottomMargin=3*cm  # Marge basse augmentée
        )
        
        # Styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='CustomHeading1',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=12
        ))
        styles.add(ParagraphStyle(
            name='CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6
        ))
        
        # Convertir le contenu en éléments Platypus
        elements = []
        
        # Traitement basique du contenu
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                elements.append(Spacer(1, 0.5*cm))
            elif line.startswith('# '):
                elements.append(Paragraph(html.escape(line[2:]), styles['CustomHeading1']))
            elif line.startswith('## '):
                elements.append(Paragraph(html.escape(line[3:]), styles['Heading2']))
            else:
                elements.append(Paragraph(html.escape(line), styles['CustomNormal']))
        
        # Générer le PDF
        doc.build(elements)
    
    def _generate_docx(self, content: str, output_path: str, template_data: Dict) -> None:
        """
        Génère un document DOCX à partir du contenu.
        
        Args:
            content (str): Contenu à mettre dans le DOCX
            output_path (str): Chemin du fichier de sortie
            template_data (Dict): Données du modèle pour le formatage
        """
        from docx import Document
        from docx.shared import Inches
        
        # Créer un nouveau document
        doc = Document()
        
        # Définir la largeur maximale des marges
        for section in doc.sections:
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            section.top_margin = Inches(1.2)
            section.bottom_margin = Inches(1.2)
        
        # Traitement basique du contenu
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            else:
                p = doc.add_paragraph(line)
                # Limiter la largeur des paragraphes en ajustant les retraits
                p.paragraph_format.first_line_indent = Inches(0)
                p.paragraph_format.left_indent = Inches(0)
                p.paragraph_format.right_indent = Inches(0)
        
        # Enregistrer le document
        doc.save(output_path)

    def _create_metadata(self, doc: Dict[str, Any]) -> ModelMetadata:
        """Crée les métadonnées d'un document"""
        return ModelMetadata(
            id=str(doc['id']),
            name=doc.get('title', 'Sans titre'),
            type=ModelType.DOCUMENT,
            created_at=doc['created_at'],
            updated_at=doc['updated_at'],
            summary=f"{doc.get('title', 'Sans titre')} - {doc.get('type', 'Inconnu')}"
        )
        
    def _load_document_metadata(self) -> List[ModelMetadata]:
        """Charge les métadonnées des documents"""
        docs = self._load_documents()
        return [self._create_metadata(doc) for doc in docs]
        
    def _load_document_by_id(self, doc_id: str) -> Optional[Dict[str, Any]]:
        """Charge un document spécifique par son ID"""
        doc = next((d for d in self._load_documents() if d.get('id') == doc_id), None)
        return doc.copy() if doc else None


if __name__ == "__main__":
    # Exemple d'utilisation du modèle
    model = AppModel()
    
    # Exemple d'ajout d'un client
    client_data = {
        "name": "Jean Dupont",
        "company": "Société Exemple",
        "email": "jean.dupont@example.com",
        "phone": "+33 6 12 34 56 78",
        "address": "123 Rue Exemple, 75000 Paris"
    }
    client_id = model.add_client(client_data)
    print(f"Client ajouté avec l'ID: {client_id}")
    
    # Exemple d'ajout d'un modèle
    template_data = {
        "name": "Contrat Standard",
        "description": "Modèle de contrat standard pour nouveaux clients",
        "content": """# CONTRAT DE PRESTATION DE SERVICES

## Entre les soussignés :

**{{client_company}}**, représentée par {{client_name}},
Ci-après dénommée "LE CLIENT"

Et

**Vynal Docs**, 
Ci-après dénommée "LE PRESTATAIRE"

Il a été convenu ce qui suit :

1. Le prestataire s'engage à fournir les services suivants : {{services}}
2. Le montant de la prestation s'élève à {{montant}} euros HT.
3. Le présent contrat est conclu pour une durée de {{duree}} mois.

Fait à Paris, le {{current_date}}

Pour le Client,                     Pour le Prestataire,
{{client_name}}                     Vynal Docs
""",
        "variables": [
            {
                "name": "services",
                "description": "Description des services fournis",
                "default": "Rédaction de documents juridiques"
            },
            {
                "name": "montant",
                "description": "Montant HT de la prestation",
                "default": "1000"
            },
            {
                "name": "duree",
                "description": "Durée du contrat en mois",
                "default": "12"
            }
        ]
    }
    template_id = model.add_template(template_data)
    print(f"Modèle ajouté avec l'ID: {template_id}")
    
    # Exemple de génération de document
    custom_variables = {
        "services": "Développement d'une application sur mesure",
        "montant": "5000",
        "duree": "6"
    }
    document_path = model.generate_document(template_id, client_id, custom_variables)
    print(f"Document généré: {document_path}")
    
    # Exemple de sauvegarde
    backup_path = model.create_backup()
    print(f"Sauvegarde créée: {backup_path}")