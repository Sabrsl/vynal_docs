#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Vue de création de documents pour l'application Vynal Docs Automator
Interface moderne et fluide pour la création de documents
"""

import logging
import customtkinter as ctk
from typing import Any, Dict, List, Optional
import json
import os
import threading
import shutil
import tkinter.filedialog as filedialog
from ai.document_processor import AIDocumentProcessor
from utils.document_generator import DocumentGenerator
from PIL import Image
import datetime
import re
import uuid
import platform
import subprocess

logger = logging.getLogger("VynalDocsAutomator.DocumentCreatorView")

class DocumentCreatorView:
    """
    Vue de création de documents
    Interface utilisateur pour créer des documents avec assistance IA
    """
    
    def __init__(self, parent: ctk.CTk, app_model: Any) -> None:
        """
        Initialise la vue de création de document
        
        Args:
            parent: Fenêtre parente
            app_model: Modèle de l'application
        """
        self.parent = parent
        self.model = app_model
        
        # Créer le cadre principal
        self.frame = ctk.CTkFrame(parent)
        
        # Initialiser les variables d'état
        self.current_step = 0
        self.steps = [
            "Démarrer",
            "Type de document",
            "Client",
            "Validation",
            "Personnalisation", 
            "Finalisation"
        ]
        
        # État du processus de création
        self.selected_template = None
        self.client_info = {}
        self.document_data = {}
        self.missing_variables = {}
        self.preview_content = ""
        self.client_selected = False
        
        # Barre supérieure avec titre et bouton de fermeture
        self.header_frame = ctk.CTkFrame(self.frame, height=50)
        self.header_frame.pack(fill="x", padx=10, pady=5)
        self.header_frame.pack_propagate(False)
        
        # Titre de la fenêtre
        self.title_label = ctk.CTkLabel(
            self.header_frame,
            text="Création de document",
            font=("", 20, "bold")
        )
        self.title_label.pack(side="left", padx=20)
        
        # Bouton de fermeture
        self.close_btn = ctk.CTkButton(
            self.header_frame,
            text="✕",
            width=30,
            height=30,
            command=self.on_close_button,
            fg_color="transparent",
            text_color=("gray10", "gray90"),
            hover_color=("gray70", "gray30")
        )
        self.close_btn.pack(side="right", padx=5)
        
        # Zone principale avec défilement
        self.scroll_container = ctk.CTkScrollableFrame(self.frame)
        self.scroll_container.pack(fill="both", expand=True, padx=20, pady=10)
        
        # Zone de contenu principal
        self.main_frame = ctk.CTkFrame(self.scroll_container, fg_color="transparent")
        self.main_frame.pack(fill="both", expand=True)
        
        # Créer l'interface
        self.create_widgets()
        
        logger.info("Vue de création de documents initialisée")
        
    def create_widgets(self) -> None:
        """
        Crée les widgets de l'interface
        """
        # Barre de progression
        self.progress_frame = ctk.CTkFrame(self.main_frame)
        self.progress_frame.pack(fill="x", pady=(0, 20))
        
        # Créer les indicateurs d'étape
        self.step_indicators = []
        for i, step in enumerate(self.steps):
            step_frame = ctk.CTkFrame(self.progress_frame)
            step_frame.grid(row=0, column=i*2, padx=5, pady=5)
            
            # Numéro de l'étape
            number = ctk.CTkLabel(
                step_frame,
                text=str(i + 1),
                width=30,
                height=30,
                corner_radius=15,
                fg_color=("gray75", "gray25"),
                text_color="white"
            )
            number.pack(pady=5)
            
            # Nom de l'étape
            label = ctk.CTkLabel(step_frame, text=step, font=("", 12))
            label.pack(pady=2)
            
            self.step_indicators.append({"number": number, "label": label})
            
            # Ajouter un séparateur sauf pour la dernière étape
            if i < len(self.steps) - 1:
                separator = ctk.CTkFrame(
                    self.progress_frame,
                    width=50,
                    height=2,
                    fg_color=("gray75", "gray25")
                )
                separator.grid(row=0, column=i*2 + 1, sticky="ew", pady=20)
        
        # Zone de contenu principal
        self.content_frame = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        self.content_frame.pack(fill="both", expand=True, pady=10)
        
        # Zone de navigation
        self.nav_frame = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        self.nav_frame.pack(fill="x", pady=10)
        
        # Afficher la première étape
        self.show_step(0)
        
    def show_step(self, step: int) -> None:
        """
        Affiche l'étape spécifiée du processus
        
        Args:
            step: Numéro de l'étape à afficher
        """
        # S'assurer que l'indice est valide
        if step < 0 or step >= len(self.steps):
            logger.error(f"Indice d'étape invalide: {step}")
            return
        
        # Initialiser selected_template comme un dictionnaire vide s'il n'existe pas
        # pour éviter les erreurs AttributeError: 'NoneType' object has no attribute 'get'
        if not hasattr(self, 'selected_template') or self.selected_template is None:
            self.selected_template = {}
            
        # Mettre à jour les indicateurs d'étape
        for i, indicator in enumerate(self.step_indicators):
            if i < step:
                # Étape terminée
                indicator["number"].configure(
                    fg_color="green",
                    text="✓",
                    text_color="white"
                )
            elif i == step:
                # Étape en cours
                indicator["number"].configure(
                    fg_color="#1f538d",
                    text=str(i + 1),
                    text_color="white"
                )
            else:
                # Étape à venir
                indicator["number"].configure(
                    fg_color=("gray75", "gray25"),
                    text=str(i + 1),
                    text_color=("gray20", "gray80")
                )
                
        # Mettre à jour le titre de la fenêtre
        self.title_label.configure(text=f"Création de document - {self.steps[step]}")
        
        # Nettoyer la zone de contenu
        for widget in self.content_frame.winfo_children():
            widget.destroy()
            
        # Afficher le contenu de l'étape
        if step == 0:
            self.show_initial_options()
        elif step == 1:
            self.show_document_types()
        elif step == 2:
            self.show_client_form()
        elif step == 3:
            self.show_validation()
        elif step == 4:
            self.show_customization()
        elif step == 5:
            self.show_finalization()
            
        self.current_step = step
        
        # Mettre à jour les boutons de navigation
        self.update_navigation()
        
        # Faire défiler vers le haut du contenu - méthode corrigée
        try:
            # Utiliser _parent_canvas qui est l'attribut interne de CTkScrollableFrame
            if hasattr(self.scroll_container, '_parent_canvas'):
                self.scroll_container._parent_canvas.yview_moveto(0)
        except Exception as e:
            logger.warning(f"Impossible de faire défiler vers le haut: {e}")
        
        logger.info(f"Affichage de l'étape {step}: {self.steps[step]}")
        
    def show_initial_options(self) -> None:
        """
        Affiche les options initiales (nouveau document ou modèle existant)
        """
        # Nettoyer la zone de contenu
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Titre
        title = ctk.CTkLabel(
            self.content_frame,
            text="Choisissez une option",
            font=("", 24, "bold")  # Taille réduite
        )
        title.pack(pady=(10, 20))  # Padding vertical réduit
        
        # Ajouter un message de débogage pour s'assurer que cette méthode est bien appelée
        logger.info("Affichage des options initiales - traiter un document")

        # Frame pour les boutons
        buttons_frame = ctk.CTkFrame(self.content_frame, fg_color="transparent")
        buttons_frame.pack(pady=20, expand=True)  # Padding vertical réduit

        # Fonction pour créer un bouton avec une icône plus grande
        def create_button_with_large_icon(parent, icon, text, command, **kwargs):
            # Créer un cadre pour contenir l'icône et le texte
            button_frame = ctk.CTkFrame(parent, fg_color=kwargs.get('fg_color', ("gray90", "gray20")))
            button_frame.pack(side="left", padx=kwargs.get('padx', 20))  # Espacement horizontal réduit
            
            # Ajouter l'icône en grand
            icon_label = ctk.CTkLabel(
                button_frame,
                text=icon,
                font=("", 80),  # Taille d'icône réduite
                fg_color=kwargs.get('fg_color', ("gray90", "gray20")),
                text_color=kwargs.get('text_color', ("gray10", "gray90"))
            )
            icon_label.pack(pady=(20, 10))  # Padding vertical réduit
            
            # Ajouter le texte
            text_label = ctk.CTkLabel(
                button_frame,
                text=text,
                font=("", 18),  # Taille de police réduite
                fg_color=kwargs.get('fg_color', ("gray90", "gray20")),
                text_color=kwargs.get('text_color', ("gray10", "gray90"))
            )
            text_label.pack(pady=(0, 20))  # Padding vertical réduit
            
            # Configurer les dimensions du cadre
            button_frame.configure(
                width=kwargs.get('width', 450),
                height=kwargs.get('height', 450),
                corner_radius=kwargs.get('corner_radius', 20)
            )
            
            # Rendre le bouton cliquable
            button_frame.bind("<Button-1>", lambda e: command())
            icon_label.bind("<Button-1>", lambda e: command())
            text_label.bind("<Button-1>", lambda e: command())
            
            # Effet de survol
            hover_color = kwargs.get('hover_color', ("gray80", "gray30"))
            orig_color = kwargs.get('fg_color', ("gray90", "gray20"))
            
            def on_enter(e):
                button_frame.configure(fg_color=hover_color)
                icon_label.configure(fg_color=hover_color)
                text_label.configure(fg_color=hover_color)
                
            def on_leave(e):
                button_frame.configure(fg_color=orig_color)
                icon_label.configure(fg_color=orig_color)
                text_label.configure(fg_color=orig_color)
            
            button_frame.bind("<Enter>", on_enter)
            button_frame.bind("<Leave>", on_leave)
            icon_label.bind("<Enter>", on_enter)
            icon_label.bind("<Leave>", on_leave)
            text_label.bind("<Enter>", on_enter)
            text_label.bind("<Leave>", on_leave)
            
            # Configurer le curseur
            button_frame.configure(cursor="hand2")
            icon_label.configure(cursor="hand2")
            text_label.configure(cursor="hand2")
            
            return button_frame
        
        # Créer le bouton pour utiliser un modèle existant
        template_btn = create_button_with_large_icon(
            buttons_frame,
            icon="📂",
            text="Utiliser un modèle\nexistant",
            command=self._use_existing_template,
            width=350,  # Réduction de la taille
            height=300,  # Réduction de la taille
            fg_color=("gray90", "gray20"),
            hover_color=("gray80", "gray30")
        )
        
        # Créer le bouton pour importer un document
        import_btn = create_button_with_large_icon(
            buttons_frame,
            icon="📄",
            text="À partir d'un\ndocument",
            command=self.import_document,
            width=350,  # Réduction de la taille
            height=300,  # Réduction de la taille
            fg_color=("gray90", "gray20"),
            hover_color=("gray80", "gray30")
        )
        
        # S'assurer que les boutons sont visibles
        buttons_frame.update()
        
        # Mettre à jour la navigation pour s'assurer que les boutons sont affichés
        self.update_navigation()
    
    def _use_existing_template(self) -> None:
        """
        Crée directement un formulaire de création de document
        """
        try:
            # Masquer cette vue
            self.hide()
            
            # Créer directement un formulaire de document
            from views.document_form_view import DocumentFormView
            
            # Définir une fonction callback pour quand l'utilisateur annule
            def on_cancel_callback():
                # Réafficher cette vue
                self.show()
            
            # Définir une fonction callback pour après la sauvegarde du document
            def on_save_callback(document_id=None, client_id=None, client_name=None, **kwargs):
                logger.info(f"Document sauvegardé: {document_id} pour le client {client_name}")
                # Rediriger vers le tableau de bord après la sauvegarde
                try:
                    # Vérifier différentes options pour trouver la méthode show_view
                    if hasattr(self.parent, 'show_view'):
                        logger.info("Redirection vers le tableau de bord après sauvegarde via self.parent")
                        self.parent.show_view('dashboard')
                    elif 'parent' in locals() and hasattr(locals()['parent'], 'show_view'):
                        logger.info("Redirection vers le tableau de bord après sauvegarde via parent local")
                        locals()['parent'].show_view('dashboard')
                    else:
                        # Tenter de récupérer le parent via self
                        root = self.parent
                        # Remonter jusqu'à trouver un objet avec show_view
                        while root is not None:
                            if hasattr(root, 'show_view'):
                                logger.info("Redirection vers le tableau de bord après sauvegarde via recherche de parent")
                                root.show_view('dashboard')
                                return
                            if hasattr(root, 'master'):
                                root = root.master
                            else:
                                break
                                
                        logger.error("Impossible de rediriger vers le tableau de bord: méthode show_view non trouvée")
                        # Si nous ne pouvons pas rediriger, au moins essayer de réafficher cette vue
                        self.show()
                except Exception as e:
                    logger.error(f"Erreur lors de la redirection: {e}")
                    self.show()
            
            # Vérifier si le modèle possède la méthode save_document
            # Si ce n'est pas le cas, ajouter un adaptateur temporaire
            if not hasattr(self.model, 'save_document'):
                logger.info("Ajout d'un adaptateur save_document au modèle")
                
                # Créer un adaptateur qui sauvegarde un document
                def save_document_adapter(document):
                    """
                    Sauvegarde un document dans la base de données et le génère sous forme de fichier
                    """
                    try:
                        # Récupérer les informations nécessaires
                        title = document.get('name', 'Document')
                        client_id = document.get('client_id')
                        client_name = document.get('client_name', 'Client non spécifié')
                        
                        # Préparer le dossier de sortie
                        output_dir = os.path.join("data", "documents", "outputs")
                        os.makedirs(output_dir, exist_ok=True)
                        
                        # Générer un nom de fichier unique
                        from datetime import datetime
                        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                        filename_base = f"{title}_{client_name}_{timestamp}"
                        
                        # Nettoyer le nom du fichier (enlever les caractères spéciaux)
                        filename = re.sub(r'[^\w\-_\. ]', '_', filename_base)
                        
                        # Chemins pour différents formats de documents
                        pdf_path = os.path.join(output_dir, f"{filename}.pdf")
                        txt_path = os.path.join(output_dir, f"{filename}.txt")
                        
                        # Obtenir le contenu du modèle
                        template_content = document.get('content', '')
                        
                        # Si le modèle a un chemin de fichier, l'utiliser pour obtenir le contenu
                        if 'file_path' in document and document['file_path']:
                            try:
                                template_file = document['file_path']
                                if not os.path.isabs(template_file):
                                    # Chemin relatif - chercher dans les dossiers standard
                                    possible_paths = [
                                        os.path.join("data", "templates", template_file),
                                        os.path.join("data", "models", template_file),
                                        template_file
                                    ]
                                    
                                    for path in possible_paths:
                                        if os.path.exists(path):
                                            with open(path, 'r', encoding='utf-8') as f:
                                                template_content = f.read()
                                            logger.info(f"Contenu du modèle chargé depuis le fichier: {path}")
                                            break
                            except Exception as e:
                                logger.warning(f"Impossible de charger le contenu du fichier modèle: {e}")
                        
                        # Remplacer les variables dans le modèle
                        variables = {}
                        if 'variables' in document:
                            for var_name, var_info in document.get('variables', {}).items():
                                if isinstance(var_info, dict):
                                    variables[var_name] = var_info.get('current_value', '')
                                else:
                                    variables[var_name] = str(var_info)
                        
                        # Ajouter les variables standard
                        variables.update({
                            "client_name": client_name,
                            "date": datetime.now().strftime('%Y-%m-%d'),
                            "datetime": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                            "template_name": title
                        })
                        
                        # Tenter de récupérer les informations complètes du client
                        if client_id and hasattr(self.model, 'get_client'):
                            try:
                                client_data = self.model.get_client(client_id)
                                if client_data:
                                    # Ajouter les informations du client aux variables
                                    for key, value in client_data.items():
                                        if value:  # Ne pas ajouter les valeurs vides
                                            variables[f"client_{key}"] = value
                            except Exception as client_error:
                                logger.warning(f"Impossible de récupérer les informations du client {client_id}: {client_error}")
                        
                        # Remplacer les variables dans le contenu
                        final_content = template_content
                        for var_name, var_value in variables.items():
                            # Format {variable}
                            final_content = final_content.replace(f"{{{var_name}}}", str(var_value))
                            # Format {{variable}} (double accolades)
                            final_content = final_content.replace(f"{{{{{var_name}}}}}", str(var_value))
                        
                        # Tenter de générer un PDF
                        try:
                            # Utiliser ReportLab pour générer un PDF
                            from reportlab.lib.pagesizes import letter
                            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                            from reportlab.lib.styles import getSampleStyleSheet
                            
                            styles = getSampleStyleSheet()
                            doc = SimpleDocTemplate(pdf_path, pagesize=letter)
                            
                            # Convertir le contenu en paragraphes
                            content_lines = final_content.split('\n')
                            pdf_content = []
                            
                            for line in content_lines:
                                if line.strip():  # Ne pas ajouter les lignes vides
                                    pdf_content.append(Paragraph(line, styles['Normal']))
                                    pdf_content.append(Spacer(1, 6))  # Petit espace entre les lignes
                            
                            # Générer le PDF
                            doc.build(pdf_content)
                            
                            # Vérifier que le PDF a été généré correctement
                            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                                logger.info(f"Document PDF généré avec succès: {pdf_path}")
                                file_path = pdf_path
                            else:
                                raise Exception("Le fichier PDF généré est vide ou n'a pas été créé.")
                            
                        except Exception as pdf_error:
                            logger.error(f"Erreur lors de la génération du PDF: {pdf_error}")
                            
                            # En cas d'échec, générer un fichier texte
                            with open(txt_path, 'w', encoding='utf-8') as f:
                                f.write(final_content)
                            logger.info(f"Document texte créé: {txt_path}")
                            file_path = txt_path
                        
                        # Ajouter le chemin du fichier au document
                        document['file_path'] = file_path
                        
                        # Générer un ID si nécessaire
                        if 'id' not in document:
                            document['id'] = str(uuid.uuid4())
                            
                        # S'assurer que la liste des documents existe
                        if not hasattr(self.model, 'documents'):
                            self.model.documents = []
                            
                        # Ajouter ou mettre à jour le document
                        doc_updated = False
                        if isinstance(self.model.documents, list):
                            # Mettre à jour le document existant s'il existe
                            for i, doc in enumerate(self.model.documents):
                                if doc.get('id') == document.get('id'):
                                    self.model.documents[i] = document
                                    doc_updated = True
                                    break
                            
                            # Ajouter le document s'il n'existe pas
                            if not doc_updated:
                                self.model.documents.append(document)
                        
                        # Sauvegarder la liste des documents
                        if hasattr(self.model, 'save_documents'):
                            self.model.save_documents()
                        
                        # Planifier la redirection vers le tableau de bord
                        # Utiliser after pour différer l'exécution et éviter les erreurs de fenêtre
                        def redirect_to_dashboard():
                            try:
                                # Vérifier différentes options pour trouver la méthode show_view
                                if hasattr(self.parent, 'show_view'):
                                    logger.info("Redirection vers le tableau de bord via self.parent")
                                    self.parent.after(100, lambda: self.parent.show_view('dashboard'))
                                    return True
                                elif hasattr(self, 'parent') and hasattr(self.parent, 'master') and hasattr(self.parent.master, 'show_view'):
                                    logger.info("Redirection vers le tableau de bord via self.parent.master")
                                    self.parent.after(100, lambda: self.parent.master.show_view('dashboard'))
                                    return True
                                else:
                                    # Remonter la hiérarchie des widgets
                                    root = self.parent
                                    max_depth = 5  # Limiter la profondeur de recherche pour éviter les boucles infinies
                                    depth = 0
                                    while root is not None and depth < max_depth:
                                        if hasattr(root, 'show_view'):
                                            logger.info(f"Redirection vers le tableau de bord trouvée à la profondeur {depth}")
                                            root.after(100, lambda r=root: r.show_view('dashboard'))
                                            return True
                                        if hasattr(root, 'master'):
                                            root = root.master
                                        else:
                                            break
                                        depth += 1
                                            
                                    logger.error("Impossible de rediriger vers le tableau de bord: méthode show_view non trouvée")
                                    return False
                            except Exception as e:
                                logger.error(f"Erreur lors de la tentative de redirection: {e}")
                                return False
                        
                        # Exécuter la redirection
                        redirect_success = redirect_to_dashboard()
                        
                        # Si la redirection a échoué, tenter de revenir à l'état précédent
                        if not redirect_success:
                            self.show()
                        
                        return document
                    except Exception as e:
                        logger.error(f"Erreur dans l'adaptateur save_document: {e}")
                        # Tenter de revenir à l'état précédent
                        self.show()
                        return None
                
                # Attacher l'adaptateur au modèle
                setattr(self.model, 'save_document', save_document_adapter)
            
            # Créer le formulaire avec les données vides
            form = DocumentFormView(
                self.parent,
                self.model,
                document_data={},
                on_save_callback=on_save_callback,  # Utiliser notre nouveau callback
                on_cancel_callback=on_cancel_callback
            )
            
            logger.info("Création directe du formulaire de document")
            
        except Exception as e:
            logger.error(f"Erreur lors de la création du formulaire de document: {e}")
            self.show_error(f"Erreur: {str(e)}")
            # En cas d'erreur, réafficher cette vue
            self.show()
            
    def _on_document_saved(self, document_id=None, client_id=None, client_name=None, **kwargs):
        """
        Callback appelé après la sauvegarde d'un document
        """
        logger.info(f"Document sauvegardé: {document_id} pour le client {client_name}")
        
        # Rediriger vers le tableau de bord
        try:
            # Vérifier différentes options pour trouver la méthode show_view
            if hasattr(self.parent, 'show_view'):
                logger.info("Redirection vers le tableau de bord depuis _on_document_saved via self.parent")
                self.parent.show_view('dashboard')
            else:
                # Tenter de récupérer le parent via self
                root = self.parent
                # Remonter jusqu'à trouver un objet avec show_view
                while root is not None:
                    if hasattr(root, 'show_view'):
                        logger.info("Redirection vers le tableau de bord via recherche de parent")
                        root.show_view('dashboard')
                        return
                    if hasattr(root, 'master'):
                        root = root.master
                    else:
                        break
                        
                logger.error("Impossible de rediriger vers le tableau de bord: méthode show_view non trouvée")
                # Si on ne peut pas rediriger, au moins essayer de réafficher la vue
                self.show()
        except Exception as e:
            logger.error(f"Erreur lors de la redirection: {e}")
            # Si une erreur se produit, essayer de réafficher la vue
            self.show()

    def show_document_types(self) -> None:
        """
        Affiche les types de documents disponibles sous forme de cartes
        """
        # Nettoyer la zone de contenu
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Titre
        title = ctk.CTkLabel(
            self.content_frame,
            text="Types de documents disponibles",
            font=("", 24, "bold")
        )
        title.pack(pady=(0, 20))

        # Grille pour les cartes
        grid_frame = ctk.CTkFrame(self.content_frame, fg_color="transparent")
        grid_frame.pack(expand=True, fill="both", padx=20, pady=20)

        # Configuration de la grille (3 colonnes)
        grid_frame.grid_columnconfigure((0, 1, 2), weight=1)

        # Charger les types depuis le dossier data/documents/types
        types_dir = os.path.join("data", "documents", "types")
        if not os.path.exists(types_dir):
            # Créer le dossier s'il n'existe pas
            os.makedirs(types_dir)
            logger.info(f"Dossier créé : {types_dir}")

        # Liste pour stocker les types de documents
        document_types = []
        
        # Dictionnaire de correspondance type -> icône
        type_icons = {
            "contrat": "📜",
            "facture": "💰",
            "devis": "📋",
            "lettre": "✉️",
            "rapport": "📊",
            "presentation": "🎯",
            "proposition": "💡",
            "convention": "🤝",
            "certificat": "🏆",
            "attestation": "📝",
            "formulaire": "📄",
            "note": "📌",
            "procès-verbal": "📋",
            "plan": "🗺️",
            "budget": "💰",
            "planning": "📅",
            "inventaire": "📦",
            "catalogue": "📚",
            "manuel": "📖",
            "guide": "📗",
            "tutoriel": "📘",
            "documentation": "📑",
            "spécification": "📋",
            "analyse": "🔍",
            "étude": "📊",
            "projet": "🎯",
            "résumé": "📝",
            "synthèse": "📊",
            "évaluation": "📈",
            "audit": "🔍",
            "statistiques": "📈",
            "graphique": "📊",
            "tableau": "📋",
            "liste": "📋",
            "registre": "📑",
            "journal": "📓",
            "carnet": "📔",
            "agenda": "📅",
            "calendrier": "📅",
            "horaire": "⏰",
            "emploi du temps": "📅",
            "programme": "📋",
            "plan d'action": "🎯",
            "stratégie": "🎯",
            "objectif": "🎯",
            "mission": "🎯",
            "vision": "👀",
            "suggestion": "💡",
            "recommandation": "💡",
            "avis": "💡",
            "expertise": "🔍",
            "consultation": "💡",
            "recherche": "🔍",
            "enquête": "🔍",
            "sondage": "📊",
            "questionnaire": "📋",
            "demande": "📝",
            "requête": "📝",
            "pétition": "📝",
            "plainte": "📝",
            "réclamation": "📝"
        }
        
        # Parcourir le dossier des types
        if os.path.exists(types_dir):
            for type_name in os.listdir(types_dir):
                type_path = os.path.join(types_dir, type_name)
                if os.path.isdir(type_path):
                    # Compter le nombre de modèles dans le dossier
                    model_count = len([f for f in os.listdir(type_path) if f.endswith(('.docx', '.pdf', '.txt'))])
                    
                    # Déterminer l'icône en fonction du type de dossier
                    icon = "📁"  # Icône par défaut
                    type_lower = type_name.lower()
                    for key, value in type_icons.items():
                        if key in type_lower:
                            icon = value
                            break
                    
                    # Ajouter le type à la liste
                    document_types.append({
                        "name": type_name,
                        "icon": icon,
                        "count": f"{model_count} modèle{'s' if model_count > 1 else ''}"
                    })

        # Créer une carte pour chaque type
        for i, doc_type in enumerate(document_types):
            # Créer un cadre pour la carte
            card = ctk.CTkFrame(
                grid_frame,
                corner_radius=10,
                fg_color=("gray90", "gray20"),
                width=200,
                height=150
            )
            row = i // 3
            col = i % 3
            card.grid(row=row, column=col, padx=10, pady=10, sticky="nsew")
            card.grid_propagate(False)

            # Contenu de la carte
            icon_label = ctk.CTkLabel(
                card,
                text=doc_type["icon"],
                font=("", 32)
            )
            icon_label.pack(pady=(20, 5))

            name_label = ctk.CTkLabel(
                card,
                text=doc_type["name"],
                font=("", 16, "bold")
            )
            name_label.pack(pady=5)

            count_label = ctk.CTkLabel(
                card,
                text=doc_type["count"],
                font=("", 12),
                text_color=("gray50", "gray70")
            )
            count_label.pack(pady=5)

            # Rendre la carte cliquable
            for widget in [card, icon_label, name_label, count_label]:
                widget.bind("<Button-1>", lambda e, type_name=doc_type["name"]: self.show_models_for_type(type_name))
                widget.bind("<Enter>", lambda e, frame=card: frame.configure(fg_color=("gray80", "gray30")))
                widget.bind("<Leave>", lambda e, frame=card: frame.configure(fg_color=("gray90", "gray20")))
                widget.configure(cursor="hand2")

        # Message si aucun type n'est trouvé
        if not document_types:
            no_types_label = ctk.CTkLabel(
                grid_frame,
                text="Aucun type de document trouvé\ndans le dossier data/documents/types",
                font=("", 14),
                text_color=("gray50", "gray70")
            )
            no_types_label.pack(pady=50)

        # Bouton retour
        back_button = ctk.CTkButton(
            self.content_frame,
            text="Retour",
            command=lambda: self.show_step(0),
            width=100
        )
        back_button.pack(pady=20)

    def show_models_for_type(self, type_name: str) -> None:
        """
        Affiche les modèles disponibles pour un type de document donné
        
        Args:
            type_name: Nom du type de document sélectionné
        """
        # Nettoyer la zone de contenu
        for widget in self.content_frame.winfo_children():
            widget.destroy()

        # Titre avec le type de document
        title = ctk.CTkLabel(
            self.content_frame,
            text=f"Modèles de {type_name.lower()}",
            font=("", 24, "bold")
        )
        title.pack(pady=(0, 20))

        # Grille pour les cartes
        grid_frame = ctk.CTkFrame(self.content_frame, fg_color="transparent")
        grid_frame.pack(expand=True, fill="both", padx=20, pady=20)

        # Configuration de la grille (3 colonnes)
        grid_frame.grid_columnconfigure((0, 1, 2), weight=1)

        # Charger les modèles depuis le dossier du type
        type_path = os.path.join("data", "documents", "types", type_name)
        models = []

        if os.path.exists(type_path):
            for model_name in os.listdir(type_path):
                if model_name.endswith(('.docx', '.pdf', '.txt')):
                    # Déterminer l'icône en fonction du type de fichier
                    icon = "📄" if model_name.endswith('.txt') else "📝" if model_name.endswith('.docx') else "📋"
                    
                    # Ajouter le modèle à la liste
                    models.append({
                        "name": os.path.splitext(model_name)[0],  # Nom sans extension
                        "icon": icon,
                        "file": model_name  # Nom complet avec extension
                    })

        # Créer une carte pour chaque modèle
        for i, model in enumerate(models):
            # Créer un cadre pour la carte
            card = ctk.CTkFrame(
                grid_frame,
                corner_radius=10,
                fg_color=("gray90", "gray20"),
                width=200,
                height=150
            )
            row = i // 3
            col = i % 3
            card.grid(row=row, column=col, padx=10, pady=10, sticky="nsew")
            card.grid_propagate(False)

            # Contenu de la carte
            icon_label = ctk.CTkLabel(
                card,
                text=model["icon"],
                font=("", 32)
            )
            icon_label.pack(pady=(20, 5))

            name_label = ctk.CTkLabel(
                card,
                text=model["name"],
                font=("", 16, "bold")
            )
            name_label.pack(pady=5)

            # Rendre la carte cliquable
            for widget in [card, icon_label, name_label]:
                widget.bind("<Button-1>", lambda e, t=type_name, m=model["file"]: self.select_model(t, m))
                widget.bind("<Enter>", lambda e, frame=card: frame.configure(fg_color=("gray80", "gray30")))
                widget.bind("<Leave>", lambda e, frame=card: frame.configure(fg_color=("gray90", "gray20")))
                widget.configure(cursor="hand2")

        # Message si aucun modèle n'est trouvé
        if not models:
            no_models_label = ctk.CTkLabel(
                grid_frame,
                text=f"Aucun modèle trouvé pour le type {type_name}",
                font=("", 14),
                text_color=("gray50", "gray70")
            )
            no_models_label.pack(pady=50)

        # Bouton retour
        back_button = ctk.CTkButton(
            self.content_frame,
            text="Retour aux types de documents",
            command=self.show_document_types,
            width=200
        )
        back_button.pack(pady=20)

    def select_model(self, type_name: str, model_file: str) -> None:
        """
        Sélectionne un modèle et passe à l'étape suivante
        
        Args:
            type_name: Nom du type de document
            model_file: Nom du fichier modèle
        """
        try:
            self.selected_type = type_name
            self.selected_model = model_file
            self.selected_template = {
                "type": type_name,
                "name": os.path.splitext(model_file)[0],
                "file": model_file,
                "path": os.path.join("data", "documents", "types", type_name, model_file)
            }
            logging.info(f"Modèle sélectionné : {model_file} ({type_name})")
            self.show_step(2)  # Passer à l'étape du formulaire client
        except Exception as e:
            logger.error(f"Erreur lors de la sélection du modèle: {e}")
            self.show_error("Une erreur est survenue lors de la sélection du modèle")

    def show_client_form(self) -> None:
        """
        Affiche le formulaire simplifié de recherche client
        """
        # Si on vient d'un document importé et que l'analyse n'est pas terminée
        if (hasattr(self, 'selected_template') and 
            self.selected_template.get('from_analysis') and 
            not getattr(self, 'analysis_complete', False)):
            
            # Nettoyer la zone de contenu
            for widget in self.content_frame.winfo_children():
                widget.destroy()
            
            # Créer un cadre conteneur pour centrer le contenu
            container = ctk.CTkFrame(self.content_frame, fg_color="transparent")
            container.pack(fill="both", expand=True, padx=20, pady=20)
            
            # Titre
            title = ctk.CTkLabel(
                container,
                text="Analyse du document en cours",
                font=("", 24, "bold")
            )
            title.pack(pady=(0, 20))
            
            # Message de chargement
            loading_label = ctk.CTkLabel(
                container,
                text=f"Analyse de {self.selected_template['name']}...\nCette opération peut prendre quelques instants.",
                font=("", 14),
                justify="center"
            )
            loading_label.pack(pady=20)
            
            # Barre de progression indéterminée
            progress = ctk.CTkProgressBar(container, width=400)
            progress.pack(pady=10)
            progress.start()  # Animation de chargement
            
            # Cadre pour afficher les variables détectées (initialement vide)
            variables_frame = ctk.CTkScrollableFrame(container, fg_color=("gray95", "gray15"), height=200)
            variables_frame.pack(fill="x", padx=20, pady=10)
            variables_frame.pack_forget()  # Masquer jusqu'à ce que l'analyse soit terminée
            
            # Afficher la barre de recherche client même pendant l'analyse
            search_frame = ctk.CTkFrame(container, fg_color=("transparent"))
            search_frame.pack(fill="x", padx=20, pady=(30, 10))
            
            # Message d'information
            info_label = ctk.CTkLabel(
                search_frame,
                text="Veuillez attendre la fin de l'analyse avant de sélectionner un client.",
                font=("", 12),
                text_color=("gray50", "gray70"),
                wraplength=600,
                justify="center"
            )
            info_label.pack(pady=(0, 10))
            
            # Cadre de recherche (désactivé pendant l'analyse)
            search_input_frame = ctk.CTkFrame(search_frame, fg_color=("gray90", "gray20"))
            search_input_frame.pack(fill="x", padx=20, pady=10)
            
            # Champ de recherche (désactivé)
            self.search_entry = ctk.CTkEntry(
                search_input_frame,
                placeholder_text="Nom, email, téléphone ou entreprise du client...",
                width=400,
                state="disabled",
                fg_color=("gray80", "gray30")
            )
            self.search_entry.pack(side="left", padx=10, pady=10)
            
            # Bouton de recherche (désactivé)
            search_button = ctk.CTkButton(
                search_input_frame,
                text="Rechercher",
                width=100,
                command=self.search_client,
                state="disabled",
                fg_color=("gray75", "gray45")
            )
            search_button.pack(side="left", padx=10, pady=10)
            
            # Zone d'affichage des résultats avec défilement
            self.results_frame = ctk.CTkScrollableFrame(container, fg_color="transparent", height=200)
            self.results_frame.pack(fill="both", expand=True, padx=20, pady=10)
            
            # Message indiquant que la recherche est désactivée
            disabled_msg = ctk.CTkLabel(
                self.results_frame,
                text="La recherche de clients sera disponible une fois l'analyse terminée",
                font=("", 12),
                text_color=("gray50", "gray70")
            )
            disabled_msg.pack(pady=50)
            
            # Vérifier périodiquement si l'analyse est terminée
            def check_analysis():
                if hasattr(self, 'analysis_error') and self.analysis_error:
                    # Arrêter l'animation
                    progress.stop()
                    # Afficher l'erreur
                    error_label = ctk.CTkLabel(
                        container,
                        text=f"Erreur lors de l'analyse : {self.analysis_error}",
                        text_color="red",
                        wraplength=600
                    )
                    error_label.pack(pady=10)
                    return
                
                if getattr(self, 'analysis_complete', False):
                    # Analyse terminée, arrêter le spinner
                    progress.stop()
                    
                    # Vérifier si des variables ont été trouvées
                    if hasattr(self, 'document_data') and self.document_data.get('variables'):
                        # Variables trouvées, mettre à jour l'interface
                        title.configure(text="Analyse du document terminée")
                        loading_label.configure(
                            text="Analyse terminée avec succès.",
                            text_color=("green", "#00AA00")
                        )
                        
                        # Afficher les variables détectées
                        variables_frame.pack(fill="x", padx=20, pady=10)
                        
                        # Titre des variables
                        vars_title = ctk.CTkLabel(
                            variables_frame,
                            text="Variables détectées :",
                            font=("", 14, "bold")
                        )
                        vars_title.pack(anchor="w", padx=10, pady=(10, 5))
                        
                        # Afficher chaque variable
                        var_count = 0
                        for var_name, var_info in self.document_data.get('variables', {}).items():
                            var_count += 1
                            var_frame = ctk.CTkFrame(variables_frame, fg_color="transparent")
                            var_frame.pack(fill="x", padx=10, pady=3)
                            
                            # Nom de la variable
                            var_label = ctk.CTkLabel(
                                var_frame,
                                text=f"{var_name}:",
                                font=("", 12, "bold"),
                                width=150
                            )
                            var_label.pack(side="left")
                            
                            # Valeur courante
                            current_value = var_info.get('current_value', '') if isinstance(var_info, dict) else str(var_info)
                            value_label = ctk.CTkLabel(
                                var_frame,
                                text=current_value,
                                font=("", 12)
                            )
                            value_label.pack(side="left", padx=10)
                        
                        # Message de succès
                        count_label = ctk.CTkLabel(
                            variables_frame,
                            text=f"{var_count} variables trouvées",
                            font=("", 12),
                            text_color=("gray50", "gray70")
                        )
                        count_label.pack(anchor="e", padx=10, pady=(5, 10))
                        
                        # Changer le message d'information
                        info_label.configure(
                            text="Vous pouvez maintenant sélectionner un client pour continuer :",
                            text_color=("black", "white")
                        )
                        
                        # Activer les éléments de recherche
                        self.search_entry.configure(state="normal", fg_color=("white", "gray20"))
                        search_button.configure(state="normal", fg_color=("#3a7ebf", "#1f538d"))
                        
                        # Nettoyer les résultats (mais ne pas afficher tous les clients)
                        for widget in self.results_frame.winfo_children():
                            widget.destroy()
                            
                        # Afficher un message pour inviter à la recherche
                        search_prompt = ctk.CTkLabel(
                            self.results_frame,
                            text="Saisissez un terme pour rechercher un client",
                            font=("Arial", 12)
                        )
                        search_prompt.pack(pady=20, padx=10)
                        
                        # Activer les boutons de navigation
                        self.update_navigation()
                    else:
                        # Aucune variable trouvée
                        title.configure(text="Aucune variable détectée")
                        loading_label.configure(
                            text="Aucune variable n'a été détectée dans le document.\nVeuillez essayer avec un autre document.",
                            text_color="red"
                        )
                        
                        # Désactiver la recherche client
                        info_label.configure(
                            text="Aucune variable n'a été détectée. Veuillez réinitialiser et essayer avec un autre document.",
                            text_color="red"
                        )
                else:
                    # Vérifier à nouveau dans 500ms
                    self.content_frame.after(500, check_analysis)
            
            # Démarrer la vérification
            self.content_frame.after(500, check_analysis)
            
            # Mettre à jour la navigation (boutons suivant/précédent)
            self.update_navigation()
            return
            
        # Nettoyer la zone de contenu
        for widget in self.content_frame.winfo_children():
            widget.destroy()
            
        # Créer un cadre conteneur pour centrer le contenu
        container = ctk.CTkFrame(self.content_frame, fg_color="transparent")
        container.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Titre
        title = ctk.CTkLabel(
            container,
            text="Recherche du client",
            font=("", 24, "bold")
        )
        title.pack(pady=(0, 20))

        # Message d'information
        info_label = ctk.CTkLabel(
            container,
            text="Recherchez un client pour associer à ce document :",
            font=("", 12),
            text_color=("gray50", "gray70"),
            wraplength=600,
            justify="center"
        )
        info_label.pack(pady=(0, 20))

        # Cadre de recherche
        search_frame = ctk.CTkFrame(container)
        search_frame.pack(fill="x", padx=20, pady=10)

        # Champ de recherche
        self.search_entry = ctk.CTkEntry(
            search_frame,
            placeholder_text="Nom, email, téléphone ou entreprise du client...",
            width=400
        )
        self.search_entry.pack(side="left", padx=10, pady=10)
        
        # Bouton de recherche
        search_button = ctk.CTkButton(
            search_frame,
            text="Rechercher",
            width=100,
            command=self.search_client
        )
        search_button.pack(side="left", padx=10, pady=10)

        # Zone d'affichage des résultats avec défilement
        self.results_frame = ctk.CTkScrollableFrame(container, fg_color="transparent")
        self.results_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        # Afficher un message initial pour guider l'utilisateur
        guide_label = ctk.CTkLabel(
            self.results_frame,
            text="Saisissez un terme pour rechercher un client",
            font=("Arial", 12)
        )
        guide_label.pack(pady=20, padx=10)
        
        # Forcer la mise à jour de l'affichage
        self.content_frame.update()
        container.update()
        self.results_frame.update()
        
        # Lier l'événement Return/Enter au champ de recherche
        self.search_entry.bind("<Return>", lambda e: self.search_client())
        
        # Lier l'événement de modification du texte pour la recherche en temps réel
        self.search_entry.bind("<KeyRelease>", lambda e: self.search_client())
        
        # S'assurer que le focus est sur le champ de recherche
        self.search_entry.focus_set()
        
        # Mettre à jour la navigation (boutons suivant/précédent)
        self.update_navigation()

    def search_client(self) -> None:
        """
        Recherche un client dans la base de données
        Utilise le texte saisi dans le champ de recherche
        """
        try:
            # Nettoyer la zone de résultats
            for widget in self.results_frame.winfo_children():
                if widget.winfo_exists():  # Vérifier que le widget existe toujours
                    widget.destroy()
                
            logger.info("Recherche de clients démarrée")
                
            # Récupérer la valeur de recherche
            search_value = self.search_entry.get().strip().lower() if hasattr(self, 'search_entry') else ""
            
            # Si la recherche est vide, ne pas afficher de clients
            if not search_value:
                logger.info("Aucun terme de recherche, affichage de tous les clients")
                # Afficher un message indiquant qu'il faut effectuer une recherche
                message = ctk.CTkLabel(
                    self.results_frame,
                    text="Saisissez un terme pour rechercher un client",
                    font=("Arial", 12)
                )
                message.pack(pady=20, padx=10)
                return
                
            # Récupérer les clients depuis le modèle
            clients = []
            if hasattr(self.model, 'clients'):
                clients = getattr(self.model, 'clients')
                logger.info(f"Nombre de clients disponibles: {len(clients)}")
                
            # Si self.model.clients n'est pas accessible, essayer de charger les clients depuis le fichier JSON
            if not clients:
                clients_file = os.path.join("data", "clients", "clients.json")
                if os.path.exists(clients_file):
                    try:
                        with open(clients_file, "r", encoding="utf-8") as f:
                            clients = json.load(f)
                        logger.info(f"Clients chargés depuis le fichier: {len(clients)}")
                    except Exception as e:
                        logger.error(f"Erreur lors de la lecture du fichier clients: {e}")
            
            # Filtrer les clients selon la recherche
            filtered_clients = []
            for client in clients:
                for key, value in client.items():
                    if isinstance(value, str) and search_value in value.lower():
                        filtered_clients.append(client)
                        break
            logger.info(f"Recherche '{search_value}': {len(filtered_clients)} clients trouvés")
            
            # Afficher les résultats
            if not filtered_clients:
                no_results = ctk.CTkLabel(
                    self.results_frame,
                    text=f"Aucun client trouvé pour '{search_value}'",
                    font=("Arial", 12)
                )
                no_results.pack(pady=20, padx=10)
                return
            
            # Afficher les clients filtrés avec l'ancien design
            for client in filtered_clients:
                # Créer un cadre pour chaque client
                client_frame = ctk.CTkFrame(self.results_frame)
                client_frame.pack(fill="x", padx=10, pady=5)
                
                # Nom du client (ou société)
                name = client.get('nom', '') or client.get('name', '') or client.get('société', '') or "Client sans nom"
                
                # Email et téléphone pour l'affichage
                email = client.get('email', '')
                phone = client.get('telephone', '')
                
                # Informations du client
                info_text = f"{name}"
                if email:
                    info_text += f" - {email}"
                if phone:
                    info_text += f" - {phone}"
                    
                client_label = ctk.CTkLabel(
                    client_frame,
                    text=info_text,
                    font=("Arial", 12),
                    anchor="w",
                    justify="left"
                )
                client_label.pack(side="left", padx=10, pady=5, fill="x", expand=True)
                
                # Créer une copie locale du client pour éviter les problèmes de référence dans les lambdas
                client_data = client.copy()
                
                # Fonction sécurisée pour sélectionner un client
                def create_safe_select_callback(client_data):
                    def safe_select_callback(event=None):
                        # Vérifier que le cadre parent existe toujours avant d'appeler la sélection
                        if hasattr(self, 'results_frame') and self.results_frame.winfo_exists():
                            self._select_client(client_data)
                    return safe_select_callback
                
                # Bouton de sélection
                select_button = ctk.CTkButton(
                    client_frame,
                    text="Sélectionner",
                    width=100,
                    command=create_safe_select_callback(client_data)
                )
                select_button.pack(side="right", padx=10, pady=5)
                
                # Rendre la ligne cliquable
                safe_callback = create_safe_select_callback(client_data)
                client_frame.bind("<Button-1>", lambda e, cb=safe_callback: cb())
                client_label.bind("<Button-1>", lambda e, cb=safe_callback: cb())
                
                # Configurer le curseur
                client_frame.configure(cursor="hand2")
                client_label.configure(cursor="hand2")
                
        except Exception as e:
            logger.error(f"Erreur lors de la recherche de clients: {e}")
            error_label = ctk.CTkLabel(
                self.results_frame,
                text=f"Erreur: {str(e)}",
                text_color="red",
                font=("Arial", 12)
            )
            error_label.pack(pady=20, padx=10)

    def on_close_button(self) -> None:
        """
        Gère la fermeture de la vue lorsque l'utilisateur clique sur le bouton X
        """
        try:
            # Réinitialiser les variables
            self.current_step = 0
            self.selected_template = None
            self.client_info = {}
            self.document_data = {}
            self.missing_variables = {}
            self.preview_content = ""
            self.client_selected = False
            
            # Revenir à la vue précédente (dashboard)
            if hasattr(self.parent, 'show_view'):
                self.parent.show_view('dashboard')
            
            logger.info("Vue document_creator fermée")
        except Exception as e:
            logger.error(f"Erreur lors de la fermeture de la vue document_creator: {e}")
            # Tenter de revenir à la vue dashboard en cas d'erreur
            try:
                if hasattr(self.parent, 'show_view'):
                    self.parent.show_view('dashboard')
            except:
                pass
                
    def import_document(self) -> None:
        """
        Importe un document externe pour analyse
        Ouvre une boîte de dialogue de fichier
        """
        try:
            # Ouvrir une boîte de dialogue pour sélectionner un fichier
            filetypes = [
                ("Documents", "*.pdf *.docx *.doc *.odt *.txt"),
                ("PDF", "*.pdf"),
                ("Word", "*.docx *.doc"),
                ("OpenDocument", "*.odt"),
                ("Texte", "*.txt"),
                ("Tous les fichiers", "*.*")
            ]
            file_path = filedialog.askopenfilename(
                title="Sélectionner un document",
                filetypes=filetypes
            )
            
            if not file_path:
                # L'utilisateur a annulé la sélection
                return
                
            logger.info(f"Document sélectionné pour analyse: {file_path}")
            
            # Créer un template virtuel basé sur le document importé
            self.selected_template = {
                "type": "import",
                "name": os.path.basename(file_path),
                "file": os.path.basename(file_path),
                "path": file_path,
                "from_analysis": True
            }
            
            # Passer directement à l'étape de recherche de client
            self.show_step(2)
            
            # Démarrer l'analyse du document
            self._start_document_analysis(file_path)
            
        except Exception as e:
            logger.error(f"Erreur lors de l'importation du document: {e}")
            self.show_error(f"Erreur lors de l'importation: {str(e)}")
            
    def _start_document_analysis(self, file_path: str) -> None:
        """
        Démarre l'analyse d'un document avec l'IA
        
        Args:
            file_path: Chemin du fichier à analyser
        """
        try:
            # Initialiser le processeur de documents s'il n'existe pas déjà
            if not hasattr(self, 'document_processor'):
                self.document_processor = AIDocumentProcessor()
                
            logger.info(f"Début de l'analyse du document: {file_path}")
            
            # Indiquer que l'analyse est en cours
            self.analysis_in_progress = True
            self.analysis_complete = False
            self.analysis_error = None
            
            # Créer un thread pour l'analyse afin de ne pas bloquer l'interface
            def analyze_document_thread():
                try:
                    # Effectuer l'analyse
                    result = self.document_processor.analyze_document(file_path)
                    
                    # Stocker les données résultantes
                    self.document_data = result
                    
                    # Marquer l'analyse comme terminée
                    self.analysis_in_progress = False
                    self.analysis_complete = True
                    
                    logger.info("Analyse du document terminée")
                    
                except Exception as e:
                    self.analysis_in_progress = False
                    self.analysis_complete = True
                    self.analysis_error = str(e)
                    logger.error(f"Erreur lors de l'analyse du document: {e}")
            
            # Démarrer le thread d'analyse
            analysis_thread = threading.Thread(target=analyze_document_thread)
            analysis_thread.daemon = True
            analysis_thread.start()
            
        except Exception as e:
            logger.error(f"Erreur lors du démarrage de l'analyse: {e}")
            self.analysis_in_progress = False
            self.analysis_error = str(e)
            
    def _apply_client_to_analysis(self, client: Dict[str, Any]) -> None:
        """
        Applique les informations du client sélectionné à l'analyse en cours
        
        Args:
            client: Dictionnaire contenant les informations du client
        """
        try:
            # Vérifier que les widgets nécessaires existent toujours
            if not hasattr(self, 'content_frame') or not self.content_frame.winfo_exists():
                logger.error("Le cadre de contenu n'existe plus, abandon de l'application du client")
                return
            
            # Stocker les informations du client
            self.client_info = client
            self.client_selected = True
            
            # Passer à l'étape suivante si l'analyse est terminée, mais avec un délai pour éviter les problèmes de fenêtre
            if self.analysis_complete:
                # Utiliser after pour différer l'exécution et éviter les erreurs de fenêtre
                self.parent.after(100, lambda: self._safe_show_step(3))
                
        except Exception as e:
            logger.error(f"Erreur lors de l'application du client à l'analyse: {e}")
            self.show_error(f"Erreur: {str(e)}")
            
    def _safe_show_step(self, step: int) -> None:
        """
        Version sécurisée de show_step qui vérifie si les widgets existent toujours
        
        Args:
            step: Numéro de l'étape à afficher
        """
        try:
            # Vérifier que les widgets nécessaires existent toujours
            if hasattr(self, 'content_frame') and self.content_frame.winfo_exists():
                self.show_step(step)
            else:
                logger.error(f"Tentative d'afficher l'étape {step} mais les widgets n'existent plus")
        except Exception as e:
            logger.error(f"Erreur lors de l'affichage sécurisé de l'étape {step}: {e}")
            
    def show_error(self, message: str) -> None:
        """
        Affiche un message d'erreur dans une boîte de dialogue
        
        Args:
            message: Message d'erreur à afficher
        """
        try:
            # Utiliser une boîte de dialogue CTk
            from customtkinter import CTkToplevel, CTkLabel, CTkButton
            
            # Créer une fenêtre modale
            error_window = CTkToplevel(self.parent)
            error_window.title("Erreur")
            error_window.geometry("400x200")
            error_window.grab_set()  # Rendre la fenêtre modale
            
            # Message d'erreur
            error_label = CTkLabel(
                error_window,
                text=message,
                wraplength=350,
                text_color="red"
            )
            error_label.pack(padx=20, pady=30, expand=True)
            
            # Bouton OK
            ok_button = CTkButton(
                error_window,
                text="OK",
                width=100,
                command=error_window.destroy
            )
            ok_button.pack(pady=20)
            
            # Centrer la fenêtre
            error_window.update_idletasks()
            width = error_window.winfo_width()
            height = error_window.winfo_height()
            x = (error_window.winfo_screenwidth() // 2) - (width // 2)
            y = (error_window.winfo_screenheight() // 2) - (height // 2)
            error_window.geometry(f"{width}x{height}+{x}+{y}")
            
        except Exception as e:
            # Fallback pour les erreurs dans la gestion des erreurs
            logger.error(f"Erreur lors de l'affichage du message d'erreur: {e}")
            # Tenter d'utiliser la méthode standard
            from tkinter import messagebox
            messagebox.showerror("Erreur", message)
 
    def update_navigation(self) -> None:
        """
        Met à jour les boutons de navigation selon l'étape actuelle
        """
        try:
            # Nettoyer la zone de navigation
            for widget in self.nav_frame.winfo_children():
                widget.destroy()
                
            logger.info(f"Mise à jour de la navigation à l'étape {self.current_step}")
            
            # Bouton de retour (sauf pour la première étape)
            if self.current_step > 0:
                back_btn = ctk.CTkButton(
                    self.nav_frame,
                    text="← Retour",
                    command=lambda: self.show_step(self.current_step - 1),
                    width=100,
                    fg_color=("gray80", "gray30"),
                    text_color=("gray10", "gray90")
                )
                back_btn.pack(side="left", padx=10)
                
            # Bouton de réinitialisation
            reset_btn = ctk.CTkButton(
                self.nav_frame,
                text="↺ Réinitialiser",
                command=self.reset_process,
                width=100,
                fg_color=("gray80", "gray30"),
                text_color=("gray10", "gray90")
            )
            reset_btn.pack(side="left", padx=10)
            
            # Bouton suivant (sauf pour la première étape qui utilise directement les grands boutons)
            if self.current_step > 0:
                next_btn = ctk.CTkButton(
                    self.nav_frame,
                    text="Suivant →",
                    command=lambda: self.show_step(self.current_step + 1),
                    width=100
                )
                
                # Déterminer si le bouton suivant doit être activé
                enabled = True
                
                # Gestion de l'état du bouton selon l'étape
                if self.current_step == 2:  # Étape client
                    if not getattr(self, 'client_selected', False):
                        enabled = False
                    # Si on est en cours d'analyse, désactiver le bouton
                    if getattr(self, 'analysis_in_progress', False) and not getattr(self, 'analysis_complete', False):
                        enabled = False
                
                # Configurer l'état du bouton
                if not enabled:
                    next_btn.configure(state="disabled", fg_color=("gray75", "gray45"))
                
                next_btn.pack(side="right", padx=10)
            
        except Exception as e:
            logger.error(f"Erreur lors de la mise à jour de la navigation: {e}")
            
    def reset_process(self) -> None:
        """
        Réinitialise le processus de création
        """
        try:
            # Réinitialiser les variables
            self.current_step = 0
            self.selected_template = None
            self.client_info = {}
            self.document_data = {}
            self.missing_variables = {}
            self.preview_content = ""
            self.client_selected = False
            self.analysis_in_progress = False
            self.analysis_complete = False
            self.analysis_error = None
            
            # Retourner à la première étape
            self.show_step(0)
            
            logger.info("Processus de création réinitialisé")
        except Exception as e:
            logger.error(f"Erreur lors de la réinitialisation: {e}")

    def _select_client(self, client: Dict[str, Any]) -> None:
        """
        Sélectionne un client et passe à l'étape suivante
        
        Args:
            client: Dictionnaire contenant les informations du client sélectionné
        """
        try:
            if not client:
                logger.warning("Tentative de sélection d'un client vide")
                return
                
            # Masquer l'interface de recherche si elle est affichée
            if hasattr(self, 'search_frame') and self.search_frame:
                self.search_frame.destroy()
                
            # Nettoyer les données précédentes si nécessaire
            if hasattr(self, 'client_fields'):
                self.client_fields = {}
                
            # Stocker les informations du client
            self.client_info = client
            self.client_selected = True
            logger.info(f"Client sélectionné: {client.get('nom', 'Inconnu')} {client.get('prénom', '')}")
            
            # Si nous n'avons pas de document_data, initialiser avec un dictionnaire vide
            if not hasattr(self, 'document_data'):
                self.document_data = {'variables': {}}
            elif not self.document_data:
                self.document_data = {'variables': {}}
            elif 'variables' not in self.document_data:
                self.document_data['variables'] = {}
                
            # Si nous sommes dans l'étape de sélection client, passer à l'étape suivante
            # Sinon, mettre à jour les champs avec les informations du client
            if self.current_step == 2:
                # Animation de transition avant de passer à l'étape suivante
                self._clean_interface_for_transition(3)
                
                # Fonction pour gérer la transition après l'animation
                def handle_transition():
                    # Mettre à jour les informations du client dans document_data
                    if 'client_info' not in self.document_data:
                        self.document_data['client_info'] = self.client_info
                    
                    # Remplir automatiquement les variables avec les informations client
                    self._auto_fill_variables()
                    
                    # Passer à l'étape suivante
                    self._complete_transition()
                
                # Définir le délai pour la transition
                transition_delay = 300  # milliseconds
                self.parent.after(transition_delay, handle_transition)
            else:
                # Nous sommes dans une autre étape, mettre simplement à jour le client
                # et remplir les variables avec les nouvelles informations
                
                # Mettre à jour les informations du client dans document_data
                if 'client_info' not in self.document_data:
                    self.document_data['client_info'] = self.client_info
                
                # Remplir automatiquement les variables avec les informations client
                self._auto_fill_variables()
                
                # Mettre à jour l'interface si nécessaire
                if self.current_step == 3 or self.current_step == 4:
                    # Rafraîchir l'affichage des variables
                    self.show_step(self.current_step)
                
                logger.info("Informations client mises à jour")
                self.show_message("Client sélectionné", 
                                 f"Client {client.get('nom', 'Inconnu')} {client.get('prénom', '')} sélectionné avec succès.",
                                 "success")
                
        except Exception as e:
            logger.error(f"Erreur lors de la sélection du client: {e}")
            self.show_error(f"Erreur lors de la sélection du client: {str(e)}")

    def _auto_fill_variables(self) -> None:
        """
        Remplit automatiquement les variables du document avec les informations du client
        """
        if not hasattr(self, 'document_data') or not self.document_data or not hasattr(self, 'client_info') or not self.client_info:
            return
            
        # Vérifier si nous avons des variables
        if 'variables' not in self.document_data or not self.document_data['variables']:
            return
            
        # Dictionnaire des correspondances possibles entre noms de variables et champs client
        var_to_client_field = {
            # Mappings directs (nom variable -> champ client)
            'nom': ['nom', 'name', 'nomcomplet', 'nom_complet', 'fullname', 'full_name'],
            'prénom': ['prénom', 'prenom', 'firstname', 'first_name'],
            'adresse': ['adresse', 'address', 'rue', 'street'],
            'ville': ['ville', 'city', 'town'],
            'code_postal': ['code_postal', 'codepostal', 'cp', 'zip', 'zipcode', 'zip_code'],
            'pays': ['pays', 'country'],
            'email': ['email', 'courriel', 'mail', 'e_mail'],
            'téléphone': ['téléphone', 'telephone', 'tel', 'phone', 'phonenumber', 'phone_number'],
            'mobile': ['mobile', 'cellulaire', 'portable', 'cell', 'cellphone'],
            'société': ['société', 'societe', 'entreprise', 'company', 'organization'],
            'siret': ['siret', 'siren', 'numéro_siret', 'numero_siret'],
            'tva': ['tva', 'vat', 'numéro_tva', 'numero_tva'],
            'date_naissance': ['date_naissance', 'datenaissance', 'birthday', 'birthdate', 'birth_date'],
            'profession': ['profession', 'métier', 'metier', 'job', 'occupation'],
        }
        
        # Pour les mappings avec préfixes
        client_prefixes = ['client', 'clt', 'customer', 'beneficiary', 'bénéficiaire']
        
        # Parcourir toutes les variables du document
        for var_name in list(self.document_data['variables'].keys()):
            # Variable déjà avec une valeur?
            var_info = self.document_data['variables'][var_name]
            current_value = ""
            if isinstance(var_info, dict):
                current_value = var_info.get('current_value', '')
            else:
                current_value = str(var_info)
                
            # Ne pas remplacer si déjà une valeur
            if current_value.strip():
                continue
                
            # Cas 1: Variable avec préfixe client_xxx (le plus facile)
            if var_name.startswith('client_') or var_name.startswith('clt_') or var_name.startswith('customer_'):
                # Extraire le nom du champ après le préfixe
                field_name = var_name.split('_', 1)[1]
                
                # Vérifier si ce champ existe dans les données client
                if field_name in self.client_info and self.client_info[field_name]:
                    self._set_variable_value(var_name, self.client_info[field_name], True)
                    continue
            
            # Cas 2: Vérifier si la variable correspond directement à un champ client
            if var_name in self.client_info and self.client_info[var_name]:
                self._set_variable_value(var_name, self.client_info[var_name], True)
                continue
                
            # Cas 3: Reconnaissance intelligente basée sur le nom de la variable
            var_name_lower = var_name.lower()
            matched = False
            
            # Rechercher dans les mappings standard
            for target_field, possible_names in var_to_client_field.items():
                # Vérifier si le nom de variable correspond à une des possibilités
                if var_name_lower in possible_names or any(name in var_name_lower for name in possible_names):
                    # Chercher le champ correspondant dans les données client
                    for client_field in self.client_info:
                        client_field_lower = client_field.lower()
                        if client_field_lower in possible_names or target_field == client_field_lower:
                            if self.client_info[client_field]:
                                self._set_variable_value(var_name, self.client_info[client_field], True)
                                matched = True
                                break
                if matched:
                    break
                    
            if matched:
                continue
                
            # Cas 4: Vérifier les variables avec d'autres préfixes courants
            for prefix in client_prefixes:
                if var_name_lower.startswith(f"{prefix}_") or var_name_lower.startswith(f"{prefix}."):
                    # Extraire le nom du champ après le préfixe
                    separator = "_" if var_name_lower.startswith(f"{prefix}_") else "."
                    field_name = var_name_lower.split(separator, 1)[1]
                    
                    # Chercher une correspondance dans les champs client
                    for client_field in self.client_info:
                        if field_name == client_field.lower() or client_field.lower().endswith(field_name):
                            if self.client_info[client_field]:
                                self._set_variable_value(var_name, self.client_info[client_field], True)
                                matched = True
                                break
                if matched:
                    break
            
            # Cas 5: Cas spéciaux
            if not matched:
                # Nom complet
                if var_name_lower in ['nom_complet', 'nomcomplet', 'fullname', 'full_name', 'nom_et_prénom', 'nom_prenom']:
                    if 'nom' in self.client_info and 'prénom' in self.client_info:
                        full_name = f"{self.client_info['prénom']} {self.client_info['nom']}"
                        self._set_variable_value(var_name, full_name, True)
                    elif 'nom' in self.client_info:
                        self._set_variable_value(var_name, self.client_info['nom'], True)
                
                # Adresse complète
                elif var_name_lower in ['adresse_complete', 'adressecomplete', 'full_address', 'fulladdress']:
                    address_parts = []
                    if 'adresse' in self.client_info and self.client_info['adresse']:
                        address_parts.append(self.client_info['adresse'])
                    if 'code_postal' in self.client_info and self.client_info['code_postal']:
                        address_parts.append(self.client_info['code_postal'])
                    if 'ville' in self.client_info and self.client_info['ville']:
                        address_parts.append(self.client_info['ville'])
                    if 'pays' in self.client_info and self.client_info['pays']:
                        address_parts.append(self.client_info['pays'])
                        
                    if address_parts:
                        full_address = " ".join(address_parts)
                        self._set_variable_value(var_name, full_address, True)
                        
        logger.info(f"Variables automatiquement remplies avec les informations client")
        
    def _set_variable_value(self, var_name: str, value: Any, auto_detected: bool = False) -> None:
        """
        Définit la valeur d'une variable dans le document
        
        Args:
            var_name: Nom de la variable
            value: Valeur à définir
            auto_detected: Indique si la valeur a été détectée automatiquement
        """
        if not hasattr(self, 'document_data') or not self.document_data or 'variables' not in self.document_data:
            return
            
        if var_name not in self.document_data['variables']:
            return
            
        # Conversion en chaîne si nécessaire
        str_value = str(value) if value is not None else ""
        
        # Mise à jour selon le type de stockage
        if isinstance(self.document_data['variables'][var_name], dict):
            self.document_data['variables'][var_name]['current_value'] = str_value
            if auto_detected:
                self.document_data['variables'][var_name]['auto_detected'] = True
        else:
            # Convertir en dictionnaire pour stocker plus d'informations
            self.document_data['variables'][var_name] = {
                'current_value': str_value,
                'auto_detected': auto_detected
            }
            
        logger.debug(f"Variable '{var_name}' définie à '{str_value}' (auto-détecté: {auto_detected})")

    def _clean_interface_for_transition(self, target_step: int) -> None:
        """
        Méthode spéciale pour nettoyer l'interface avant une transition d'étape délicate
        Cette méthode évite les erreurs "bad window path name"
        
        Args:
            target_step: Numéro de l'étape cible
        """
        try:
            # Enregistrer l'étape cible
            self.pending_step = target_step
            
            # Nettoyer entièrement la zone de contenu
            if hasattr(self, 'content_frame') and self.content_frame.winfo_exists():
                for widget in self.content_frame.winfo_children():
                    if widget.winfo_exists():
                        widget.destroy()
                
                # Créer un message de transition
                transition_label = ctk.CTkLabel(
                    self.content_frame,
                    text="Chargement...",
                    font=("", 18)
                )
                transition_label.pack(pady=100)
                
                # Forcer la mise à jour de l'interface
                self.content_frame.update_idletasks()
                
                # Planifier l'affichage de l'étape suivante
                self.parent.after(100, lambda: self._complete_transition())
            else:
                logger.error("Le cadre de contenu n'existe plus lors de la transition")
        except Exception as e:
            logger.error(f"Erreur lors du nettoyage pour la transition: {e}")
            # Tenter quand même d'afficher l'étape cible
            self.parent.after(200, lambda: self.show_step(target_step))
            
    def _complete_transition(self) -> None:
        """
        Complète la transition vers l'étape cible après une animation de chargement
        """
        try:
            # Vérifier si nous avons une étape cible enregistrée
            if hasattr(self, 'pending_step'):
                target_step = self.pending_step
                # Supprimer l'attribut pour éviter de réutiliser cette valeur
                delattr(self, 'pending_step')
                
                # Afficher l'étape cible si elle est valide
                if 0 <= target_step < len(self.steps):
                    # Appeler directement la méthode show_step avec l'étape cible
                    self.show_step(target_step)
                    logger.info(f"Transition complétée vers l'étape {target_step}")
                else:
                    logger.error(f"Étape cible invalide pour la transition: {target_step}")
            else:
                logger.warning("Tentative de compléter une transition sans étape cible définie")
        except Exception as e:
            logger.error(f"Erreur lors de la complétion de la transition: {e}")
            
    def show_validation(self) -> None:
        """
        Affiche la page de validation du document
        """
        try:
            # Nettoyer d'abord l'interface
            self._clear_content_frame()
            
            # Vérifier si nous avons des données de document
            if not hasattr(self, 'document_data') or not self.document_data:
                self.show_error("Aucune donnée de document disponible.")
                return
                
            # Mettre à jour la barre de progression
            self.update_navigation()
            
            # Configuration du titre
            title_frame = ctk.CTkFrame(self.content_frame, fg_color="transparent")
            title_frame.pack(fill="x", padx=20, pady=(20, 10), anchor="n")
            
            title_label = ctk.CTkLabel(
                title_frame,
                text="Validation du document",
                font=("", 24, "bold")
            )
            title_label.pack(side="left", pady=10)
            
            # Définir des couleurs de thème par défaut si self.model.theme_colors n'existe pas
            theme_colors = {
                'frame_secondary_bg': ("gray90", "gray20"),
                'separator': ("gray75", "gray35"),
                'text_muted': ("gray50", "gray70"),
                'auto_detected_bg': ("#e6f7ff", "#1a3c4d"),
                'heading_bg': ("#f0f0f0", "#2d2d2d"),
                'field_bg': ("#ffffff", "#333333"),
                'field_border': ("#e0e0e0", "#555555"),
                'success_bg': ("#e6ffe6", "#1a4d1a")
            }
            
            # Utilisez les couleurs du modèle si disponibles
            if hasattr(self.model, 'theme_colors'):
                theme_colors = self.model.theme_colors
            
            # Cadre principal de validation
            validation_frame = ctk.CTkFrame(self.content_frame)
            validation_frame.pack(fill="both", expand=True, padx=20, pady=20)
            
            # Créer un en-tête explicatif avec icône
            header_frame = ctk.CTkFrame(
                validation_frame,
                fg_color=theme_colors['heading_bg'],
                corner_radius=10,
                height=80
            )
            header_frame.pack(fill="x", padx=20, pady=(20, 15))
            header_frame.pack_propagate(False)
            
            # Icône d'information
            check_icon = "✓"
            info_label = ctk.CTkLabel(
                header_frame,
                text=check_icon,
                font=("", 24)
            )
            info_label.pack(side="left", padx=(15, 0))
            
            # Cadre pour le texte
            header_text_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
            header_text_frame.pack(side="left", fill="both", expand=True, padx=15, pady=10)
            
            # Titre d'information
            header_title = ctk.CTkLabel(
                header_text_frame,
                text="Vérification finale",
                font=("", 16, "bold"),
                anchor="w"
            )
            header_title.pack(anchor="w")
            
            # Description
            description_text = (
                "Veuillez vérifier les informations ci-dessous avant de générer le document final."
            )
            description_label = ctk.CTkLabel(
                header_text_frame,
                text=description_text,
                font=("", 12),
                wraplength=700,
                justify="left",
                anchor="w"
            )
            description_label.pack(anchor="w", fill="x", pady=(5, 0))
            
            # Cadre pour le résumé du document
            summary_container = ctk.CTkScrollableFrame(
                validation_frame,
                fg_color="transparent",
                corner_radius=0,
                height=400
            )
            summary_container.pack(fill="both", expand=True, padx=20, pady=(15, 20))
            
            # Cadre des informations du modèle
            model_frame = ctk.CTkFrame(
                summary_container,
                fg_color=theme_colors['field_bg'],
                corner_radius=10,
                border_width=1,
                border_color=theme_colors['field_border']
            )
            model_frame.pack(fill="x", padx=10, pady=10)
            
            # Titre du cadre modèle
            model_title = ctk.CTkLabel(
                model_frame,
                text="Informations du document",
                font=("", 16, "bold")
            )
            model_title.pack(anchor="w", padx=15, pady=(15, 10))
            
            # Grille pour les informations du modèle
            model_grid = ctk.CTkFrame(model_frame, fg_color="transparent")
            model_grid.pack(fill="x", padx=15, pady=(0, 15))
            model_grid.columnconfigure(0, weight=1)
            model_grid.columnconfigure(1, weight=3)
            
            # Récupérer les informations du modèle
            model_info = [
                ("Type de document", self.document_data.get('type_name', 'Non spécifié')),
                ("Modèle", self.document_data.get('model_name', 'Non spécifié')),
                ("Date de création", datetime.datetime.now().strftime('%Y-%m-%d %H:%M'))
            ]
            
            # Afficher les informations du modèle
            for idx, (label, value) in enumerate(model_info):
                # Libellé
                ctk.CTkLabel(
                    model_grid,
                    text=label,
                    font=("", 12, "bold"),
                    anchor="w"
                ).grid(row=idx, column=0, sticky="w", padx=5, pady=3)
                
                # Valeur
                ctk.CTkLabel(
                    model_grid,
                    text=str(value),
                    font=("", 12),
                    anchor="w"
                ).grid(row=idx, column=1, sticky="w", padx=5, pady=3)
            
            # Cadre des informations client
            client_frame = ctk.CTkFrame(
                summary_container,
                fg_color=theme_colors['field_bg'],
                corner_radius=10,
                border_width=1,
                border_color=theme_colors['field_border']
            )
            client_frame.pack(fill="x", padx=10, pady=10)
            
            # Titre du cadre client
            client_title = ctk.CTkLabel(
                client_frame,
                text="Informations du client",
                font=("", 16, "bold")
            )
            client_title.pack(anchor="w", padx=15, pady=(15, 10))
            
            # Grille pour les informations client
            client_grid = ctk.CTkFrame(client_frame, fg_color="transparent")
            client_grid.pack(fill="x", padx=15, pady=(0, 15))
            client_grid.columnconfigure(0, weight=1)
            client_grid.columnconfigure(1, weight=3)
            
            # Afficher les informations client si disponibles
            if hasattr(self, 'client_info') and self.client_info:
                client_display_fields = [
                    ('nom', 'Nom'),
                    ('prénom', 'Prénom'),
                    ('société', 'Société/Organisation'),
                    ('email', 'Email'),
                    ('téléphone', 'Téléphone'),
                    ('adresse', 'Adresse')
                ]
                
                row_idx = 0
                for field_key, field_label in client_display_fields:
                    # Vérifier si le champ existe et a une valeur
                    if field_key in self.client_info and self.client_info[field_key]:
                        # Libellé
                        ctk.CTkLabel(
                            client_grid,
                            text=field_label,
                            font=("", 12, "bold"),
                            anchor="w"
                        ).grid(row=row_idx, column=0, sticky="w", padx=5, pady=3)
                        
                        # Valeur
                        ctk.CTkLabel(
                            client_grid,
                            text=str(self.client_info[field_key]),
                            font=("", 12),
                            anchor="w"
                        ).grid(row=row_idx, column=1, sticky="w", padx=5, pady=3)
                        
                        row_idx += 1
            else:
                # Aucune information client
                ctk.CTkLabel(
                    client_grid,
                    text="Aucune information client disponible",
                    font=("", 12),
                    text_color=theme_colors['text_muted']
                ).grid(row=0, column=0, columnspan=2, sticky="w", padx=5, pady=10)
            
            # Cadre des variables
            variables_frame = ctk.CTkFrame(
                summary_container,
                fg_color=theme_colors['field_bg'],
                corner_radius=10,
                border_width=1,
                border_color=theme_colors['field_border']
            )
            variables_frame.pack(fill="x", padx=10, pady=10)
            
            # Titre du cadre variables
            variables_title = ctk.CTkLabel(
                variables_frame,
                text="Variables du document",
                font=("", 16, "bold")
            )
            variables_title.pack(anchor="w", padx=15, pady=(15, 10))
            
            # Grille pour les variables
            variables_grid = ctk.CTkFrame(variables_frame, fg_color="transparent")
            variables_grid.pack(fill="x", padx=15, pady=(0, 15))
            variables_grid.columnconfigure(0, weight=1)
            variables_grid.columnconfigure(1, weight=3)
            
            # Afficher les variables si disponibles
            if 'variables' in self.document_data and self.document_data['variables']:
                row_idx = 0
                for var_name, var_info in self.document_data['variables'].items():
                    # Récupérer la valeur actuelle
                    current_value = ""
                    is_auto_detected = False
                    
                    if isinstance(var_info, dict):
                        current_value = var_info.get('current_value', '')
                        is_auto_detected = var_info.get('auto_detected', False)
                    else:
                        current_value = str(var_info)
                    
                    # Libellé avec indication auto-détecté si applicable
                    label_text = var_name
                    if is_auto_detected:
                        label_text = f"{var_name} (auto)"
                    
                    ctk.CTkLabel(
                        variables_grid,
                        text=label_text,
                        font=("", 12, "bold"),
                        anchor="w"
                    ).grid(row=row_idx, column=0, sticky="w", padx=5, pady=3)
                    
                    # Valeur
                    ctk.CTkLabel(
                        variables_grid,
                        text=str(current_value),
                        font=("", 12),
                        anchor="w"
                    ).grid(row=row_idx, column=1, sticky="w", padx=5, pady=3)
                    
                    row_idx += 1
            else:
                # Aucune variable
                ctk.CTkLabel(
                    variables_grid,
                    text="Aucune variable disponible",
                    font=("", 12),
                    text_color=theme_colors['text_muted']
                ).grid(row=0, column=0, columnspan=2, sticky="w", padx=5, pady=10)
            
            # Message de confirmation
            confirmation_frame = ctk.CTkFrame(
                validation_frame,
                fg_color=theme_colors['success_bg'],
                corner_radius=10
            )
            confirmation_frame.pack(fill="x", padx=20, pady=(0, 15))
            
            confirmation_label = ctk.CTkLabel(
                confirmation_frame,
                text="Toutes les informations nécessaires ont été collectées. Vous pouvez maintenant confirmer le document.",
                font=("", 12),
                wraplength=700,
                justify="center"
            )
            confirmation_label.pack(pady=10)
            
            # Boutons de navigation
            buttons_frame = ctk.CTkFrame(self.content_frame, fg_color="transparent")
            buttons_frame.pack(fill="x", padx=20, pady=(10, 20))
            
            # Bouton retour
            back_button = ctk.CTkButton(
                buttons_frame,
                text="Retour",
                width=120,
                height=40,
                fg_color=("gray70", "gray30"),
                hover_color=("gray60", "gray40"),
                command=lambda: self.show_step(self.current_step - 1)
            )
            back_button.pack(side="left", padx=10, pady=10)
            
            # Bouton de confirmation
            confirm_button = ctk.CTkButton(
                buttons_frame,
                text="Confirmer et continuer",
                width=180,
                height=40,
                font=("", 14, "bold"),
                command=self._confirm_document
            )
            confirm_button.pack(side="right", padx=10, pady=10)
            
            # Mettre à jour l'étape actuelle
            self.current_step = 3
            
            logger.info("Affichage de l'étape de validation")
            
        except Exception as e:
            logger.error(f"Erreur lors de l'affichage de l'étape de validation: {e}")
            self.show_error(f"Erreur: {str(e)}")
            
    def _confirm_document(self) -> None:
        """
        Confirme la création du document et passe à l'étape de personnalisation
        """
        try:
            # Passer à l'étape de personnalisation
            self.show_step(4)
            
            logger.info("Document confirmé, passage à l'étape de personnalisation")
            
        except Exception as e:
            logger.error(f"Erreur lors de la confirmation du document: {e}")
            self.show_error(f"Erreur: {str(e)}")
            
    def show_customization(self) -> None:
        """
        Affiche l'étape de personnalisation du document
        """
        try:
            # Nettoyer d'abord l'interface
            self._clear_content_frame()
            
            # Vérifier si nous avons des données de document
            if not hasattr(self, 'document_data') or not self.document_data:
                self.show_error("Aucune donnée de document disponible.")
                return
            
            # Mettre à jour la barre de progression
            self.update_navigation()
            
            # Configuration du titre
            title_frame = ctk.CTkFrame(self.content_frame, fg_color="transparent")
            title_frame.pack(fill="x", padx=20, pady=(20, 10), anchor="n")
            
            title_label = ctk.CTkLabel(
                title_frame,
                text="Personnalisez votre document",
                font=("", 24, "bold")
            )
            title_label.pack(side="left", pady=10)
            
            # S'assurer que les variables sont remplies avec les informations client
            if hasattr(self, 'client_info') and self.client_info:
                self._auto_fill_variables()
            
            # Créer le cadre principal pour la personnalisation
            customization_frame = ctk.CTkFrame(self.content_frame)
            customization_frame.pack(fill="both", expand=True, padx=20, pady=20)
            
            # Définir des couleurs de thème par défaut si self.model.theme_colors n'existe pas
            theme_colors = {
                'frame_secondary_bg': ("gray90", "gray20"),
                'separator': ("gray75", "gray35"),
                'text_muted': ("gray50", "gray70"),
                'auto_detected_bg': ("#e6f7ff", "#1a3c4d"),
                'heading_bg': ("#f0f0f0", "#2d2d2d"),
                'field_bg': ("#ffffff", "#333333"),
                'field_border': ("#e0e0e0", "#555555")
            }
            
            # Utilisez les couleurs du modèle si disponibles
            if hasattr(self.model, 'theme_colors'):
                theme_colors = self.model.theme_colors
            
            # Créer un en-tête explicatif avec icône
            header_frame = ctk.CTkFrame(
                customization_frame,
                fg_color=theme_colors['heading_bg'],
                corner_radius=10,
                height=80
            )
            header_frame.pack(fill="x", padx=20, pady=(20, 15))
            header_frame.pack_propagate(False)
            
            # Icône d'information
            info_icon = "ℹ️"
            info_label = ctk.CTkLabel(
                header_frame,
                text=info_icon,
                font=("", 24)
            )
            info_label.pack(side="left", padx=(15, 0))
            
            # Cadre pour le texte
            header_text_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
            header_text_frame.pack(side="left", fill="both", expand=True, padx=15, pady=10)
            
            # Titre d'information
            header_title = ctk.CTkLabel(
                header_text_frame,
                text="Personnalisation du document",
                font=("", 16, "bold"),
                anchor="w"
            )
            header_title.pack(anchor="w")
            
            # Description
            description_text = (
                "Vérifiez et complétez les informations ci-dessous. Les variables pré-remplies "
                "sont affichées en haut (surlignées en bleu) et celles à compléter sont en bas."
            )
            description_label = ctk.CTkLabel(
                header_text_frame,
                text=description_text,
                font=("", 12),
                wraplength=700,
                justify="left",
                anchor="w"
            )
            description_label.pack(anchor="w", fill="x", pady=(5, 0))
            
            # Créer un cadre pour les variables
            variables_container = ctk.CTkScrollableFrame(
                customization_frame,
                fg_color="transparent",
                corner_radius=0,
                height=400  # Hauteur fixe pour un meilleur contrôle de l'interface
            )
            variables_container.pack(fill="both", expand=True, padx=20, pady=(15, 20))
            
            # Si aucune variable n'est disponible, afficher un message
            if not self.document_data.get('variables') or not isinstance(self.document_data['variables'], dict):
                no_vars_frame = ctk.CTkFrame(
                    variables_container,
                    corner_radius=10,
                    fg_color=theme_colors['field_bg'],
                    border_width=1,
                    border_color=theme_colors['field_border']
                )
                no_vars_frame.pack(fill="x", padx=10, pady=10)
                
                no_vars_label = ctk.CTkLabel(
                    no_vars_frame,
                    text="Aucune variable n'a été détectée dans ce document.",
                    font=("", 14),
                    text_color=theme_colors['text_muted']
                )
                no_vars_label.pack(pady=20)
            else:
                # Afficher les variables
                self.variable_entries = {}
                
                # Séparer les variables auto-détectées de celles à remplir
                auto_detected_vars = []
                to_fill_vars = []
                
                # Trier les variables 
                for var_name, var_info in self.document_data['variables'].items():
                    is_auto_detected = False
                    if isinstance(var_info, dict):
                        is_auto_detected = var_info.get('auto_detected', False)
                    
                    if is_auto_detected:
                        auto_detected_vars.append((var_name, var_info))
                    else:
                        to_fill_vars.append((var_name, var_info))
                
                # Trier les variables par ordre alphabétique dans chaque groupe
                auto_detected_vars.sort(key=lambda x: x[0].lower())
                to_fill_vars.sort(key=lambda x: x[0].lower())
                
                # Afficher d'abord les variables auto-détectées si elles existent
                if auto_detected_vars:
                    # En-tête pour les variables auto-détectées
                    auto_detected_header = ctk.CTkFrame(
                        variables_container,
                        fg_color=theme_colors['heading_bg'],
                        corner_radius=10,
                        height=30
                    )
                    auto_detected_header.pack(fill="x", padx=10, pady=(15, 5))
                    auto_detected_header.pack_propagate(False)
                    
                    auto_detected_label = ctk.CTkLabel(
                        auto_detected_header,
                        text="Variables pré-remplies",
                        font=("", 14, "bold")
                    )
                    auto_detected_label.pack(side="left", padx=15, pady=5)
                    
                    # Afficher les variables auto-détectées
                    for var_name, var_info in auto_detected_vars:
                        current_value = ""
                        if isinstance(var_info, dict):
                            current_value = var_info.get('current_value', '')
                        else:
                            current_value = str(var_info)
                        
                        # Créer un cadre pour cette variable avec un fond bleu
                        var_frame = ctk.CTkFrame(
                            variables_container,
                            fg_color=theme_colors['auto_detected_bg'],
                            corner_radius=8,
                            border_width=1,
                            border_color=theme_colors['field_border']
                        )
                        var_frame.pack(fill="x", pady=5, ipady=5, padx=10)
                        
                        # Conteneur pour le nom et la valeur
                        content_frame = ctk.CTkFrame(var_frame, fg_color="transparent")
                        content_frame.pack(fill="x", padx=10, pady=5)
                        content_frame.grid_columnconfigure(0, weight=2)
                        content_frame.grid_columnconfigure(1, weight=3)
                        
                        # Étiquette pour le nom de la variable
                        var_label = ctk.CTkLabel(
                            content_frame,
                            text=var_name,
                            font=("", 13, "bold"),
                            wraplength=250,
                            justify="left"
                        )
                        var_label.grid(row=0, column=0, sticky="w", padx=(5, 15), pady=5)
                        
                        # Champ de saisie pour la valeur
                        var_entry = ctk.CTkEntry(
                            content_frame,
                            width=300,
                            height=35,
                            placeholder_text="Entrez une valeur"
                        )
                        var_entry.grid(row=0, column=1, sticky="ew", padx=5, pady=5)
                        
                        # Ajouter un indicateur auto-détecté
                        auto_label = ctk.CTkLabel(
                            content_frame,
                            text="Auto-détecté",
                            font=("", 10),
                            text_color=("gray60", "gray70")
                        )
                        auto_label.grid(row=1, column=0, sticky="w", padx=(5, 15))
                        
                        # Mettre la valeur actuelle dans le champ
                        if current_value:
                            var_entry.insert(0, current_value)
                        
                        # Stocker l'entrée pour pouvoir récupérer la valeur plus tard
                        self.variable_entries[var_name] = var_entry
                
                # Ajouter un séparateur si les deux types de variables sont présents
                if auto_detected_vars and to_fill_vars:
                    separator = ctk.CTkFrame(
                        variables_container,
                        height=2,
                        fg_color=theme_colors['separator']
                    )
                    separator.pack(fill="x", padx=20, pady=15)
                
                # Afficher ensuite les variables à remplir
                if to_fill_vars:
                    # En-tête pour les variables à remplir
                    to_fill_header = ctk.CTkFrame(
                        variables_container,
                        fg_color=theme_colors['heading_bg'],
                        corner_radius=10,
                        height=30
                    )
                    to_fill_header.pack(fill="x", padx=10, pady=(15 if not auto_detected_vars else 5, 5))
                    to_fill_header.pack_propagate(False)
                    
                    to_fill_label = ctk.CTkLabel(
                        to_fill_header,
                        text="Variables à compléter",
                        font=("", 14, "bold")
                    )
                    to_fill_label.pack(side="left", padx=15, pady=5)
                    
                    # Afficher les variables à remplir
                    for var_name, var_info in to_fill_vars:
                        current_value = ""
                        if isinstance(var_info, dict):
                            current_value = var_info.get('current_value', '')
                        else:
                            current_value = str(var_info)
                        
                        # Créer un cadre pour cette variable
                        var_frame = ctk.CTkFrame(
                            variables_container,
                            fg_color=theme_colors['field_bg'],
                            corner_radius=8,
                            border_width=1,
                            border_color=theme_colors['field_border']
                        )
                        var_frame.pack(fill="x", pady=5, ipady=5, padx=10)
                        
                        # Conteneur pour le nom et la valeur
                        content_frame = ctk.CTkFrame(var_frame, fg_color="transparent")
                        content_frame.pack(fill="x", padx=10, pady=5)
                        content_frame.grid_columnconfigure(0, weight=2)
                        content_frame.grid_columnconfigure(1, weight=3)
                        
                        # Étiquette pour le nom de la variable
                        var_label = ctk.CTkLabel(
                            content_frame,
                            text=var_name,
                            font=("", 13, "bold"),
                            wraplength=250,
                            justify="left"
                        )
                        var_label.grid(row=0, column=0, sticky="w", padx=(5, 15), pady=5)
                        
                        # Champ de saisie pour la valeur
                        var_entry = ctk.CTkEntry(
                            content_frame,
                            width=300,
                            height=35,
                            placeholder_text="Entrez une valeur"
                        )
                        var_entry.grid(row=0, column=1, sticky="ew", padx=5, pady=5)
                        
                        # Mettre la valeur actuelle dans le champ
                        if current_value:
                            var_entry.insert(0, current_value)
                        
                        # Stocker l'entrée pour pouvoir récupérer la valeur plus tard
                        self.variable_entries[var_name] = var_entry
            
            # Boutons de navigation
            buttons_frame = ctk.CTkFrame(self.content_frame, fg_color="transparent")
            buttons_frame.pack(fill="x", padx=20, pady=(10, 20))
            
            # Bouton retour
            back_button = ctk.CTkButton(
                buttons_frame,
                text="Retour",
                width=120,
                height=40,
                fg_color=("gray70", "gray30"),
                hover_color=("gray60", "gray40"),
                command=lambda: self.show_step(self.current_step - 1)
            )
            back_button.pack(side="left", padx=10, pady=10)
            
            # Bouton suivant
            next_button = ctk.CTkButton(
                buttons_frame,
                text="Générer le document",
                width=180,
                height=40,
                font=("", 14, "bold"),
                command=self._generate_and_finalize_document
            )
            next_button.pack(side="right", padx=10, pady=10)
            
            # Mettre à jour l'étape actuelle
            self.current_step = 4
            
            logger.info("Affichage de l'étape de personnalisation")
            
        except Exception as e:
            logger.error(f"Erreur lors de l'affichage de l'étape de personnalisation: {e}")
            self.show_error(f"Erreur: {str(e)}")
            
    def _generate_and_finalize_document(self) -> None:
        """
        Récupère les variables saisies, génère le document et passe à l'étape de finalisation
        """
        try:
            # Récupérer les valeurs saisies si des entrées existent
            if hasattr(self, 'variable_entries') and self.variable_entries:
                for var_name, entry_widget in self.variable_entries.items():
                    # Récupérer la valeur saisie
                    value = entry_widget.get()
                    
                    # Mettre à jour la valeur dans les données du document
                    if hasattr(self, 'document_data') and self.document_data and 'variables' in self.document_data:
                        if var_name in self.document_data['variables']:
                            if isinstance(self.document_data['variables'][var_name], dict):
                                self.document_data['variables'][var_name]['current_value'] = value
                            else:
                                self.document_data['variables'][var_name] = value
            
            # Générer le document
            self._generate_document()
            
            # Passer à l'étape de finalisation
            self.show_step(5)
            
            logger.info("Document généré, passage à l'étape de finalisation")
            
        except Exception as e:
            logger.error(f"Erreur lors de la génération du document: {e}")
            self.show_error(f"Erreur: {str(e)}")
            
    def _generate_document(self) -> None:
        """
        Génère le document final avec les variables saisies
        """
        try:
            # Préparer le dossier de sortie
            output_dir = os.path.join("data", "documents", "outputs")
            os.makedirs(output_dir, exist_ok=True)
            
            # Créer un nom de fichier basé sur le template et le client
            template_name = self.selected_template.get('name', 'document')
            client_name = "client"
            if hasattr(self, 'client_info') and self.client_info:
                client_name = self.client_info.get('nom', '') or self.client_info.get('name', '') or self.client_info.get('société', '') or "client"
            
            # Générer un nom de fichier unique
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename_base = f"{template_name}_{client_name}_{timestamp}"
            
            # Nettoyer le nom du fichier (enlever les caractères spéciaux)
            filename_base = re.sub(r'[^\w\-_\. ]', '_', filename_base)
            
            # Chemins pour différents formats de documents
            pdf_path = os.path.join(output_dir, f"{filename_base}.pdf")
            txt_path = os.path.join(output_dir, f"{filename_base}.txt")
            docx_path = os.path.join(output_dir, f"{filename_base}.docx")
            
            # Obtenir le contenu du modèle
            template_content = self.selected_template.get('content', '')
            
            # Si le modèle a un chemin de fichier, l'utiliser pour obtenir le contenu
            if 'file_path' in self.selected_template and self.selected_template['file_path']:
                try:
                    template_file = self.selected_template['file_path']
                    if not os.path.isabs(template_file):
                        # Chemin relatif - chercher dans les dossiers standard
                        possible_paths = [
                            os.path.join("data", "templates", template_file),
                            os.path.join("data", "models", template_file),
                            template_file
                        ]
                        
                        for path in possible_paths:
                            if os.path.exists(path):
                                with open(path, 'r', encoding='utf-8') as f:
                                    template_content = f.read()
                                logger.info(f"Contenu du modèle chargé depuis le fichier: {path}")
                                break
                except Exception as e:
                    logger.warning(f"Impossible de charger le contenu du fichier modèle: {e}")
            
            # Remplacer les variables dans le modèle
            variables = {}
            
            # Récupérer les variables du document
            if hasattr(self, 'document_data') and self.document_data and 'variables' in self.document_data:
                for var_name, var_info in self.document_data.get('variables', {}).items():
                    if isinstance(var_info, dict):
                        variables[var_name] = var_info.get('current_value', '')
                    else:
                        variables[var_name] = str(var_info)
            
            # Ajouter les variables standard
            variables.update({
                "client_name": client_name,
                "date": datetime.datetime.now().strftime('%Y-%m-%d'),
                "datetime": datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "template_name": template_name
            })
            
            # Ajouter les informations du client
            if hasattr(self, 'client_info') and self.client_info:
                for key, value in self.client_info.items():
                    if value:  # Ne pas ajouter les valeurs vides
                        variables[f"client_{key}"] = value
            
            # Remplacer les variables dans le contenu
            final_content = template_content
            for var_name, var_value in variables.items():
                # Format {variable}
                final_content = final_content.replace(f"{{{var_name}}}", str(var_value))
                # Format {{variable}} (double accolades)
                final_content = final_content.replace(f"{{{{{var_name}}}}}", str(var_value))
            
            # Tenter d'abord de générer directement un PDF avec FPDF, plus fiable
            try:
                from fpdf import FPDF
                
                class PDF(FPDF):
                    def header(self):
                        pass
                        
                    def footer(self):
                        self.set_y(-15)
                        self.set_font('Arial', 'I', 8)
                        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')
                
                # Créer un document PDF
                pdf = PDF()
                pdf.add_page()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.set_font('Arial', '', 11)
                
                # Ajouter les lignes du document
                for line in final_content.split('\n'):
                    if line.strip():
                        if line.strip().startswith('#'):
                            # Titre
                            clean_line = line.strip().replace('#', '').strip()
                            pdf.set_font('Arial', 'B', 16)
                            pdf.cell(0, 10, txt=clean_line, ln=True)
                            pdf.set_font('Arial', '', 11)
                        else:
                            # Texte normal - utiliser multi_cell pour le retour à la ligne automatique
                            pdf.multi_cell(0, 5, txt=line)
                    else:
                        # Ligne vide
                        pdf.ln(5)
                
                # Sauvegarder le PDF
                pdf.output(pdf_path)
                
                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                    logger.info(f"Document PDF généré avec succès via FPDF: {pdf_path}")
                    self.generated_document_path = pdf_path
                    
                    # Créer également une version DOCX pour les utilisateurs qui préfèrent Word
                    # mais nous utilisons le PDF comme document principal
                    try:
                        from docx import Document
                        from docx.shared import Pt, Inches
                        
                        # Créer un nouveau document
                        doc = Document()
                        
                        # Définir les marges du document (en pouces)
                        sections = doc.sections
                        for section in sections:
                            section.top_margin = Inches(1)
                            section.bottom_margin = Inches(1)
                            section.left_margin = Inches(1)
                            section.right_margin = Inches(1)
                        
                        # Parcourir chaque ligne du contenu
                        for line in final_content.split('\n'):
                            if not line.strip():
                                # Ligne vide
                                doc.add_paragraph()
                            elif line.strip().startswith('#'):
                                # Titre
                                level = line.count('#', 0, 3)  # Compter les # au début (max 3)
                                clean_line = line.strip().replace('#', '').strip()
                                
                                if level == 1:
                                    # Titre principal
                                    heading = doc.add_heading(clean_line, level=1)
                                elif level == 2:
                                    # Sous-titre
                                    heading = doc.add_heading(clean_line, level=2)
                                else:
                                    # Autre niveau de titre
                                    heading = doc.add_paragraph(clean_line)
                                    heading.style = 'Heading3'
                            else:
                                # Texte normal
                                p = doc.add_paragraph(line)
                        
                        # Sauvegarder le document Word comme option alternative
                        doc.save(docx_path)
                        logger.info(f"Document Word créé avec succès: {docx_path}")
                    except Exception as docx_error:
                        logger.warning(f"Impossible de créer la version DOCX: {docx_error}")
                else:
                    raise Exception("Le fichier PDF généré est vide ou n'a pas été créé.")
                    
            except Exception as pdf_error:
                logger.error(f"Erreur lors de la génération du PDF: {pdf_error}")
                
                # En cas d'échec de FPDF, essayer la génération directe DOCX
                try:
                    from docx import Document
                    from docx.shared import Pt, Inches
                    
                    # Créer un nouveau document
                    doc = Document()
                    
                    # Définir les marges du document (en pouces)
                    sections = doc.sections
                    for section in sections:
                        section.top_margin = Inches(1)
                        section.bottom_margin = Inches(1)
                        section.left_margin = Inches(1)
                        section.right_margin = Inches(1)
                    
                    # Parcourir chaque ligne du contenu
                    for line in final_content.split('\n'):
                        if not line.strip():
                            # Ligne vide
                            doc.add_paragraph()
                        elif line.strip().startswith('#'):
                            # Titre
                            level = line.count('#', 0, 3)  # Compter les # au début (max 3)
                            clean_line = line.strip().replace('#', '').strip()
                            
                            if level == 1:
                                # Titre principal
                                heading = doc.add_heading(clean_line, level=1)
                            elif level == 2:
                                # Sous-titre
                                heading = doc.add_heading(clean_line, level=2)
                            else:
                                # Autre niveau de titre
                                heading = doc.add_paragraph(clean_line)
                                heading.style = 'Heading3'
                        else:
                            # Texte normal
                            p = doc.add_paragraph(line)
                    
                    # Sauvegarder le document Word
                    doc.save(docx_path)
                    logger.info(f"Document Word créé avec succès: {docx_path}")
                    self.generated_document_path = docx_path
                    
                except Exception as docx_error:
                    logger.error(f"Erreur lors de la création du document Word: {docx_error}")
                    
                    # En dernier recours, créer un fichier texte
                    try:
                        with open(txt_path, 'w', encoding='utf-8') as f:
                            f.write(final_content)
                        
                        logger.info(f"Document texte créé comme solution de repli: {txt_path}")
                        self.generated_document_path = txt_path
                    except Exception as txt_error:
                        logger.error(f"Erreur lors de la création du fichier texte: {txt_error}")
                        raise Exception(f"Impossible de générer le document dans aucun format.")
            
            # Enregistrer les informations du document dans la base de données
            document_info = {
                'id': str(uuid.uuid4()),
                'title': template_name,
                'client_id': self.client_info.get('id') if hasattr(self, 'client_info') and self.client_info else None,
                'client_name': client_name,
                'created_at': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'file_path': self.generated_document_path,
                'status': 'completed'
            }
            
            # Si le modèle a une méthode pour sauvegarder des documents, l'utiliser
            if hasattr(self.model, 'save_document'):
                self.model.save_document(document_info)
                logger.info(f"Document enregistré dans la base de données: {document_info['id']}")
                
        except Exception as e:
            logger.error(f"Erreur lors de la génération du document: {e}")
            raise e

    def show_finalization(self) -> None:
        """
        Affiche l'étape de finalisation du document
        """
        try:
            # Nettoyer la zone de contenu
            self._clear_content_frame()
                
            # Mettre à jour la barre de progression
            self.update_navigation()
            
            # Configuration du titre
            title_frame = ctk.CTkFrame(self.content_frame, fg_color="transparent")
            title_frame.pack(fill="x", padx=20, pady=(20, 10), anchor="n")
            
            title_label = ctk.CTkLabel(
                title_frame,
                text="Finalisation du document",
                font=("", 24, "bold")
            )
            title_label.pack(side="left", pady=10)
            
            # Définir des couleurs de thème par défaut si self.model.theme_colors n'existe pas
            theme_colors = {
                'frame_secondary_bg': ("gray90", "gray20"),
                'separator': ("gray75", "gray35"),
                'text_muted': ("gray50", "gray70"),
                'auto_detected_bg': ("#e6f7ff", "#1a3c4d"),
                'heading_bg': ("#f0f0f0", "#2d2d2d"),
                'field_bg': ("#ffffff", "#333333"),
                'field_border': ("#e0e0e0", "#555555"),
                'success_bg': ("#e6ffe6", "#1a4d1a")
            }
            
            # Utilisez les couleurs du modèle si disponibles
            if hasattr(self.model, 'theme_colors'):
                theme_colors = self.model.theme_colors
            
            # Cadre principal de finalisation
            finalization_frame = ctk.CTkFrame(self.content_frame)
            finalization_frame.pack(fill="both", expand=True, padx=20, pady=20)
            
            # Créer un en-tête explicatif avec icône
            header_frame = ctk.CTkFrame(
                finalization_frame,
                fg_color=theme_colors['success_bg'],
                corner_radius=10,
                height=80
            )
            header_frame.pack(fill="x", padx=20, pady=(20, 15))
            header_frame.pack_propagate(False)
            
            # Icône de succès
            success_icon = "✓"
            icon_label = ctk.CTkLabel(
                header_frame,
                text=success_icon,
                font=("", 32, "bold")
            )
            icon_label.pack(side="left", padx=(15, 0))
            
            # Cadre pour le texte
            header_text_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
            header_text_frame.pack(side="left", fill="both", expand=True, padx=15, pady=10)
            
            # Titre de succès
            header_title = ctk.CTkLabel(
                header_text_frame,
                text="Document généré avec succès !",
                font=("", 18, "bold"),
                anchor="w"
            )
            header_title.pack(anchor="w")
            
            # Description
            description_text = (
                "Votre document est prêt. Vous pouvez maintenant le télécharger, l'envoyer par email "
                "ou le consulter directement."
            )
            description_label = ctk.CTkLabel(
                header_text_frame,
                text=description_text,
                font=("", 12),
                wraplength=700,
                justify="left",
                anchor="w"
            )
            description_label.pack(anchor="w", fill="x", pady=(5, 0))
            
            # Cadre pour les options de finalisation
            options_frame = ctk.CTkFrame(
                finalization_frame,
                fg_color=theme_colors['field_bg'],
                corner_radius=10,
                border_width=1,
                border_color=theme_colors['field_border']
            )
            options_frame.pack(fill="x", padx=20, pady=(20, 10))
            
            # Titre des options
            options_title = ctk.CTkLabel(
                options_frame,
                text="Que souhaitez-vous faire avec votre document ?",
                font=("", 16, "bold")
            )
            options_title.pack(anchor="w", padx=20, pady=(15, 20))
            
            # Grille pour les boutons d'action
            buttons_grid = ctk.CTkFrame(options_frame, fg_color="transparent")
            buttons_grid.pack(fill="x", padx=20, pady=(0, 20))
            buttons_grid.columnconfigure(0, weight=1)
            buttons_grid.columnconfigure(1, weight=1)
            buttons_grid.columnconfigure(2, weight=1)
            
            # Fonction pour créer un bouton d'action avec icône
            def create_action_button(parent, icon, text, command, row, column):
                button_frame = ctk.CTkFrame(parent, fg_color="transparent")
                button_frame.grid(row=row, column=column, padx=10, pady=10, sticky="nsew")
                
                # Bouton principal
                button = ctk.CTkButton(
                    button_frame,
                    text="",
                    width=80,
                    height=80,
                    command=command,
                    corner_radius=10
                )
                button.pack(anchor="center", pady=(0, 10))
                
                # Icône sur le bouton
                icon_label = ctk.CTkLabel(
                    button,
                    text=icon,
                    font=("", 36)
                )
                icon_label.place(relx=0.5, rely=0.4, anchor="center")
                
                # Texte descriptif sous le bouton
                text_label = ctk.CTkLabel(
                    button_frame,
                    text=text,
                    font=("", 12),
                    wraplength=120,
                    justify="center"
                )
                text_label.pack(pady=(0, 5))
                
                return button
            
            # Bouton de téléchargement
            download_button = create_action_button(
                buttons_grid,
                "⬇️",
                "Télécharger",
                self._download_document,
                0, 0
            )
            
            # Bouton de visualisation
            view_button = create_action_button(
                buttons_grid,
                "👁️",
                "Visualiser",
                self._view_document,
                0, 1
            )
            
            # Bouton d'envoi par email
            email_button = create_action_button(
                buttons_grid,
                "✉️",
                "Envoyer par email",
                self._send_email,
                0, 2
            )
            
            # Afficher le chemin du fichier généré si disponible
            if hasattr(self, 'generated_document_path') and self.generated_document_path:
                file_frame = ctk.CTkFrame(
                    finalization_frame,
                    fg_color=theme_colors['field_bg'],
                    corner_radius=10,
                    border_width=1,
                    border_color=theme_colors['field_border']
                )
                file_frame.pack(fill="x", padx=20, pady=10)
                
                file_title = ctk.CTkLabel(
                    file_frame,
                    text="Informations sur le fichier",
                    font=("", 16, "bold")
                )
                file_title.pack(anchor="w", padx=20, pady=(15, 10))
                
                file_path_frame = ctk.CTkFrame(file_frame, fg_color="transparent")
                file_path_frame.pack(fill="x", padx=20, pady=(0, 15))
                
                # Chemin du fichier
                path_label = ctk.CTkLabel(
                    file_path_frame,
                    text="Emplacement :",
                    font=("", 12, "bold"),
                    anchor="w",
                    width=100
                )
                path_label.pack(side="left", padx=(0, 10))
                
                # Tronquer le chemin s'il est trop long
                display_path = self.generated_document_path
                if len(display_path) > 50:
                    parts = os.path.split(display_path)
                    display_path = os.path.join('...', parts[1])
                
                path_value = ctk.CTkLabel(
                    file_path_frame,
                    text=display_path,
                    font=("", 12),
                    anchor="w"
                )
                path_value.pack(side="left", fill="x", expand=True)
            
            # Boutons de navigation
            nav_buttons_frame = ctk.CTkFrame(self.content_frame, fg_color="transparent")
            nav_buttons_frame.pack(fill="x", padx=20, pady=(10, 20))
            
            # Bouton pour créer un nouveau document
            new_doc_button = ctk.CTkButton(
                nav_buttons_frame,
                text="Créer un nouveau document",
                width=200,
                height=40,
                font=("", 14, "bold"),
                command=self.reset_process
            )
            new_doc_button.pack(side="right", padx=10, pady=10)
            
            # Bouton pour revenir au tableau de bord
            dashboard_button = ctk.CTkButton(
                nav_buttons_frame,
                text="Retour au tableau de bord",
                width=200,
                height=40,
                fg_color=("gray70", "gray30"),
                hover_color=("gray60", "gray40"),
                command=lambda: self.parent.winfo_toplevel().show_view("dashboard")
            )
            dashboard_button.pack(side="left", padx=10, pady=10)
            
            # Mettre à jour l'étape actuelle
            self.current_step = 5
            
            logger.info("Affichage de l'étape de finalisation")
            
        except Exception as e:
            logger.error(f"Erreur lors de l'affichage de l'étape de finalisation: {e}")
            self.show_error(f"Erreur: {str(e)}")
            
    def _download_document(self) -> None:
        """
        Télécharge le document créé
        """
        try:
            # Vérifier que le document a été généré et existe
            if not hasattr(self, 'generated_document_path') or not hasattr(self, 'selected_template'):
                logger.error("Pas de document généré disponible pour le téléchargement")
                self.show_error("Le document n'a pas encore été généré ou n'est pas disponible pour le téléchargement.")
                return
                
            file_path = self.generated_document_path
            
            # Vérifier si le fichier existe
            if not os.path.exists(file_path):
                logger.error(f"Le fichier du document est introuvable: {file_path}")
                self.show_error("Le fichier du document est introuvable.")
                return
                
            # Vérifier si le fichier a une taille valide
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                logger.error(f"Le fichier du document est vide: {file_path}")
                self.show_error("Le fichier du document est vide ou corrompu.")
                return
                
            # Déterminer l'extension du fichier
            _, ext = os.path.splitext(file_path)
            
            # Déterminer le type de fichier et les options appropriées
            file_type_desc = "Document"
            if ext.lower() == '.pdf':
                file_type_desc = "Fichiers PDF"
            elif ext.lower() == '.txt':
                file_type_desc = "Fichiers texte"
            elif ext.lower() == '.html':
                file_type_desc = "Fichiers HTML"
            
            # Ouvrir une boîte de dialogue pour choisir l'emplacement de sauvegarde
            dest_path = filedialog.asksaveasfilename(
                title="Enregistrer le document",
                defaultextension=ext,
                initialfile=os.path.basename(file_path),
                filetypes=[(f"{file_type_desc}", f"*{ext}"), ("Tous les fichiers", "*.*")]
            )
            
            if not dest_path:
                logger.info("Téléchargement annulé par l'utilisateur")
                return
                
            # Copier le fichier
            try:
                shutil.copy2(file_path, dest_path)
                logger.info(f"Document téléchargé: {dest_path}")
                self.show_message("Succès", "Document téléchargé avec succès", "info")
            except Exception as e:
                logger.error(f"Erreur lors de la copie du fichier: {e}")
                
                # Essayer une approche alternative en cas d'échec
                try:
                    # Lire et écrire manuellement le fichier
                    with open(file_path, 'rb') as src_file, open(dest_path, 'wb') as dst_file:
                        dst_file.write(src_file.read())
                    
                    logger.info(f"Document téléchargé (méthode alternative): {dest_path}")
                    self.show_message("Succès", "Document téléchargé avec succès", "info")
                except Exception as e2:
                    logger.error(f"Erreur lors de la copie alternative du fichier: {e2}")
                    self.show_error(f"Erreur lors du téléchargement: {str(e2)}")
            
        except Exception as e:
            logger.error(f"Erreur lors du téléchargement du document: {e}")
            self.show_error(f"Erreur lors du téléchargement: {str(e)}")
            
    def _view_document(self) -> None:
        """
        Ouvre le document dans l'application par défaut
        """
        try:
            # Vérifier si le document a été généré
            if hasattr(self, 'generated_document_path') and self.generated_document_path:
                # Ouvrir le document avec l'application par défaut
                if os.path.exists(self.generated_document_path):
                    # Utiliser la commande appropriée selon l'OS
                    if platform.system() == 'Windows':
                        os.startfile(self.generated_document_path)
                    elif platform.system() == 'Darwin':  # macOS
                        subprocess.run(['open', self.generated_document_path])
                    else:  # Linux et autres
                        subprocess.run(['xdg-open', self.generated_document_path])
                    
                    logger.info(f"Document ouvert: {self.generated_document_path}")
                else:
                    raise Exception(f"Le fichier n'existe pas: {self.generated_document_path}")
            else:
                raise Exception("Aucun document n'a été généré")
                
        except Exception as e:
            logger.error(f"Erreur lors de l'ouverture du document: {e}")
            self.show_error(f"Erreur: {str(e)}")
            
    def _clear_content_frame(self) -> None:
        """
        Nettoie la zone de contenu en supprimant tous les widgets
        """
        try:
            if hasattr(self, 'content_frame'):
                for widget in self.content_frame.winfo_children():
                    widget.destroy()
        except Exception as e:
            logger.error(f"Erreur lors du nettoyage de la zone de contenu: {e}")
            
    def _send_email(self) -> None:
        """
        Envoie le document par email
        """
        try:
            # Implémenter la logique d'envoi par email
            logger.info("Envoi du document par email")
            self.show_error("Cette fonctionnalité n'est pas encore implémentée")
            
        except Exception as e:
            logger.error(f"Erreur lors de l'envoi du document par email: {e}")
            self.show_error(f"Erreur: {str(e)}")
    
    def show_message(self, title: str, message: str, message_type: str = "info") -> None:
        """
        Affiche un message dans une boîte de dialogue
        
        Args:
            title: Titre de la boîte de dialogue
            message: Message à afficher
            message_type: Type de message (info, error, warning)
        """
        try:
            # Utiliser DialogUtils si disponible
            if 'DialogUtils' in globals():
                DialogUtils.show_message(self.parent, title, message, message_type)
            else:
                # Fallback avec messagebox standard
                from tkinter import messagebox
                
                if message_type == "error":
                    messagebox.showerror(title, message)
                elif message_type == "warning":
                    messagebox.showwarning(title, message)
                else:
                    messagebox.showinfo(title, message)
        except Exception as e:
            logger.error(f"Erreur lors de l'affichage du message: {e}")
            # Dernier recours
            from tkinter import messagebox
            messagebox.showerror("Erreur", message)

    def _reset_client_selection(self) -> None:
        """
        Réinitialise la sélection du client et affiche à nouveau le formulaire de recherche
        """
        try:
            # Réinitialiser la sélection
            self.client_selected = False
            self.client_info = {}
            
            # Afficher à nouveau la recherche
            self.show_client_form()
            
        except Exception as e:
            logger.error(f"Erreur lors de la réinitialisation de la sélection du client: {e}")
            self.show_error(f"Erreur: {str(e)}")
            
    def show(self) -> None:
        """
        Affiche la vue
        """
        # S'assurer que le conteneur principal est visible
        self.frame.pack(fill="both", expand=True)
        # Réinitialiser l'état si nécessaire
        self.current_step = 0
        # Afficher l'étape initiale
        self.show_step(0)
        logger.info("Vue document_creator affichée")

    def hide(self) -> None:
        """
        Cache la vue
        """
        # Masquer le conteneur principal
        self.frame.pack_forget()
        logger.info("Vue document_creator masquée")
