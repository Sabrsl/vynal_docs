import os
import platform
import subprocess
import logging
import customtkinter as ctk
from PIL import Image, ImageTk
import fitz  # PyMuPDF pour les PDFs
import docx  # python-docx pour les documents Word
import chardet  # Pour la détection d'encodage
from typing import Dict, Optional, Any

logger = logging.getLogger("VynalDocsAutomator.DocumentPreview")

class DocumentPreview:
    """Classe pour gérer la prévisualisation des documents"""
    
    def __init__(self, parent):
        """
        Initialise le prévisualiseur de documents
        
        Args:
            parent: Widget parent
        """
        self.parent = parent
        self.preview_window: Optional[ctk.CTkToplevel] = None
        self.current_document: Optional[Dict[str, Any]] = None
        self.current_page = 0
        self.doc_object = None
        self.page_label = None
        self.content_frame = None
        self.photo_references = []  # Garder une référence aux images pour éviter le garbage collection
        
    def preview(self, document: Dict[str, Any]):
        """
        Prévisualise un document
        
        Args:
            document: Dictionnaire contenant les informations du document
        """
        try:
            # Nettoyer les prévisualisations précédentes
            if self.preview_window and self.preview_window.winfo_exists():
                self.preview_window.destroy()
            
            # Vérifier si le document existe
            if not document:
                self._show_error("Document non trouvé")
                return
            
            # Vérifier si le fichier existe
            file_path = document.get("file_path", "")
            if not file_path or not os.path.exists(file_path):
                self._show_error(f"Fichier non trouvé: {file_path}")
                return
            
            # Créer la fenêtre de prévisualisation
            self.preview_window = ctk.CTkToplevel(self.parent)
            self.preview_window.title(f"Prévisualisation - {document.get('title', 'Document')}")
            
            # Configurer la fenêtre en plein écran
            screen_width = self.preview_window.winfo_screenwidth()
            screen_height = self.preview_window.winfo_screenheight()
            
            if platform.system() == 'Windows':
                # Pour Windows, utiliser state('zoomed')
                self.preview_window.state('zoomed')
            elif platform.system() == 'Darwin':  # macOS
                # Pour macOS, utiliser fullscreen
                self.preview_window.attributes('-fullscreen', True)
            else:  # Linux
                # Pour Linux, utiliser la géométrie complète
                self.preview_window.geometry(f"{screen_width}x{screen_height}+0+0")
            
            # Forcer la fenêtre au premier plan
            self.preview_window.attributes('-topmost', True)
            self.preview_window.lift()
            self.preview_window.focus_force()
            
            # Configuration de la grille principale
            self.preview_window.grid_columnconfigure(0, weight=1)
            self.preview_window.grid_rowconfigure(1, weight=1)
            
            # Ajouter une barre d'outils
            toolbar = ctk.CTkFrame(self.preview_window)
            toolbar.grid(row=0, column=0, sticky="ew", padx=10, pady=(10, 5))
            toolbar.grid_columnconfigure(4, weight=1)  # Pour pousser les boutons de navigation à droite
            
            # Bouton pour ouvrir avec l'application par défaut
            open_btn = ctk.CTkButton(
                toolbar,
                text="Ouvrir avec l'application par défaut",
                command=lambda: self._open_with_default_app(file_path),
                height=32
            )
            open_btn.grid(row=0, column=0, padx=(5, 10), pady=5)
            
            # Déterminer le type de fichier
            file_type = file_path.lower().split('.')[-1]
            
            # Zone de contenu avec défilement
            self.content_frame = ctk.CTkScrollableFrame(self.preview_window)
            self.content_frame.grid(row=1, column=0, sticky="nsew", padx=10, pady=10)
            self.content_frame.grid_columnconfigure(0, weight=1)
            
            # Barre de statut
            status_bar = ctk.CTkFrame(self.preview_window, height=30)
            status_bar.grid(row=2, column=0, sticky="ew", padx=10, pady=(0, 10))
            status_bar.grid_columnconfigure(0, weight=1)
            
            file_info = ctk.CTkLabel(
                status_bar,
                text=f"Type: {file_type.upper()} | Taille: {self._format_file_size(os.path.getsize(file_path))}",
                font=ctk.CTkFont(size=12)
            )
            file_info.grid(row=0, column=0, sticky="w", padx=10, pady=5)
            
            # Prévisualiser selon le type de fichier
            if file_type == 'pdf':
                self._preview_pdf(file_path, toolbar)
            elif file_type in ['doc', 'docx']:
                self._preview_word(file_path)
            elif file_type == 'txt':
                self._preview_text(file_path)
            else:
                info_label = ctk.CTkLabel(
                    self.content_frame,
                    text=f"Type de fichier {file_type} non pris en charge pour la prévisualisation.\nUtilisez le bouton 'Ouvrir avec l'application par défaut'.",
                    font=ctk.CTkFont(size=14),
                    wraplength=500
                )
                info_label.pack(pady=50)
            
            # Stocker le document courant
            self.current_document = document
            
            # Forcer la fenêtre au premier plan et maintenir le focus
            self._force_window_to_front()
            
            # Démarrer un timer pour maintenir la fenêtre au premier plan
            self._maintain_window_focus()
            
        except Exception as e:
            logger.error(f"Erreur lors de la prévisualisation du document: {e}", exc_info=True)
            self._show_error(f"Erreur lors de la prévisualisation: {str(e)}")
    
    def _preview_pdf(self, file_path: str, toolbar: ctk.CTkFrame):
        """
        Prévisualise un fichier PDF avec navigation entre pages
        
        Args:
            file_path: Chemin vers le fichier PDF
            toolbar: Barre d'outils pour ajouter les boutons de navigation
        """
        try:
            # Vérifier que le fichier est bien un PDF valide
            try:
                # Vérifier la taille du fichier
                file_size = os.path.getsize(file_path)
                if file_size == 0:
                    raise ValueError("Le fichier PDF est vide")
                
                # Lire les premiers octets pour vérifier la signature PDF
                with open(file_path, 'rb') as f:
                    pdf_header = f.read(5)
                    if pdf_header != b'%PDF-':
                        raise ValueError("Le fichier n'est pas un PDF valide (signature incorrecte)")
                
                # Ouvrir le PDF avec PyMuPDF
                self.doc_object = fitz.open(file_path)
                
                # Vérifier qu'il y a au moins une page
                if len(self.doc_object) < 1:
                    raise ValueError("Le PDF ne contient aucune page")
                
            except Exception as open_error:
                logger.error(f"Impossible d'ouvrir le PDF {file_path}: {open_error}")
                error_frame = ctk.CTkFrame(self.content_frame)
                error_frame.pack(fill="both", expand=True, padx=20, pady=20)
                
                error_label = ctk.CTkLabel(
                    error_frame,
                    text=f"Impossible d'ouvrir le PDF : {str(open_error)}",
                    font=ctk.CTkFont(size=14),
                    wraplength=600,
                    text_color="red"
                )
                error_label.pack(pady=20)
                
                # Proposer d'afficher le contenu en tant que texte
                try:
                    # Essayer de lire les premiers Ko du fichier
                    with open(file_path, 'rb') as f:
                        content = f.read(4096)
                    
                    if b'%PDF-' in content:
                        suggestion = "Le fichier semble être un PDF mais la bibliothèque ne peut pas l'ouvrir."
                    else:
                        suggestion = "Le fichier ne semble pas être un PDF valide."
                    
                    suggestion_label = ctk.CTkLabel(
                        error_frame,
                        text=f"{suggestion}\nEssayez d'ouvrir le fichier avec l'application par défaut.",
                        font=ctk.CTkFont(size=13),
                        wraplength=600
                    )
                    suggestion_label.pack(pady=10)
                except:
                    pass
                
                return  # Arrêter le traitement
            
            # Continuer avec le traitement normal...
            page_count = len(self.doc_object)
            self.current_page = 0
            
            # Ajouter les boutons de navigation
            prev_btn = ctk.CTkButton(
                toolbar,
                text="< Précédent",
                command=self._prev_page,
                width=100,
                height=32,
                state="disabled" if page_count <= 1 else "normal"
            )
            prev_btn.grid(row=0, column=1, padx=5, pady=5)
            
            self.page_label = ctk.CTkLabel(
                toolbar,
                text=f"Page 1 sur {page_count}",
                font=ctk.CTkFont(size=13)
            )
            self.page_label.grid(row=0, column=2, padx=10, pady=5)
            
            next_btn = ctk.CTkButton(
                toolbar,
                text="Suivant >",
                command=self._next_page,
                width=100,
                height=32,
                state="disabled" if page_count <= 1 else "normal"
            )
            next_btn.grid(row=0, column=3, padx=5, pady=5)
            
            # Bouton de zoom
            zoom_label = ctk.CTkLabel(toolbar, text="Zoom:", font=ctk.CTkFont(size=13))
            zoom_label.grid(row=0, column=4, padx=(20, 5), pady=5, sticky="e")
            
            zoom_values = ["50%", "75%", "100%", "125%", "150%", "200%"]
            self.zoom_var = ctk.StringVar(value="100%")
            
            # Créer un frame pour le menu de zoom
            zoom_frame = ctk.CTkFrame(toolbar, fg_color="transparent")
            zoom_frame.grid(row=0, column=5, padx=(0, 10), pady=5, sticky="e")
            
            # Créer le menu de zoom avec un style personnalisé
            zoom_menu = ctk.CTkOptionMenu(
                zoom_frame,
                values=zoom_values,
                variable=self.zoom_var,
                command=self._update_pdf_view,
                width=80,
                height=32,
                fg_color=("gray75", "gray25"),  # Couleur de fond adaptative
                button_color=("gray70", "gray30"),  # Couleur du bouton adaptative
                button_hover_color=("gray65", "gray35")  # Couleur au survol adaptative
            )
            zoom_menu.pack(side="right")
            
            # Afficher la première page
            self._update_pdf_view()
            
        except Exception as e:
            logger.error(f"Erreur lors de la prévisualisation du PDF: {e}", exc_info=True)
            self._show_error(f"Erreur lors de la prévisualisation du PDF: {str(e)}")
    
    def _update_pdf_view(self, *args):
        """Met à jour l'affichage du PDF avec le zoom actuel"""
        try:
            # Nettoyer le contenu précédent
            for widget in self.content_frame.winfo_children():
                widget.destroy()
            
            # Mettre à jour le label de page
            if self.page_label:
                page_count = len(self.doc_object)
                self.page_label.configure(text=f"Page {self.current_page + 1} sur {page_count}")
            
            # Obtenir la page actuelle
            page = self.doc_object[self.current_page]
            
            # Déterminer le facteur de zoom
            zoom_text = self.zoom_var.get()
            zoom_factor = float(zoom_text.rstrip('%')) / 100.0
            
            # Créer une matrice de zoom avec une meilleure qualité
            zoom_matrix = fitz.Matrix(zoom_factor, zoom_factor)
            
            # Rendre la page avec le zoom et une meilleure qualité
            pix = page.get_pixmap(matrix=zoom_matrix, alpha=False)
            
            # Convertir en format PhotoImage avec une meilleure gestion de la mémoire
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # Calculer les dimensions de la fenêtre de prévisualisation
            preview_width = self.content_frame.winfo_width()
            preview_height = self.content_frame.winfo_height()
            
            # Si la fenêtre n'a pas encore de dimensions, utiliser les dimensions de l'écran
            if preview_width <= 1 or preview_height <= 1:
                preview_width = self.preview_window.winfo_screenwidth() - 40
                preview_height = self.preview_window.winfo_screenheight() - 200
            
            # Calculer le ratio de redimensionnement pour s'adapter à la fenêtre
            width_ratio = preview_width / img.width
            height_ratio = preview_height / img.height
            resize_ratio = min(width_ratio, height_ratio)
            
            # Redimensionner l'image pour s'adapter à la fenêtre
            new_width = int(img.width * resize_ratio)
            new_height = int(img.height * resize_ratio)
            img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
            
            # Convertir en PhotoImage
            photo = ImageTk.PhotoImage(img)
            
            # Garder une référence à l'image
            self.photo_references.clear()
            self.photo_references.append(photo)
            
            # Créer un label pour afficher l'image avec un fond blanc
            label = ctk.CTkLabel(
                self.content_frame,
                image=photo,
                text="",
                fg_color="white"  # Fond blanc pour une meilleure visibilité
            )
            label.photo = photo  # Garder une référence supplémentaire
            label.pack(pady=5)
            
            # Forcer la mise à jour de l'affichage
            self.content_frame.update_idletasks()
            
        except Exception as e:
            logger.error(f"Erreur lors de la mise à jour de la vue PDF: {e}", exc_info=True)
            ctk.CTkLabel(
                self.content_frame,
                text=f"Erreur d'affichage: {str(e)}",
                text_color="red",
                wraplength=500
            ).pack(pady=20)
    
    def _next_page(self):
        """Affiche la page suivante du PDF"""
        if self.doc_object and self.current_page < len(self.doc_object) - 1:
            self.current_page += 1
            self._update_pdf_view()
    
    def _prev_page(self):
        """Affiche la page précédente du PDF"""
        if self.doc_object and self.current_page > 0:
            self.current_page -= 1
            self._update_pdf_view()
    
    def _preview_word(self, file_path: str):
        """
        Prévisualise un document Word
        
        Args:
            file_path: Chemin vers le fichier Word
        """
        try:
            # Ouvrir le document Word
            doc = docx.Document(file_path)
            
            # Créer un widget de texte
            text_widget = ctk.CTkTextbox(self.content_frame, wrap="word", font=ctk.CTkFont(family="Segoe UI", size=13))
            text_widget.pack(fill=ctk.BOTH, expand=True, padx=5, pady=5)
            
            # Extraire et afficher le texte
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    text_widget.insert("end", para.text + "\n", ('heading',))
                    text_widget.tag_configure('heading', font=ctk.CTkFont(family="Segoe UI", size=16, weight="bold"))
                else:
                    text_widget.insert("end", para.text + "\n")
            
            # Ajouter les tables
            for table in doc.tables:
                text_widget.insert("end", "\n--- TABLE ---\n", ('table_header',))
                text_widget.tag_configure('table_header', font=ctk.CTkFont(size=12, weight="bold"))
                
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text_widget.insert("end", row_text + "\n")
                
                text_widget.insert("end", "--- FIN TABLE ---\n\n")
            
            text_widget.configure(state="disabled")
            
        except Exception as e:
            logger.error(f"Erreur lors de la prévisualisation du document Word: {e}", exc_info=True)
            self._show_error("Erreur lors de la prévisualisation du document Word")
    
    def _preview_text(self, file_path: str):
        """
        Prévisualise un fichier texte
        
        Args:
            file_path: Chemin vers le fichier texte
        """
        try:
            # Détecter l'encodage du fichier
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'utf-8'  # Fallback à utf-8 si None
            
            # Créer un widget de texte
            text_widget = ctk.CTkTextbox(
                self.content_frame,
                wrap="word",
                font=ctk.CTkFont(family="Consolas", size=13)
            )
            text_widget.pack(fill=ctk.BOTH, expand=True, padx=5, pady=5)
            
            # Lire et afficher le contenu
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    text_widget.insert("1.0", content)
            except UnicodeDecodeError:
                # Tentative de secours avec utf-8 et ignorer les erreurs
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                    text_widget.insert("1.0", content)
                    text_widget.insert("1.0", "[AVERTISSEMENT: Problèmes d'encodage détectés. L'affichage peut être incorrect.]\n\n", "warning")
                    text_widget.tag_configure("warning", foreground="red")
            
            text_widget.configure(state="disabled")
            
        except Exception as e:
            logger.error(f"Erreur lors de la prévisualisation du fichier texte: {e}", exc_info=True)
            self._show_error("Erreur lors de la prévisualisation du fichier texte")
    
    def _open_with_default_app(self, file_path: str):
        """
        Ouvre le fichier avec l'application par défaut du système
        
        Args:
            file_path: Chemin vers le fichier à ouvrir
        """
        try:
            logger.info(f"Ouverture du fichier {file_path} avec l'application par défaut")
            
            # Vérifier si le fichier existe
            if not os.path.exists(file_path):
                self._show_error(f"Le fichier n'existe pas: {file_path}")
                return
            
            # Vérifier si c'est un fichier (pas un dossier)
            if not os.path.isfile(file_path):
                self._show_error(f"Le chemin spécifié n'est pas un fichier: {file_path}")
                return
            
            # Vérifier les permissions
            if not os.access(file_path, os.R_OK):
                self._show_error(f"Impossible d'accéder au fichier. Vérifiez les permissions: {file_path}")
                return
            
            # Convertir le chemin en chemin absolu
            abs_path = os.path.abspath(file_path)
            
            # Ouvrir le fichier selon le système d'exploitation
            if platform.system() == 'Windows':
                try:
                    # Essayer d'abord avec os.startfile
                    os.startfile(abs_path)
                except Exception as e:
                    try:
                        # En cas d'échec, essayer avec subprocess et explorer
                        subprocess.Popen(['explorer', abs_path], shell=True)
                    except Exception as e2:
                        # Dernière tentative avec cmd
                        subprocess.Popen(['cmd', '/c', 'start', '', abs_path], shell=True)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.Popen(['open', abs_path])
            else:  # Linux
                subprocess.Popen(['xdg-open', abs_path])
                
        except Exception as e:
            logger.error(f"Erreur lors de l'ouverture du fichier: {e}", exc_info=True)
            self._show_error(f"Erreur lors de l'ouverture du fichier: {str(e)}")
    
    def _show_error(self, message: str):
        """
        Affiche un message d'erreur dans une fenêtre modale
        
        Args:
            message: Message d'erreur à afficher
        """
        # Fermer la fenêtre de prévisualisation si elle existe
        if self.preview_window and self.preview_window.winfo_exists():
            self.preview_window.destroy()
            self.preview_window = None
        
        # Créer une fenêtre d'erreur
        error_window = ctk.CTkToplevel(self.parent)
        error_window.title("Erreur")
        error_window.geometry("400x200")
        error_window.resizable(False, False)
        
        # Configuration de la grille
        error_window.grid_columnconfigure(0, weight=1)
        error_window.grid_rowconfigure(0, weight=1)
        
        # Cadre principal
        main_frame = ctk.CTkFrame(error_window)
        main_frame.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        main_frame.grid_columnconfigure(0, weight=1)
        main_frame.grid_rowconfigure(0, weight=1)
        
        # Message d'erreur
        error_label = ctk.CTkLabel(
            main_frame,
            text=message,
            font=ctk.CTkFont(size=14),
            wraplength=350,
            text_color=("red", "#FF5555")
        )
        error_label.grid(row=0, column=0, padx=20, pady=(20, 10), sticky="nsew")
        
        # Bouton OK
        ok_button = ctk.CTkButton(
            main_frame,
            text="OK",
            command=error_window.destroy,
            width=100,
            height=35
        )
        ok_button.grid(row=1, column=0, pady=(5, 15))
        
        # Centrer la fenêtre
        self._center_window(error_window)
        
        # Rendre modale
        error_window.transient(self.parent)
        error_window.grab_set()
        self.parent.wait_window(error_window)
    
    def _center_window(self, window: ctk.CTkToplevel):
        """
        Centre une fenêtre sur l'écran
        
        Args:
            window: Fenêtre à centrer
        """
        window.update_idletasks()
        width = window.winfo_width()
        height = window.winfo_height()
        x = (window.winfo_screenwidth() // 2) - (width // 2)
        y = (window.winfo_screenheight() // 2) - (height // 2)
        window.geometry(f"{width}x{height}+{x}+{y}")
    
    def _format_file_size(self, size_bytes: int) -> str:
        """
        Formate une taille en octets en format lisible (KB, MB, etc.)
        
        Args:
            size_bytes: Taille en octets
            
        Returns:
            Taille formatée
        """
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.2f} {unit}"
            size_bytes /= 1024.0
            
        return f"{size_bytes:.2f} PB"
    
    def __del__(self):
        """Nettoyage lors de la destruction de l'objet"""
        try:
            # Fermer les objets de document ouverts
            if self.doc_object:
                self.doc_object.close()
                
            # Fermer la fenêtre si elle existe encore
            if self.preview_window and self.preview_window.winfo_exists():
                self.preview_window.destroy()
                
        except Exception:
            pass  # Ignorer les erreurs pendant le nettoyage
    
    def _force_window_to_front(self):
        """Force la fenêtre au premier plan"""
        if self.preview_window and self.preview_window.winfo_exists():
            self.preview_window.attributes('-topmost', True)
            self.preview_window.lift()
            self.preview_window.focus_force()
            self.preview_window.update()
    
    def _maintain_window_focus(self):
        """Maintient la fenêtre au premier plan"""
        if self.preview_window and self.preview_window.winfo_exists():
            self._force_window_to_front()
            # Répéter toutes les 100ms
            self.preview_window.after(100, self._maintain_window_focus)