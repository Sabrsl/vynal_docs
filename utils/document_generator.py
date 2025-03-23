import os
import re
import logging
from datetime import datetime
import io

# Pour les documents PDF
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, mm, inch
    from reportlab.platypus import Paragraph, Table, TableStyle, Spacer, Image
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from PIL import Image as PILImage
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logging.warning("ReportLab n'est pas installé. La génération de PDF sera limitée.")

# Pour les documents Word
try:
    import docx
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx n'est pas installé. La génération de DOCX sera limitée.")

logger = logging.getLogger("VynalDocsAutomator.DocumentGenerator")

class DocumentGenerator:
    """
    Générateur de documents pour l'application
    Version améliorée avec gestion d'erreur robuste
    """
    
    def __init__(self, config_manager=None):
        """
        Initialise le générateur de documents
        
        Args:
            config_manager: Gestionnaire de configuration
        """
        self.config = config_manager
        logger.info("DocumentGenerator initialisé avec les capacités suivantes:")
        logger.info(f"- PDF (ReportLab): {'Disponible' if REPORTLAB_AVAILABLE else 'Non disponible'}")
        logger.info(f"- DOCX (python-docx): {'Disponible' if DOCX_AVAILABLE else 'Non disponible'}")
    
    def strip_html(self, html_content):
        """
        Enlève les balises HTML du contenu tout en préservant le texte
        et extrait les images pour le traitement ultérieur
        
        Args:
            html_content: Contenu HTML
            
        Returns:
            tuple: (texte sans balises HTML, liste d'images extraites)
        """
        # Images extraites
        images = []
        
        # Si le contenu n'est pas une chaîne, le convertir
        if not isinstance(html_content, str):
            logger.warning(f"Le contenu n'est pas une chaîne, conversion en: {str(html_content)}")
            return str(html_content), images
        
        # Journaliser le contenu HTML pour le débogage
        logger.info(f"Contenu HTML à traiter (début): {html_content[:100]}...")
        
        # Extraire les balises d'image pour traitement ultérieur
        img_pattern = r'<img\s+[^>]*src\s*=\s*["\']([^"\']+)["\'][^>]*>'
        img_matches = list(re.finditer(img_pattern, html_content))
        logger.info(f"Nombre d'images trouvées dans le HTML: {len(img_matches)}")
        
        for img_match in img_matches:
            img_src = img_match.group(1)
            # Journaliser pour le débogage
            logger.info(f"Image trouvée avec src: {img_src[:50]}...")
            
            # Extraire les dimensions si disponibles
            width_match = re.search(r'width\s*=\s*["\']([^"\']+)["\']', img_match.group(0))
            height_match = re.search(r'height\s*=\s*["\']([^"\']+)["\']', img_match.group(0))
            width = width_match.group(1) if width_match else None
            height = height_match.group(1) if height_match else None
            
            # Si c'est une image en base64, la sauvegarder dans un fichier temporaire
            if img_src.startswith('data:image/'):
                try:
                    import base64
                    import tempfile
                    
                    # Extraire le type et les données
                    header, data = img_src.split(',', 1)
                    mime_type = header.split(';')[0].split(':')[1]
                    img_ext = mime_type.split('/')[1]
                    
                    # Créer un fichier temporaire
                    temp_fd, temp_img_path = tempfile.mkstemp(suffix=f'.{img_ext}')
                    os.close(temp_fd)
                    
                    # Décoder et sauvegarder l'image
                    img_data = base64.b64decode(data)
                    with open(temp_img_path, 'wb') as f:
                        f.write(img_data)
                    
                    # Utiliser ce chemin temporaire au lieu de l'URL base64
                    img_src = temp_img_path
                    logger.info(f"Image base64 convertie en fichier temporaire: {temp_img_path}")
                except Exception as e:
                    logger.warning(f"Erreur lors du traitement de l'image base64: {e}")
            
            # Ajouter l'image à la liste avec sa position et dimensions
            image_info = {
                'src': img_src,
                'position': img_match.start()
            }
            if width:
                image_info['width'] = width
                logger.info(f"Largeur de l'image: {width}")
            if height:
                image_info['height'] = height
                logger.info(f"Hauteur de l'image: {height}")
                
            images.append(image_info)
        
        # Supprimer les balises d'image mais ajouter un marqueur pour leur emplacement
        html_content = re.sub(img_pattern, '[IMAGE]', html_content)
        
        # Compter les occurrences de [IMAGE] pour vérifier la cohérence
        image_markers = html_content.count('[IMAGE]')
        logger.info(f"Nombre de marqueurs [IMAGE] dans le texte: {image_markers}")
        if image_markers != len(images):
            logger.warning(f"Incohérence: {len(images)} images extraites mais {image_markers} marqueurs [IMAGE]")
        
        # Traiter certaines balises spéciales pour conserver la mise en forme
        # Remplacer les listes à puces par des caractères spéciaux
        html_content = re.sub(r'<li>(.*?)</li>', r'• \1\n', html_content)
        
        # Remplacer les sauts de ligne HTML par des sauts de ligne réels
        html_content = html_content.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
        
        # Remplacer les paragraphes par des lignes avec double saut de ligne
        html_content = re.sub(r'<p>(.*?)</p>', r'\1\n\n', html_content)
        
        # Remplacer les titres avec mise en forme
        html_content = re.sub(r'<h1>(.*?)</h1>', r'\1\n\n', html_content)
        html_content = re.sub(r'<h2>(.*?)</h2>', r'\1\n\n', html_content)
        html_content = re.sub(r'<h3>(.*?)</h3>', r'\1\n\n', html_content)
        
        # Remplacer les balises de style courantes
        html_content = re.sub(r'<strong>(.*?)</strong>', r'\1', html_content)
        html_content = re.sub(r'<b>(.*?)</b>', r'\1', html_content)
        html_content = re.sub(r'<em>(.*?)</em>', r'\1', html_content)
        html_content = re.sub(r'<i>(.*?)</i>', r'\1', html_content)
        html_content = re.sub(r'<u>(.*?)</u>', r'\1', html_content)
        
        # Supprimer les balises de listes
        html_content = html_content.replace('<ul>', '').replace('</ul>', '\n')
        html_content = html_content.replace('<ol>', '').replace('</ol>', '\n')
        
        # Supprimer toutes les autres balises HTML
        html_content = re.sub(r'<[^>]*>', '', html_content)
        
        # Supprimer les espaces multiples
        html_content = re.sub(r' +', ' ', html_content)
        
        # Supprimer les sauts de ligne multiples
        html_content = re.sub(r'\n{3,}', '\n\n', html_content)
        
        # Gérer les entités HTML courantes et avancées
        html_entities = {
            '&amp;': '&', 
            '&lt;': '<', 
            '&gt;': '>', 
            '&quot;': '"', 
            '&apos;': "'",
            '&nbsp;': ' ',
            '&copy;': '©',
            '&reg;': '®',
            '&trade;': '™',
            '&euro;': '€',
            '&pound;': '£',
            '&yen;': '¥',
            '&cent;': '¢',
            '&deg;': '°',
            '&plusmn;': '±',
            '&divide;': '÷',
            '&times;': '×',
            '&sect;': '§',
            '&para;': '¶',
            '&micro;': 'µ',
            '&middot;': '·',
            '&bull;': '•',
            '&hellip;': '…',
            '&ndash;': '–',
            '&mdash;': '—',
            '&lsquo;': ''',
            '&rsquo;': ''',
            '&ldquo;': '"',
            '&rdquo;': '"',
            '&laquo;': '«',
            '&raquo;': '»',
            '&frac14;': '¼',
            '&frac12;': '½',
            '&frac34;': '¾',
            '&larr;': '←',
            '&uarr;': '↑',
            '&rarr;': '→',
            '&darr;': '↓',
            '&infin;': '∞',
            '&ne;': '≠',
            '&asymp;': '≈',
            '&le;': '≤',
            '&ge;': '≥',
            '&sum;': '∑',
            '&int;': '∫',
            '&alpha;': 'α',
            '&beta;': 'β',
            '&gamma;': 'γ',
            '&delta;': 'δ',
            '&epsilon;': 'ε',
            '&theta;': 'θ',
            '&lambda;': 'λ',
            '&mu;': 'μ',
            '&pi;': 'π',
            '&sigma;': 'σ',
            '&phi;': 'φ',
            '&omega;': 'ω'
        }
        
        for entity, char in html_entities.items():
            html_content = html_content.replace(entity, char)
        
        # Traiter les entités HTML numériques (décimales et hexadécimales)
        def replace_numeric_entity(match):
            try:
                value = match.group(1)
                if value.startswith('x'):
                    # Entité hexadécimale
                    code = int(value[1:], 16)
                else:
                    # Entité décimale
                    code = int(value)
                return chr(code)
            except:
                return match.group(0)  # Retourner l'entité originale en cas d'erreur
        
        # Remplacer les entités numériques
        html_content = re.sub(r'&#(\d+);', lambda m: chr(int(m.group(1))), html_content)
        html_content = re.sub(r'&#x([0-9a-fA-F]+);', lambda m: chr(int(m.group(1), 16)), html_content)
        
        return html_content.strip(), images
    
    def replace_variables(self, content, variables):
        """
        Remplace les variables dans un contenu de manière sécurisée
        Supporte à la fois les formats {variable} et {{variable}}
        
        Args:
            content: Contenu avec variables
            variables: Dictionnaire des variables et leurs valeurs
            
        Returns:
            str: Contenu avec variables remplacées
        """
        import re
        
        # Si le contenu n'est pas une chaîne, le convertir
        if not isinstance(content, str):
            content = str(content)
        
        # Préparation des valeurs pour éviter les valeurs None
        safe_variables = {}
        for key, value in variables.items():
            if value is None:
                safe_variables[key] = ""
            else:
                safe_variables[key] = str(value)
        
        # Remplacer d'abord les variables au format {{variable}}
        for var_name, var_value in safe_variables.items():
            # Format {{variable}}
            pattern = r"{{" + re.escape(var_name) + r"}}"
            content = re.sub(pattern, var_value, content)
        
        # Ensuite remplacer les variables au format {variable}
        for var_name, var_value in safe_variables.items():
            # Format {variable} (pour la rétrocompatibilité)
            # Utiliser une expression régulière pour éviter les faux positifs
            pattern = r"{" + re.escape(var_name) + r"}"
            content = re.sub(pattern, var_value, content)
        
        # Rechercher les variables non remplacées au format {{variable}}
        remaining_vars = re.findall(r'{{([^{}]*?)}}', content)
        if remaining_vars:
            logger.warning(f"Variables non remplacées (format {{variable}}): {remaining_vars}")
            # Remplacer les variables non trouvées par une chaîne vide
            for var in remaining_vars:
                content = content.replace(f"{{{{{var}}}}}", "")
        
        # Rechercher les variables non remplacées au format {variable}
        remaining_simple_vars = re.findall(r'{([^{}]*)}', content)
        if remaining_simple_vars:
            logger.warning(f"Variables non remplacées (format {variable}): {remaining_simple_vars}")
            # Remplacer les variables non trouvées par une chaîne vide
            for var in remaining_simple_vars:
                content = content.replace(f"{{{var}}}", "")
        
        return content
    
    def clean_filename(self, name):
        """
        Nettoie un nom pour qu'il soit utilisable dans un nom de fichier
        
        Args:
            name: Nom à nettoyer
            
        Returns:
            str: Nom nettoyé
        """
        # Supprimer les caractères spéciaux et remplacer les espaces par des underscores
        name = re.sub(r'[\\/*?:"<>|]', '', name)
        name = name.replace(' ', '_')
        name = name.replace('/', '_')
        name = name.replace('\\', '_')
        name = name.strip('.')  # Supprime les points en début et fin
        
        return name
    
    def generate_document(self, file_path, template, client, company_info, variables, format_type=None):
        """
        Génère un document à partir d'un modèle avec gestion d'erreur robuste
        
        Args:
            file_path: Chemin du fichier à créer
            template: Modèle de document (dict)
            client: Informations du client (dict)
            company_info: Informations de l'entreprise (dict)
            variables: Variables spécifiques pour le document
            format_type: Format du document (pdf, docx ou txt)
            
        Returns:
            bool: True si le document a été généré avec succès, False sinon
        """
        try:
            # Déterminer le format si non spécifié
            if format_type is None:
                format_type = "pdf"
                if self.config:
                    format_type = self.config.get("document.default_format", "pdf")
            
            format_type = format_type.lower()
            
            # Vérifier la disponibilité des modules nécessaires
            if format_type == "pdf" and not REPORTLAB_AVAILABLE:
                logger.warning("ReportLab n'est pas installé, utilisation du format TXT à la place")
                format_type = "txt"
            elif format_type == "docx" and not DOCX_AVAILABLE:
                logger.warning("python-docx n'est pas installé, utilisation du format TXT à la place")
                format_type = "txt"
            
            # S'assurer que l'extension correspond au format
            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext[1:] != format_type:
                file_path = os.path.splitext(file_path)[0] + f".{format_type}"
            
            # S'assurer que le répertoire parent existe
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Préparer les variables
            all_variables = {}
            
            # Variables standard du client
            all_variables.update({
                "client_name": client.get("name", ""),
                "client_company": client.get("company", ""),
                "client_email": client.get("email", ""),
                "client_phone": client.get("phone", ""),
                "client_address": client.get("address", "")
            })
            
            # Variables de l'entreprise
            all_variables.update({
                "company_name": company_info.get("name", ""),
                "company_address": company_info.get("address", ""),
                "company_email": company_info.get("email", ""),
                "company_phone": company_info.get("phone", ""),
                "company_website": company_info.get("website", "")
            })
            
            # Variables spécifiques
            all_variables.update(variables)
            
            # Date actuelle
            all_variables["date"] = datetime.now().strftime("%Y-%m-%d")
            
            # Obtenir le contenu du modèle
            content = template.get("content", "")
            
            # Journaliser la taille du contenu pour le débogage
            content_sample = content[:100] + "..." if len(content) > 100 else content
            logger.info(f"Contenu du modèle (échantillon): {content_sample}")
            logger.info(f"Taille du contenu du modèle: {len(content)} caractères")
            logger.info(f"Le contenu contient des balises img: {'<img' in content}")
            
            # Si le modèle a un chemin de fichier, l'utiliser
            if "file_path" in template and template["file_path"]:
                try:
                    template_file = template["file_path"]
                    if not os.path.isabs(template_file):
                        # Construire le chemin absolu si le chemin est relatif
                        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                        template_file = os.path.join(base_dir, "data", "templates", template_file)
                    
                    if os.path.exists(template_file):
                        with open(template_file, 'r', encoding='utf-8') as f:
                            content = f.read()
                except Exception as e:
                    logger.error(f"Erreur lors de la lecture du fichier modèle: {e}")
                    # Continuer avec le contenu stocké dans la base de données
            
            # Remplacer les variables dans le contenu
            content = self.replace_variables(content, all_variables)
            
            # Vérifier si le contenu après remplacement des variables contient toujours des balises d'image
            logger.info(f"Après remplacement des variables, contient des balises img: {'<img' in content}")
            if '<img' in content:
                # Compter les balises d'image
                img_count = content.count('<img')
                logger.info(f"Nombre de balises <img> dans le contenu: {img_count}")
                # Extraire un exemple de balise d'image
                img_example = content[content.find('<img'):content.find('>', content.find('<img'))+1]
                logger.info(f"Exemple de balise image: {img_example}")
            
            # Titre du document
            title = template.get("name", "Document")
            if "title" in variables:
                title = variables["title"]
            
            # Obtenir le chemin du logo
            logo_path = None
            if self.config:
                logo_path = self.config.get("app.company_logo", "")
                
                # Vérifier que le logo existe
                if logo_path and not os.path.exists(logo_path):
                    logger.warning(f"Logo spécifié mais introuvable: {logo_path}")
                    logo_path = None
            
            # Générer le document selon le format avec gestion d'erreur
            try:
                if format_type == "pdf":
                    success = self.generate_pdf(file_path, content, title, client, company_info, logo_path)
                elif format_type == "docx":
                    success = self.generate_docx(file_path, content, title, client, company_info, logo_path)
                else:
                    success = self.generate_txt(file_path, content, title, client, company_info)
                
                if not success:
                    raise Exception(f"Échec lors de la génération au format {format_type}")
                
                logger.info(f"Document généré avec succès: {file_path}")
                return True
            
            except Exception as format_error:
                logger.error(f"Erreur lors de la génération au format {format_type}: {format_error}")
                
                # Solution de repli: générer un document texte
                try:
                    txt_path = os.path.splitext(file_path)[0] + ".txt"
                    success = self.generate_txt(txt_path, content, title, client, company_info)
                    
                    if success:
                        logger.info(f"Document texte de secours généré: {txt_path}")
                        return True
                    else:
                        raise Exception("Échec de la génération du document texte de secours")
                
                except Exception as fallback_error:
                    logger.error(f"Erreur lors de la génération du document texte de secours: {fallback_error}")
                    
                    # Dernière tentative: créer un document texte minimal
                    try:
                        minimal_path = os.path.splitext(file_path)[0] + ".txt"
                        with open(minimal_path, 'w', encoding='utf-8') as f:
                            f.write(f"Titre: {title}\n")
                            f.write(f"Date: {datetime.now().strftime('%Y-%m-%d')}\n")
                            f.write(f"Client: {client.get('name', '')}\n\n")
                            f.write("Une erreur est survenue lors de la génération du document.\n")
                            f.write("Ce document est une version de secours minimale.\n")
                        
                        logger.info(f"Document texte minimal créé: {minimal_path}")
                        return True
                    
                    except Exception as minimal_error:
                        logger.error(f"Erreur lors de la création du document minimal: {minimal_error}")
                        return False
        
        except Exception as e:
            logger.error(f"Erreur globale lors de la génération du document: {e}")
            
            try:
                # Tentative ultime: créer un fichier texte d'erreur
                error_path = os.path.join(os.path.dirname(file_path), f"error_{datetime.now().strftime('%Y%m%d%H%M%S')}.txt")
                with open(error_path, 'w', encoding='utf-8') as f:
                    f.write(f"ERREUR DE GÉNÉRATION DE DOCUMENT\n")
                    f.write(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    f.write(f"Erreur: {str(e)}\n")
                
                logger.info(f"Fichier d'erreur créé: {error_path}")
                return False
            except:
                return False
    
    def generate_txt(self, file_path, content, title, client, company_info):
        """
        Génère un document texte simple en retirant les balises HTML
        
        Args:
            file_path: Chemin du fichier à créer
            content: Contenu du document (peut contenir du HTML)
            title: Titre du document
            client: Informations du client
            company_info: Informations de l'entreprise
            
        Returns:
            bool: True si réussi, False sinon
        """
        try:
            # Créer le répertoire parent si nécessaire
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Nettoyer le contenu HTML
            clean_content = self.strip_html(content)[0]
            
            with open(file_path, 'w', encoding='utf-8') as f:
                # En-tête
                f.write(f"Titre: {title}\n\n")
                
                # Contenu
                f.write("=== CONTENU ===\n")
                f.write(clean_content)
            
            return True
        
        except Exception as e:
            logger.error(f"Erreur lors de la génération du fichier texte: {e}")
            return False
    
    def generate_pdf(self, file_path, content, title, client, company_info, logo_path=None):
        """
        Génère un document PDF en retirant les balises HTML
        
        Args:
            file_path: Chemin du fichier à créer
            content: Contenu du document (peut contenir du HTML)
            title: Titre du document
            client: Informations du client
            company_info: Informations de l'entreprise
            logo_path: Chemin du logo de l'entreprise
            
        Returns:
            bool: True si réussi, False sinon
        """
        try:
            if not REPORTLAB_AVAILABLE:
                logger.error("ReportLab n'est pas installé, impossible de générer un PDF")
                return False
            
            # Créer le répertoire parent si nécessaire
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Nettoyer le contenu HTML et extraire les images
            clean_content, images = self.strip_html(content)
            logger.info(f"Génération PDF: {len(images)} images extraites")
            
            # Créer un PDF
            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4
            
            # Définir la police à utiliser - Helvetica par défaut
            font_name = "Helvetica"
            
            # Essayer d'utiliser une police qui supporte mieux l'Unicode si disponible
            try:
                # Vérifier l'existence de DejaVuSans.ttf
                dejavu_path = None
                
                # Chemins possibles pour DejaVuSans.ttf
                possible_paths = [
                    os.path.join(os.path.dirname(__file__), 'fonts', 'DejaVuSans.ttf'),  # Dans le dossier fonts/ relatif à ce script
                    os.path.join(os.path.dirname(os.path.dirname(__file__)), 'fonts', 'DejaVuSans.ttf'),  # Dans le dossier fonts/ du projet
                    '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',  # Chemin Linux commun
                    'C:\\Windows\\Fonts\\DejaVuSans.ttf',  # Chemin Windows
                    os.path.expanduser('~/Library/Fonts/DejaVuSans.ttf')  # Chemin macOS
                ]
                
                for path in possible_paths:
                    if os.path.exists(path):
                        dejavu_path = path
                        break
                
                if dejavu_path:
                    pdfmetrics.registerFont(TTFont('DejaVuSans', dejavu_path))
                    c.setFont("DejaVuSans", 10)
                    font_name = "DejaVuSans"
                    logger.info(f"Utilisation de la police DejaVuSans pour le PDF")
            except Exception as font_error:
                logger.warning(f"Impossible d'utiliser la police DejaVuSans, utilisation de Helvetica: {font_error}")
            
            # Ajouter le logo si fourni
            if logo_path and os.path.exists(logo_path):
                try:
                    # Tenter de charger le logo
                    logo_img = PILImage.open(logo_path)
                    logo_width, logo_height = logo_img.size
                    # Redimensionner si trop grand
                    max_logo_width = 150  # Taille maximale en points
                    max_logo_height = 70
                    if logo_width > max_logo_width or logo_height > max_logo_height:
                        ratio = min(max_logo_width / logo_width, max_logo_height / logo_height)
                        logo_width, logo_height = int(logo_width * ratio), int(logo_height * ratio)
                    
                    # Positionner le logo en haut à droite
                    c.drawInlineImage(logo_path, width - logo_width - 50, height - 50 - logo_height, logo_width, logo_height)
                except Exception as logo_error:
                    logger.warning(f"Impossible d'ajouter le logo: {logo_error}")
            
            # Définir les marges et la largeur disponible
            margin_left = 50
            margin_right = 50
            available_width = width - margin_left - margin_right
            
            # Titre du document - avec gestion de titre long
            c.setFont(font_name + "-Bold" if font_name == "Helvetica" else font_name, 16)
            title = self._sanitize_text(title)
            
            # Vérifier si le titre est trop long
            if c.stringWidth(title, font_name + "-Bold" if font_name == "Helvetica" else font_name, 16) > available_width:
                # Découper le titre en plusieurs lignes
                words = title.split(' ')
                current_line = ""
                y_position = height - 70
                
                for word in words:
                    test_line = current_line + " " + word if current_line else word
                    if c.stringWidth(test_line, font_name + "-Bold" if font_name == "Helvetica" else font_name, 16) <= available_width:
                        current_line = test_line
                    else:
                        # Dessiner la partie actuelle du titre
                        c.drawString(margin_left, y_position, current_line)
                        y_position -= 20  # Espacement pour le titre
                        current_line = word
                
                # Dessiner la dernière partie du titre
                if current_line:
                    c.drawString(margin_left, y_position, current_line)
                    y_position -= 20
                
                # Ligne séparatrice après la dernière ligne du titre
                c.line(50, y_position - 5, width - 50, y_position - 5)
                y_position -= 20  # Espacement après la ligne
            else:
                # Titre normal sur une seule ligne
                c.drawString(margin_left, height - 70, title)
                
                # Ligne séparatrice
                c.line(50, height - 85, width - 50, height - 85)
                
                y_position = height - 105
            
            # Contenu du document
            # Diviser le contenu en lignes
            lines = clean_content.split('\n')
            c.setFont(font_name, 10)
            
            font_size = 10
            image_index = 0  # Pour suivre quelle image afficher
            
            # Journaliser les lignes qui contiennent [IMAGE] pour débogage
            image_lines = [i for i, line in enumerate(lines) if line.strip() == '[IMAGE]']
            logger.info(f"Lignes avec marqueur [IMAGE]: {image_lines}")
            
            for line in lines:
                # Vérifier si c'est un marqueur d'image
                if line.strip() == '[IMAGE]' and image_index < len(images):
                    logger.info(f"Traitement de l'image {image_index+1}/{len(images)}")
                    # Tenter d'ajouter l'image
                    try:
                        img_src = images[image_index]['src']
                        logger.info(f"Source de l'image: {img_src[:50]}...")
                        # Récupérer les dimensions si présentes
                        img_width_attr = images[image_index].get('width')
                        img_height_attr = images[image_index].get('height')
                        image_index += 1
                        
                        # Vérifier si l'image est une URL ou un chemin local
                        is_url = img_src.startswith(('http://', 'https://', 'ftp://'))
                        is_temp = False  # Pour suivre si le fichier est temporaire (base64)
                        
                        # Chemin où télécharger/copier l'image si nécessaire
                        temp_img_path = None
                        
                        if is_url:
                            # Télécharger l'image si c'est une URL
                            import tempfile
                            import urllib.request
                            
                            try:
                                # Créer un fichier temporaire
                                temp_fd, temp_img_path = tempfile.mkstemp(suffix='.jpg')
                                os.close(temp_fd)
                                
                                # Télécharger l'image
                                urllib.request.urlretrieve(img_src, temp_img_path)
                                logger.info(f"Image téléchargée de {img_src} vers {temp_img_path}")
                                is_temp = True  # Marquer comme fichier temporaire
                            except Exception as url_error:
                                logger.warning(f"Erreur lors du téléchargement de l'image {img_src}: {url_error}")
                                temp_img_path = None
                        else:
                            # Vérifier si c'est déjà un fichier temporaire (cas des images base64)
                            if os.path.exists(img_src) and 'temp' in img_src:
                                temp_img_path = img_src
                                is_temp = True  # Marquer comme fichier temporaire
                            else:
                                # Chemin local
                                temp_img_path = img_src if os.path.isabs(img_src) else os.path.join(os.path.dirname(file_path), img_src)
                        
                        # Ajouter l'image si disponible
                        if temp_img_path and os.path.exists(temp_img_path):
                            try:
                                img = PILImage.open(temp_img_path)
                                img_width, img_height = img.size
                                
                                # Utiliser les dimensions spécifiées si disponibles
                                if img_width_attr and img_height_attr:
                                    try:
                                        specified_width = int(img_width_attr)
                                        specified_height = int(img_height_attr)
                                        # Vérifier que les dimensions sont raisonnables
                                        if 0 < specified_width <= available_width and 0 < specified_height <= 500:
                                            img_width, img_height = specified_width, specified_height
                                    except (ValueError, TypeError):
                                        # Ignorer si les dimensions ne sont pas valides
                                        pass
                                
                                # Redimensionner si trop grande
                                max_img_width = available_width
                                max_img_height = 300  # Points
                                
                                if img_width > max_img_width or img_height > max_img_height:
                                    ratio = min(max_img_width / img_width, max_img_height / img_height)
                                    img_width, img_height = int(img_width * ratio), int(img_height * ratio)
                                
                                # Vérifier s'il reste assez d'espace sur la page
                                if y_position - img_height < 70:
                                    c.showPage()
                                    c.setFont(font_name, font_size)
                                    y_position = height - 70
                                
                                # Centrer l'image
                                x_pos = margin_left + (available_width - img_width) / 2
                                c.drawInlineImage(temp_img_path, x_pos, y_position - img_height, img_width, img_height)
                                
                                # Déplacer le curseur y
                                y_position -= (img_height + 15)  # 15 points d'espace supplémentaire
                                
                                # Supprimer le fichier temporaire si c'en est un
                                if is_temp and temp_img_path:
                                    try:
                                        os.remove(temp_img_path)
                                    except:
                                        pass
                            except Exception as img_error:
                                logger.warning(f"Erreur lors de l'ajout de l'image {temp_img_path}: {img_error}")
                    except Exception as img_process_error:
                        logger.warning(f"Erreur lors du traitement de l'image: {img_process_error}")
                else:
                    # Sanitize line text
                    line = self._sanitize_text(line)
                    
                    # Vérifier si la ligne est trop longue 
                    if c.stringWidth(line, font_name, font_size) > available_width:
                        # Découpage de la ligne en plusieurs lignes
                        words = line.split(' ')
                        current_line = ""
                        
                        for word in words:
                            test_line = current_line + " " + word if current_line else word
                            if c.stringWidth(test_line, font_name, font_size) <= available_width:
                                current_line = test_line
                            else:
                                # Vérifier s'il reste assez d'espace sur la page
                                if y_position < 70:  # Marge de bas de page augmentée
                                    c.showPage()  # Nouvelle page
                                    c.setFont(font_name, 10)
                                    y_position = height - 70  # Marge supérieure augmentée
                                
                                # Écrire la ligne actuelle
                                c.drawString(margin_left, y_position, current_line)
                                y_position -= 15  # Espacement des lignes
                                current_line = word
                        
                        # Écrire la dernière partie de la ligne
                        if current_line:
                            # Vérifier s'il reste assez d'espace sur la page
                            if y_position < 70:
                                c.showPage()
                                c.setFont(font_name, 10)
                                y_position = height - 70
                            
                            c.drawString(margin_left, y_position, current_line)
                            y_position -= 15
                    else:
                        # Vérifier s'il reste assez d'espace sur la page
                        if y_position < 70:  # Marge de bas de page augmentée
                            c.showPage()  # Nouvelle page
                            c.setFont(font_name, 10)
                            y_position = height - 70  # Marge supérieure augmentée
                        
                        # Dessiner la ligne
                        c.drawString(margin_left, y_position, line)
                        y_position -= 15  # Espacement des lignes
            
            # Finaliser le document
            c.showPage()
            c.save()
            
            return True
            
        except Exception as e:
            logger.error(f"Erreur lors de la génération du PDF: {e}")
            return False
    
    def _sanitize_text(self, text):
        """
        Nettoie le texte pour éviter les problèmes avec les caractères spéciaux dans les PDF
        
        Args:
            text: Texte à nettoyer
            
        Returns:
            str: Texte nettoyé
        """
        if not isinstance(text, str):
            text = str(text)
        
        # Remplacer les caractères problématiques par leurs équivalents compatibles
        special_chars_map = {
            '…': '...',
            '–': '-',
            '—': '-',
            ''': "'",
            ''': "'",
            '"': '"',
            '"': '"',
            '«': '"',
            '»': '"',
            '•': '*'
        }
        
        for char, replacement in special_chars_map.items():
            text = text.replace(char, replacement)
        
        # Supprimer les caractères non imprimables ou non supportés
        text = ''.join(c for c in text if ord(c) >= 32 or c in '\n\r\t')
        
        return text
    
    def generate_docx(self, file_path, content, title, client, company_info, logo_path=None):
        """
        Génère un document DOCX en retirant les balises HTML
        
        Args:
            file_path: Chemin du fichier à créer
            content: Contenu du document (peut contenir du HTML)
            title: Titre du document
            client: Informations du client
            company_info: Informations de l'entreprise
            logo_path: Chemin du logo de l'entreprise
            
        Returns:
            bool: True si réussi, False sinon
        """
        try:
            if not DOCX_AVAILABLE:
                logger.error("python-docx n'est pas installé, impossible de générer un DOCX")
                return False
            
            # Créer le répertoire parent si nécessaire
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Nettoyer le contenu HTML et extraire les images
            clean_content, images = self.strip_html(content)
            logger.info(f"Génération DOCX: {len(images)} images extraites")
            
            # Créer un document Word
            doc = docx.Document()
            
            # Définir la largeur maximale des marges
            for section in doc.sections:
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)
                section.top_margin = Inches(1.2)
                section.bottom_margin = Inches(1.2)
            
            # Ajouter le logo si fourni
            if logo_path and os.path.exists(logo_path):
                try:
                    # Ajouter un paragraphe pour le logo
                    logo_para = doc.add_paragraph()
                    logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    logo_run = logo_para.add_run()
                    logo_run.add_picture(logo_path, width=Inches(2.0))
                except Exception as logo_error:
                    logger.warning(f"Impossible d'ajouter le logo: {logo_error}")
            
            # Ajouter une police Unicode compatible à l'ensemble du document
            doc_style = doc.styles['Normal']
            doc_style.font.name = 'Arial Unicode MS'
            
            # Titre du document
            title_para = doc.add_heading(title, level=1)
            
            # Ligne séparatrice
            doc.add_paragraph("_______________________________________________________________")
            
            # Contenu
            # Diviser le contenu en paragraphes
            paragraphs = clean_content.split('\n\n')
            image_index = 0  # Pour suivre quelle image afficher
            
            # Journaliser les paragraphes qui contiennent [IMAGE] pour débogage
            image_paragraphs = [i for i, para in enumerate(paragraphs) if para.strip() == '[IMAGE]']
            logger.info(f"Paragraphes avec marqueur [IMAGE]: {image_paragraphs}")
            
            for paragraph_text in paragraphs:
                if paragraph_text.strip() == '[IMAGE]' and image_index < len(images):
                    logger.info(f"Traitement de l'image DOCX {image_index+1}/{len(images)}")
                    # Tenter d'ajouter l'image
                    try:
                        img_src = images[image_index]['src']
                        logger.info(f"Source de l'image DOCX: {img_src[:50]}...")
                        # Récupérer les dimensions si présentes
                        img_width_attr = images[image_index].get('width')
                        img_height_attr = images[image_index].get('height')
                        image_index += 1
                        
                        # Vérifier si l'image est une URL ou un chemin local
                        is_url = img_src.startswith(('http://', 'https://', 'ftp://'))
                        is_temp = False  # Pour suivre si le fichier est temporaire (base64)
                        
                        # Chemin où télécharger/copier l'image si nécessaire
                        temp_img_path = None
                        
                        if is_url:
                            # Télécharger l'image si c'est une URL
                            import tempfile
                            import urllib.request
                            
                            try:
                                # Créer un fichier temporaire
                                temp_fd, temp_img_path = tempfile.mkstemp(suffix='.jpg')
                                os.close(temp_fd)
                                
                                # Télécharger l'image
                                urllib.request.urlretrieve(img_src, temp_img_path)
                                logger.info(f"Image téléchargée de {img_src} vers {temp_img_path}")
                                is_temp = True  # Marquer comme fichier temporaire
                            except Exception as url_error:
                                logger.warning(f"Erreur lors du téléchargement de l'image {img_src}: {url_error}")
                                temp_img_path = None
                        else:
                            # Vérifier si c'est déjà un fichier temporaire (cas des images base64)
                            if os.path.exists(img_src) and 'temp' in img_src:
                                temp_img_path = img_src
                                is_temp = True  # Marquer comme fichier temporaire
                            else:
                                # Chemin local
                                temp_img_path = img_src if os.path.isabs(img_src) else os.path.join(os.path.dirname(file_path), img_src)
                        
                        # Ajouter l'image si disponible
                        if temp_img_path and os.path.exists(temp_img_path):
                            img_para = doc.add_paragraph()
                            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            img_run = img_para.add_run()
                            
                            # Calculer la largeur en pouces
                            img_width_inches = 5.0  # Taille par défaut (5 pouces)
                            
                            # Utiliser les dimensions spécifiées si disponibles
                            if img_width_attr and img_height_attr:
                                try:
                                    specified_width = int(img_width_attr)
                                    max_width_pixels = 500  # Largeur maximale en pixels
                                    
                                    # Limiter la largeur à une valeur raisonnable
                                    if 0 < specified_width <= max_width_pixels:
                                        # Convertir en pouces (approximativement)
                                        img_width_inches = min(specified_width / 96, 6.0)  # max 6 pouces, 96 DPI
                                except (ValueError, TypeError):
                                    # Utiliser la taille par défaut
                                    pass
                            
                            # Ajouter l'image avec la largeur calculée
                            img_run.add_picture(temp_img_path, width=Inches(img_width_inches))
                            
                            # Supprimer le fichier temporaire si c'en est un
                            if is_temp and temp_img_path:
                                try:
                                    os.remove(temp_img_path)
                                except:
                                    pass
                    except Exception as img_error:
                        logger.warning(f"Erreur lors de l'ajout de l'image: {img_error}")
                elif paragraph_text.strip():
                    # Pour Word, nous n'avons pas besoin de sanitizer le texte comme pour PDF
                    # car DOCX gère bien les caractères Unicode
                    p = doc.add_paragraph(paragraph_text)
                    # Limiter la largeur des paragraphes en ajustant les retraits
                    p.paragraph_format.first_line_indent = Inches(0)
                    p.paragraph_format.left_indent = Inches(0)
                    p.paragraph_format.right_indent = Inches(0)
                    
                    # S'assurer que la police prend en charge l'Unicode
                    for run in p.runs:
                        run.font.name = 'Arial Unicode MS'
            
            # Enregistrer le document
            doc.save(file_path)
            
            return True
            
        except Exception as e:
            logger.error(f"Erreur lors de la génération du DOCX: {e}")
            return False